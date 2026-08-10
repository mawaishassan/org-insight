import re
import io
import zipfile

def append_to_service():
    with open('app/reports/custom_service.py', 'r') as f:
        content = f.read()

    new_code = """
async def export_custom_report_attachments(
    db, custom_report_id: int, org_id: int, year: int, format: str, attachment_ids: list[int]
) -> tuple[bytes, str, str]:
    import io
    import zipfile
    import datetime
    import re
    from sqlalchemy import select
    from app.core.models import CustomReport, KPI, KPIEntry, KpiMultiLineRow, KpiMultiLineCell, KPIFieldSubField
    from sqlalchemy.orm import selectinload, noload

    report = await get_custom_report(db, custom_report_id, org_id)
    if not report:
        raise ValueError("Report not found")
        
    attachments = [a for a in (report.attachments or []) if a.id in attachment_ids]
    if not attachments:
        raise ValueError("No matching attachments found")

    # Generate file for each attachment
    files = [] # list of (filename, bytes)
    yr = year if year is not None else datetime.date.today().year

    for att in attachments:
        # Load entries for this KPI
        kfield = att.kpi_field
        kpi = att.kpi
        
        entries_res = await db.execute(
            select(KPIEntry)
            .where(
                KPIEntry.organization_id == org_id,
                KPIEntry.kpi_id == kpi.id,
                KPIEntry.year == yr,
                KPIEntry.is_draft == False
            )
        )
        entries = entries_res.scalars().all()
        if not entries:
            continue
            
        entry = entries[-1] # Simplification, typically we sort by period_key

        rows_stmt = (
            select(KpiMultiLineRow.id, KpiMultiLineRow.row_index)
            .where(
                KpiMultiLineRow.entry_id == entry.id,
                KpiMultiLineRow.field_id == kfield.id,
            )
            .order_by(KpiMultiLineRow.row_index)
        )
        rows_list = (await db.execute(rows_stmt)).all()
        
        chunk_rows = []
        if rows_list:
            row_ids = [r[0] for r in rows_list]
            cells_res = await db.execute(
                select(
                    KpiMultiLineCell.row_id, KpiMultiLineCell.value_text, KpiMultiLineCell.value_number,
                    KpiMultiLineCell.value_boolean, KpiMultiLineCell.value_date, KpiMultiLineCell.value_json,
                    KPIFieldSubField.key, KPIFieldSubField.name
                )
                .join(KPIFieldSubField, KPIFieldSubField.id == KpiMultiLineCell.sub_field_id)
                .where(KpiMultiLineCell.row_id.in_(row_ids))
            )
            cells_list = cells_res.all()
            
            cells_by_row = {}
            for row_id, vt, vn, vb, vd, vj, sf_key, sf_name in cells_list:
                if row_id not in cells_by_row:
                    cells_by_row[row_id] = {}
                raw_val = vj if vj is not None else vt if vt is not None else vn if vn is not None else vb if vb is not None else (vd.isoformat() if vd else None)
                cells_by_row[row_id][sf_key] = raw_val

            for rid, r_idx in rows_list:
                chunk_rows.append(cells_by_row.get(rid, {}))

        # Build Document
        sub_fields = [{"key": sf.key, "name": sf.name or sf.key} for sf in getattr(kfield, "sub_fields", [])]
        clean_title = re.sub(r'[^\w\s-]', '', att.title).strip().replace(' ', '_')
        
        if format == "xlsx":
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = clean_title[:30]
            
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
            
            # Headers
            for col_idx, sf in enumerate(sub_fields):
                c = ws.cell(row=1, column=col_idx+1, value=sf["name"])
                c.font = header_font
                c.fill = header_fill
            
            for r_idx, item in enumerate(chunk_rows):
                for col_idx, sf in enumerate(sub_fields):
                    ws.cell(row=r_idx+2, column=col_idx+1, value=str(item.get(sf["key"], "—")))
            
            out_io = io.BytesIO()
            wb.save(out_io)
            files.append((f"{clean_title}.xlsx", out_io.getvalue()))
            
        elif format == "pdf":
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            
            out_io = io.BytesIO()
            pdf_doc = SimpleDocTemplate(out_io, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            story.append(Paragraph(att.title, styles["Heading1"]))
            story.append(Spacer(1, 12))
            
            if sub_fields and chunk_rows:
                table_data = [[Paragraph(sf["name"], styles["Normal"]) for sf in sub_fields]]
                for item in chunk_rows:
                    row = [Paragraph(str(item.get(sf["key"], "—")), styles["Normal"]) for sf in sub_fields]
                    table_data.append(row)
                
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey)
                ]))
                story.append(t)
            else:
                story.append(Paragraph("No data available.", styles["Normal"]))
                
            pdf_doc.build(story)
            files.append((f"{clean_title}.pdf", out_io.getvalue()))
            
        elif format == "docx":
            import docx
            doc = docx.Document()
            doc.add_heading(att.title, level=1)
            
            if sub_fields and chunk_rows:
                table = doc.add_table(rows=1, cols=len(sub_fields))
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                for col_idx, sf in enumerate(sub_fields):
                    hdr_cells[col_idx].text = sf["name"]
                
                for item in chunk_rows:
                    row_cells = table.add_row().cells
                    for col_idx, sf in enumerate(sub_fields):
                        row_cells[col_idx].text = str(item.get(sf["key"], "—"))
            else:
                doc.add_paragraph("No data available.")
                
            out_io = io.BytesIO()
            doc.save(out_io)
            files.append((f"{clean_title}.docx", out_io.getvalue()))

    if len(files) == 1:
        fname, fbytes = files[0]
        content_type = "application/pdf" if format == "pdf" else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" if format == "xlsx" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return fbytes, fname, content_type
    elif len(files) > 1:
        zip_io = io.BytesIO()
        with zipfile.ZipFile(zip_io, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, fbytes in files:
                zf.writestr(fname, fbytes)
        return zip_io.getvalue(), "Attachments.zip", "application/zip"
    else:
        raise ValueError("No attachments could be generated")
"""
    if "async def export_custom_report_attachments" not in content:
        with open('app/reports/custom_service.py', 'a') as f:
            f.write(new_code)

append_to_service()
