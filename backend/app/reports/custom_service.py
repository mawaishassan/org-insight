import datetime
import html
import logging
from collections import defaultdict
from typing import Callable, Any
from sqlalchemy import select, delete, func
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload, noload

from app.core.models import (
    CustomReport,
    CustomReportSection,
    CustomReportField,
    CustomReportAssignment,
    CustomReportAttachment,
    KPI,
    KPIField,
    KPIEntry,
    Organization,
    User,
    FieldType,
    KpiMultiLineRow,
    KpiMultiLineCell,
    KPIFieldSubField
)
from app.reports.custom_schemas import (
    CustomReportCreate,
    CustomReportUpdate,
    CustomReportSectionLayout,
    CustomReportAttachmentLayout
)
from app.formula_engine.evaluator import evaluate_formula, apply_conditional_logic


from decimal import Decimal

import re
_ILLEGAL_CHARACTERS_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')

def clean_excel_value(val: Any) -> Any:
    if isinstance(val, str):
        return _ILLEGAL_CHARACTERS_RE.sub("", val)
    return val


def clean_numeric_value_string(val) -> str:
    if val is None:
        return "—"
    if isinstance(val, (int, float, Decimal)):
        if float(val).is_integer():
            return str(int(float(val)))
        return str(val)
    if isinstance(val, str):
        try:
            if "." in val:
                f_val = float(val)
                if f_val.is_integer():
                    return str(int(f_val))
        except ValueError:
            pass
    return str(val)


def _parse_year_int(year_val: float | int | str | None) -> int:
    if year_val is None:
        return datetime.date.today().year
    try:
        return int(year_val)
    except (ValueError, TypeError):
        import re
        match = re.search(r'\d{4}', str(year_val))
        return int(match.group(0)) if match else datetime.date.today().year


def calc_auto_header_font_size(text: str, available_w: float = 390.0, desired_size: float = 16.0) -> float:
    if not text:
        return float(desired_size)
    length = len(text.strip())
    if length <= 0:
        return float(desired_size)
    target_w = max(available_w - 8.0, 180.0)
    calc_size = target_w / (length * 0.44)
    max_limit = max(float(desired_size), 18.5)
    return round(max(12.0, min(calc_size, max_limit)), 1)

from app.reports.service import (
    _load_multi_line_items_rows_batch,
    _formulas_need_other_kpi_values,
    _load_other_kpi_values,
    _report_period_display,
    build_reference_resolution_map,
    _multi_raw_pieces,
    _normalize_reference_value,
    _extract_ref_label,
    period_key_sort_order,
    effective_kpi_time_dimension,
)
from app.entries.service import bulk_load_org_kpi_values

logger = logging.getLogger(__name__)


def _row_signature(r: dict) -> tuple:
    # 1. If a unique ID is present, use it as the signature
    for k in ['id', 'patent_lms_id', 'lms_id']:
        if k in r and r[k] is not None and r[k] != "":
            val = str(r[k])
            if val.endswith(".0"):
                val = val[:-2]
            return (("id", val),)
            
    # 2. If no ID is present, fall back to Title/Name + Date
    sig_keys = []
    for k in ['invention_title', 'title', 'name', 'project_title', 'book_title', 'paper_title']:
        if k in r and r[k] is not None and r[k] != "":
            sig_keys.append((k, str(r[k]).strip().lower()))
    for k in ['date_of_filing', 'date', 'filing_date', 'publication_date']:
        if k in r and r[k] is not None and r[k] != "":
            sig_keys.append((k, str(r[k]).strip()[:10]))
            
    if sig_keys:
        return tuple(sorted(sig_keys))
    else:
        non_empty = [(k, str(v).strip()) for k, v in r.items() if v is not None and str(v).strip() != ""]
        return tuple(sorted(non_empty))


def _deduplicate_rows(rows: list[dict]) -> list[dict]:
    by_sig = {}
    for r in rows:
        sig = _row_signature(r)
        if sig not in by_sig:
            by_sig[sig] = []
        by_sig[sig].append(r)
        
    deduped = []
    for sig, group in by_sig.items():
        best_row = max(group, key=lambda r: len([v for v in r.values() if v is not None and v != ""]))
        deduped.append(best_row)
    return deduped


# Standard TimeDimension import or helper definition
class TimeDimension:
    YEARLY = "yearly"
    HALF_YEARLY = "half_yearly"
    QUARTERLY = "quarterly"
    MONTHLY = "monthly"

    def __init__(self, val):
        self.val = str(val or "yearly").strip().lower()


async def create_custom_report(db: AsyncSession, org_id: int, data: CustomReportCreate) -> CustomReport:
    report = CustomReport(
        organization_id=org_id,
        group_id=data.group_id,
        name=data.name,
        description=data.description,
        report_header_id=data.report_header_id,
        show_report_name=data.show_report_name,
        branding_title=data.branding_title,
        show_odoo_button=data.show_odoo_button,
        odoo_sync_kpi_ids=data.odoo_sync_kpi_ids,
        apply_further_processing_based_on_mli_filter=data.apply_further_processing_based_on_mli_filter,
        scalar_font_family=data.scalar_font_family,
        mli_font_family=data.mli_font_family,
    )
    db.add(report)
    await db.flush()
    return report


async def get_custom_report(db: AsyncSession, id: int, org_id: int) -> CustomReport | None:
    result = await db.execute(
        select(CustomReport)
        .where(CustomReport.id == id, CustomReport.organization_id == org_id)
        .options(
            selectinload(CustomReport.sections).selectinload(CustomReportSection.fields).selectinload(CustomReportField.kpi_field).selectinload(KPIField.sub_fields),
            selectinload(CustomReport.sections).selectinload(CustomReportSection.kpi).options(
                noload(KPI.organization),
                noload(KPI.domain)
            ),
            selectinload(CustomReport.attachments).selectinload(CustomReportAttachment.kpi_field).selectinload(KPIField.sub_fields),
            selectinload(CustomReport.attachments).selectinload(CustomReportAttachment.kpi),
        )
    )
    return result.scalar_one_or_none()


async def list_custom_reports(db: AsyncSession, org_id: int) -> list[CustomReport]:
    result = await db.execute(
        select(CustomReport)
        .where(CustomReport.organization_id == org_id)
        .order_by(CustomReport.name)
    )
    return list(result.scalars().all())


async def update_custom_report(db: AsyncSession, id: int, org_id: int, data: CustomReportUpdate) -> CustomReport | None:
    report = await get_custom_report(db, id, org_id)
    if not report:
        return None
    report.name = data.name
    report.description = data.description
    if "group_id" in data.model_fields_set:
        report.group_id = data.group_id
    if getattr(data, "fetch_data_with_date", None) is not None:
        report.fetch_data_with_date = data.fetch_data_with_date
    if getattr(data, "date_fetching_config", None) is not None:
        report.date_fetching_config = data.date_fetching_config
    if getattr(data, "report_header_id", None) is not None:
        report.report_header_id = data.report_header_id
    if getattr(data, "show_report_name", None) is not None:
        report.show_report_name = data.show_report_name
    if getattr(data, "branding_title", None) is not None:
        report.branding_title = data.branding_title
    if getattr(data, "show_odoo_button", None) is not None:
        report.show_odoo_button = data.show_odoo_button
    if getattr(data, "odoo_sync_kpi_ids", None) is not None:
        report.odoo_sync_kpi_ids = data.odoo_sync_kpi_ids
    if getattr(data, "apply_further_processing_based_on_mli_filter", None) is not None:
        report.apply_further_processing_based_on_mli_filter = data.apply_further_processing_based_on_mli_filter

    await db.flush()
    CUSTOM_REPORT_CACHE.invalidate_report(id)
    return report


async def delete_custom_report(db: AsyncSession, id: int, org_id: int) -> bool:
    report = await get_custom_report(db, id, org_id)
    if not report:
        return False
    await db.delete(report)
    await db.flush()
    CUSTOM_REPORT_CACHE.invalidate_report(id)
    return True


async def duplicate_custom_report(db: AsyncSession, id: int, org_id: int) -> CustomReport | None:
    orig = await get_custom_report(db, id, org_id)
    if not orig:
        return None

    # Create new CustomReport
    new_report = CustomReport(
        organization_id=org_id,
        name=f"Copy of {orig.name}",
        description=orig.description,
        report_header_id=orig.report_header_id,
        show_report_name=orig.show_report_name,
        branding_title=orig.branding_title,
        show_odoo_button=orig.show_odoo_button,
        odoo_sync_kpi_ids=orig.odoo_sync_kpi_ids,
        apply_further_processing_based_on_mli_filter=getattr(orig, "apply_further_processing_based_on_mli_filter", False),
        scalar_font_family=getattr(orig, "scalar_font_family", "Inter"),
        mli_font_family=getattr(orig, "mli_font_family", "Inter"),
    )
    db.add(new_report)
    await db.flush()

    # Copy sections and fields
    for sec in orig.sections:
        new_sec = CustomReportSection(
            custom_report_id=new_report.id,
            kpi_id=sec.kpi_id,
            custom_header=sec.custom_header,
            sort_order=sec.sort_order
        )
        db.add(new_sec)
        await db.flush()

        for f in sec.fields:
            new_field = CustomReportField(
                custom_report_id=new_report.id,
                custom_report_section_id=new_sec.id,
                kpi_field_id=f.kpi_field_id,
                sort_order=f.sort_order,
                config=f.config
            )
            db.add(new_field)

    await db.flush()
    # Refresh to load relationships
    return await get_custom_report(db, new_report.id, org_id)


async def save_custom_report_layout(
    db: AsyncSession,
    id: int,
    org_id: int,
    sections: list[CustomReportSectionLayout],
    attachments: list[CustomReportAttachmentLayout] | None = None,
    fetch_data_with_date: bool | None = None,
    date_fetching_config: dict | None = None,
    report_header_id: int | None = None,
    show_report_name: bool | None = None,
    branding_title: str | None = None,
    scalar_bold: bool | None = None,
    scalar_font_size: int | None = None,
    mli_font_size: int | None = None,
    show_odoo_button: bool | None = None,
    odoo_sync_kpi_ids: list[int] | None = None,
    apply_further_processing_based_on_mli_filter: bool | None = None,
    scalar_font_family: str | None = None,
    mli_font_family: str | None = None,
) -> bool:
    report = await get_custom_report(db, id, org_id)
    if not report:
        return False

    if fetch_data_with_date is not None:
        report.fetch_data_with_date = fetch_data_with_date
    if date_fetching_config is not None:
        report.date_fetching_config = date_fetching_config
    if report_header_id is not None:
        report.report_header_id = report_header_id
    elif "report_header_id" in locals() or True:
        report.report_header_id = report_header_id
    if show_report_name is not None:
        report.show_report_name = show_report_name
    if branding_title is not None:
        report.branding_title = branding_title
    if scalar_bold is not None:
        report.scalar_bold = scalar_bold
    if scalar_font_size is not None:
        report.scalar_font_size = scalar_font_size
    if mli_font_size is not None:
        report.mli_font_size = mli_font_size
    if scalar_font_family is not None:
        report.scalar_font_family = scalar_font_family
    if mli_font_family is not None:
        report.mli_font_family = mli_font_family
    if show_odoo_button is not None:
        report.show_odoo_button = show_odoo_button
    # odoo_sync_kpi_ids can be set to empty list or None explicitly
    if odoo_sync_kpi_ids is not None:
        report.odoo_sync_kpi_ids = odoo_sync_kpi_ids
    if apply_further_processing_based_on_mli_filter is not None:
        report.apply_further_processing_based_on_mli_filter = apply_further_processing_based_on_mli_filter



    # Delete all existing sections, fields, and attachments
    await db.execute(delete(CustomReportSection).where(CustomReportSection.custom_report_id == id))
    await db.execute(delete(CustomReportField).where(CustomReportField.custom_report_id == id))
    await db.execute(delete(CustomReportAttachment).where(CustomReportAttachment.custom_report_id == id))
    await db.flush()

    # Re-insert the new sections and fields
    for s_idx, sec_data in enumerate(sections):
        sec = CustomReportSection(
            custom_report_id=id,
            kpi_id=sec_data.kpi_id,
            custom_header=sec_data.custom_header,
            sort_order=sec_data.sort_order
        )
        db.add(sec)
        await db.flush()

        for f_idx, field_data in enumerate(sec_data.fields):
            f = CustomReportField(
                custom_report_id=id,
                custom_report_section_id=sec.id,
                kpi_field_id=field_data.kpi_field_id,
                sort_order=field_data.sort_order,
                config=field_data.config
            )
            db.add(f)

    # Re-insert attachments if provided
    if attachments:
        for att_idx, att_data in enumerate(attachments):
            att = CustomReportAttachment(
                custom_report_id=id,
                kpi_id=att_data.kpi_id,
                kpi_field_id=att_data.kpi_field_id,
                title=att_data.title,
                selected_columns=att_data.selected_columns,
                filters=att_data.filters,
                sort_order=att_data.sort_order
            )
            db.add(att)

    await db.flush()
    CUSTOM_REPORT_CACHE.invalidate_report(id)
    return True


async def assign_custom_report(
    db: AsyncSession, custom_report_id: int, user_id: int, can_view: bool, can_print: bool, can_export: bool, can_change_period: bool = True
) -> CustomReportAssignment:
    # Check if assignment already exists
    result = await db.execute(
        select(CustomReportAssignment)
        .where(CustomReportAssignment.custom_report_id == custom_report_id, CustomReportAssignment.user_id == user_id)
    )
    perm = result.scalar_one_or_none()
    if not perm:
        perm = CustomReportAssignment(
            custom_report_id=custom_report_id,
            user_id=user_id
        )
        db.add(perm)

    perm.can_view = can_view
    perm.can_print = can_print
    perm.can_export = can_export
    perm.can_change_period = can_change_period
    await db.flush()
    CUSTOM_REPORT_CACHE.invalidate_report(custom_report_id)
    return perm


async def bulk_assign_custom_report(
    db: AsyncSession,
    custom_report_id: int,
    user_ids: list[int],
    can_view: bool = True,
    can_print: bool = True,
    can_export: bool = True,
    can_change_period: bool = True,
) -> list[CustomReportAssignment]:
    """Assign custom report to multiple users simultaneously inside a single transaction."""
    result = await db.execute(
        select(CustomReportAssignment)
        .where(
            CustomReportAssignment.custom_report_id == custom_report_id,
            CustomReportAssignment.user_id.in_(user_ids),
        )
    )
    existing_list = result.scalars().all()
    existing_by_user_id = {a.user_id: a for a in existing_list}

    out = []
    for uid in user_ids:
        perm = existing_by_user_id.get(uid)
        if not perm:
            perm = CustomReportAssignment(
                custom_report_id=custom_report_id,
                user_id=uid,
                can_change_period=can_change_period,
            )
            db.add(perm)
        perm.can_view = can_view
        perm.can_print = can_print
        perm.can_export = can_export
        perm.can_change_period = can_change_period
        out.append(perm)

    await db.flush()
    CUSTOM_REPORT_CACHE.invalidate_report(custom_report_id)
    return out


async def unassign_custom_report(db: AsyncSession, custom_report_id: int, user_id: int) -> bool:
    result = await db.execute(
        select(CustomReportAssignment)
        .where(CustomReportAssignment.custom_report_id == custom_report_id, CustomReportAssignment.user_id == user_id)
    )
    perm = result.scalar_one_or_none()
    if not perm:
        return False
    await db.delete(perm)
    await db.flush()
    CUSTOM_REPORT_CACHE.invalidate_report(custom_report_id)
    return True


async def list_custom_report_assignments(db: AsyncSession, custom_report_id: int) -> list[CustomReportAssignment]:
    result = await db.execute(
        select(CustomReportAssignment)
        .where(CustomReportAssignment.custom_report_id == custom_report_id)
        .options(selectinload(CustomReportAssignment.user))
    )
    return list(result.scalars().all())


def _parse_kpi_formula_dependencies(formula_expression: str) -> list[int]:
    if not formula_expression:
        return []
    import re
    pattern = r'\b(?:KPI_FIELD|SUM_KPI_ITEMS|AVG_KPI_ITEMS|COUNT_KPI_ITEMS|MIN_KPI_ITEMS|MAX_KPI_ITEMS|KPI_GROUP_BY|UNIQUE_COUNT_KPI_ITEMS|COUNT_UNIQUE_KPI_ITEMS|SUM_KPI_ITEMS_WHERE|AVG_KPI_ITEMS_WHERE|COUNT_KPI_ITEMS_WHERE|COUNT_UNIQUE_KPI_ITEMS_WHERE|MIN_KPI_ITEMS_WHERE|MAX_KPI_ITEMS_WHERE|FETCH_KPI_ITEMS_WHERE)\s*\(\s*(\d+)'
    matches = re.findall(pattern, formula_expression, re.IGNORECASE)
    return [int(m) for m in matches]


def _topological_sort_report_kpis(kpis: list[KPI]) -> list[KPI]:
    adj: dict[int, set[int]] = {}
    in_degree: dict[int, int] = {}
    kpi_map = {k.id: k for k in kpis}

    for k in kpis:
        adj[k.id] = set()
        in_degree[k.id] = 0

    for k in kpis:
        for f in getattr(k, "fields", []) or []:
            ft = getattr(f, "field_type", None)
            ft_str = ft.value if hasattr(ft, "value") else str(ft)
            exprs = []
            if ft_str == "formula":
                exprs.append(getattr(f, "formula_expression", "") or "")
            elif ft_str == "multi_line_items":
                for sf in getattr(f, "sub_fields", []) or []:
                    cfg = getattr(sf, "config", None) or {}
                    expr = cfg.get("formula_expression") or getattr(sf, "formula_expression", None)
                    if expr:
                        exprs.append(expr)
            for expr in exprs:
                for dep_id in _parse_kpi_formula_dependencies(expr):
                    if dep_id in kpi_map and dep_id != k.id:
                        if k.id not in adj[dep_id]:
                            adj[dep_id].add(k.id)
                            in_degree[k.id] += 1

    from collections import deque
    queue = deque([kid for kid, deg in in_degree.items() if deg == 0])
    order = []
    while queue:
        u = queue.popleft()
        order.append(u)
        for v in adj[u]:
            in_degree[v] -= 1
            if in_degree[v] == 0:
                queue.append(v)

    if len(order) != len(kpis):
        return kpis

    return [kpi_map[kid] for kid in order]


def _sort_formula_subfields(formula_sfs: list) -> list:
    sorted_sfs = []
    visited = set()
    sfs_dict = {item[0]: item for item in formula_sfs}
    
    def visit(sf_key):
        if sf_key in visited:
            return
        visited.add(sf_key)
        item = sfs_dict.get(sf_key)
        if not item:
            return
        expr = item[1]
        if expr:
            for other_key in sfs_dict:
                if other_key != sf_key and other_key in expr:
                    visit(other_key)
        sorted_sfs.append(item)

    for sf_key in sfs_dict:
        visit(sf_key)
    return sorted_sfs


def evaluate_report_table_footer_rows(
    footer_config: dict | None,
    sub_fields: list[dict],
    value_items: list[dict],
) -> list[dict] | None:
    """
    Evaluates custom report table footer config against rendered table rows.
    Formula calculations (SUM, COUNT, AVG, MIN, MAX) run purely over the report table dataset.
    """
    if not footer_config or not isinstance(footer_config, dict) or not footer_config.get("enabled"):
        return None

    raw_rows = footer_config.get("rows")
    if not isinstance(raw_rows, list) or not raw_rows:
        return None

    evaluated_rows = []
    for r_idx, row in enumerate(raw_rows):
        if not isinstance(row, dict):
            continue
        cells = row.get("cells")
        if not isinstance(cells, list):
            continue

        eval_cells = []
        for cell_idx, cell in enumerate(cells):
            if not isinstance(cell, dict):
                continue
            colspan = cell.get("colspan", 1)
            try:
                colspan = max(1, int(colspan))
            except (ValueError, TypeError):
                colspan = 1

            content_type = cell.get("content_type", "text")
            align = cell.get("align")
            if not align:
                align = "left" if cell_idx == 0 else "center"
            bold = cell.get("bold", True)
            dec_places = cell.get("decimal_places", 2)

            val_str = ""
            if content_type == "text":
                val_str = str(cell.get("text") or "")
            elif content_type == "formula":
                op = str(cell.get("formula_op") or "SUM").upper()
                col_key = str(cell.get("column_key") or "")

                nums = []
                for row_dict in value_items:
                    if isinstance(row_dict, dict) and col_key in row_dict:
                        raw_v = row_dict[col_key]
                        if raw_v is not None:
                            try:
                                nums.append(float(raw_v))
                            except (ValueError, TypeError):
                                try:
                                    cleaned = str(raw_v).replace(",", "").strip()
                                    if cleaned:
                                        nums.append(float(cleaned))
                                except (ValueError, TypeError):
                                    pass

                if op == "SUM":
                    res_val = sum(nums) if nums else 0.0
                elif op == "COUNT":
                    res_val = len(nums)
                elif op in ("AVG", "AVERAGE"):
                    res_val = (sum(nums) / len(nums)) if nums else 0.0
                elif op == "MIN":
                    res_val = min(nums) if nums else 0.0
                elif op == "MAX":
                    res_val = max(nums) if nums else 0.0
                else:
                    res_val = sum(nums) if nums else 0.0

                if isinstance(res_val, (int, float)) and res_val % 1 == 0:
                    val_str = f"{int(res_val):,}"
                elif dec_places is not None and str(dec_places).lower() != "auto":
                    try:
                        dp = int(dec_places)
                        res_val = round(res_val, dp)
                        if dp == 0:
                            val_str = f"{int(res_val):,}"
                        else:
                            val_str = f"{res_val:,.{dp}f}"
                    except (ValueError, TypeError):
                        val_str = f"{res_val:,.2f}" if isinstance(res_val, float) else str(res_val)
                else:
                    val_str = f"{res_val:g}" if isinstance(res_val, (int, float)) else str(res_val)

            eval_cells.append({
                "id": cell.get("id") or f"c_{len(eval_cells)}",
                "colspan": colspan,
                "content_type": content_type,
                "formula_op": cell.get("formula_op"),
                "column_key": cell.get("column_key"),
                "text": cell.get("text"),
                "value": val_str,
                "align": align,
                "bold": bold,
            })

        if eval_cells:
            evaluated_rows.append({
                "id": row.get("id") or f"r_{r_idx}",
                "cells": eval_cells,
            })

    return evaluated_rows if evaluated_rows else None


async def generate_custom_report_data(
    db: AsyncSession,
    id: int,
    org_id: int,
    year: str | int | None = None,
    include_drafts: bool = False,
    on_progress: Callable[[int], None] | None = None,
    preview: bool = False,
    include_attachments: bool = True,
    by_default: bool = False,
    period_type: str | None = None,
    current_user: User | None = None,
) -> dict[str, Any] | None:
    if on_progress:
        on_progress(10)
    custom_report = await get_custom_report(db, id, org_id)
    if not custom_report:
        return None

    selected_period = year
    # Only enter date-based period resolution when NOT in "Data Entry" (by_default) mode.
    # by_default=True means the user explicitly chose "Data Entry" — bypass ALL date-column
    # filtering and resolve entries purely by entry year (KPIEntry.year == yr).
    if custom_report and getattr(custom_report, "fetch_data_with_date", False) and not by_default:
        if not selected_period or selected_period == "by_default":
            config = getattr(custom_report, "date_fetching_config", None) or {}
            def_period = config.get("default_period")
            if def_period:
                selected_period = def_period
            else:
                import datetime
                import math
                today = datetime.date.today()
                start_month = int(config.get("custom_period_start_month") or config.get("start_month") or 1)
                if today.month >= start_month:
                    start_year = today.year
                else:
                    start_year = today.year - 1

                prefix = config.get("custom_period_prefix") or config.get("prefix") or ""
                suffix = config.get("custom_period_suffix") or config.get("suffix") or ""
                display_format = config.get("custom_period_display_format") or config.get("display_format") or "YYYY"
                duration = int(config.get("custom_period_duration_months") or config.get("duration_months") or 12)

                if display_format == "YYYY":
                    body = str(start_year)
                elif display_format == "YYYY/YY":
                    end_year = start_year + math.ceil(duration / 12)
                    body = f"{start_year}/{str(end_year % 100).zfill(2)}"
                elif display_format == "YYYY-YY":
                    end_year = start_year + math.ceil(duration / 12)
                    body = f"{start_year}-{str(end_year % 100).zfill(2)}"
                elif display_format == "YYYY-YYYY":
                    end_year = start_year + math.ceil(duration / 12)
                    body = f"{start_year}-{end_year}"
                elif display_format == "YYYY–YYYY":
                    end_year = start_year + math.ceil(duration / 12)
                    body = f"{start_year}–{end_year}"
                elif display_format == "YY/YYYY":
                    end_year = start_year + math.ceil(duration / 12)
                    body = f"{str(start_year % 100).zfill(2)}/{end_year}"
                else:
                    body = str(start_year)

                selected_period = f"{prefix}{body}{suffix}"

    date_range = None
    entry_start_year: int | None = None  # calendar start year of the period (e.g. 2026 for "2026/27")
    # Only build date_range for custom periods (not Data Entry / by_default mode)
    if custom_report and getattr(custom_report, "fetch_data_with_date", False) and not by_default and selected_period:
        org = (await db.execute(select(Organization).where(Organization.id == org_id))).scalar_one_or_none()
        if org:
            try:
                from app.widget_data.service import get_widget_date_col_key, resolve_date_range_for_period
                start_date, end_date, entry_year = resolve_date_range_for_period(org, str(selected_period), period_type)
                entry_start_year = start_date.year  # e.g. 2026 for July-2026 to June-2027
                year = entry_year  # e.g. 2027 (the year entries are stored under)
                date_range = (start_date, end_date)
            except Exception:
                pass

    yr = _parse_year_int(year)
    # yr_display: for custom periods show the period label; for Data Entry show the numeric year
    yr_display = selected_period if (custom_report and getattr(custom_report, "fetch_data_with_date", False) and not by_default and selected_period) else yr

    # Resolve period type and period_info metadata
    resolved_period_type = "Data Entry"
    if custom_report and getattr(custom_report, "fetch_data_with_date", False) and not by_default:
        resolved_period_type = period_type or (custom_report.date_fetching_config or {}).get("default_period_type") or (custom_report.date_fetching_config or {}).get("period_type") or "Data Entry"
        if resolved_period_type == "by_default":
            resolved_period_type = "Data Entry"

    # Capitalize first letter of each word to make it look premium (e.g. "Data Entry", "Fiscal Year")
    resolved_period_type = " ".join(word.capitalize() for word in str(resolved_period_type).split())
    period_info = f"{resolved_period_type} : {yr_display}"

    # 1. Identify all referenced KPIs and Attachment fields
    referenced_kpi_ids = set()
    attachment_kpi_field_ids = set()
    if getattr(custom_report, "attachments", None):
        for att in custom_report.attachments:
            if att.kpi_id:
                referenced_kpi_ids.add(att.kpi_id)
            if att.kpi_field_id:
                attachment_kpi_field_ids.add(att.kpi_field_id)

    for sec in custom_report.sections:
        referenced_kpi_ids.add(sec.kpi_id)
        for f in sec.fields:
            referenced_kpi_ids.add(f.kpi_field.kpi_id)

    unique_kpi_count = len(referenced_kpi_ids)

    if not referenced_kpi_ids:
        # Empty report
        if on_progress:
            on_progress(100)
        return {
            "custom_report_id": custom_report.id,
            "custom_report_name": custom_report.name,
            "custom_report_description": custom_report.description,
            "organization_id": custom_report.organization_id,
            "report_header_id": custom_report.report_header_id,
            "show_report_name": custom_report.show_report_name,
            "branding_title": custom_report.branding_title,
            "year": yr_display,
            "period_type": resolved_period_type,
            "period_info": period_info,
            "sections": [],
            "unique_kpi_count": unique_kpi_count,
        }

    if on_progress:
        on_progress(20)

    # 2. Load KPIs and their fields
    kpis_result = await db.execute(
        select(KPI)
        .where(KPI.id.in_(referenced_kpi_ids))
        .options(
            selectinload(KPI.fields).selectinload(KPIField.sub_fields),
            noload(KPI.organization),
            noload(KPI.domain)
        )
    )
    kpis_by_id = {k.id: k for k in kpis_result.scalars().unique().all()}

    # Expand referenced_kpi_ids to include formula dependencies and configured date-fetching KPIs
    config = getattr(custom_report, "date_fetching_config", None) or {}
    configured_kpis = config.get("configured_kpi_ids") or []
    for c_kid in configured_kpis:
        try:
            referenced_kpi_ids.add(int(c_kid))
        except (ValueError, TypeError):
            pass

    resolved_kpi_ids = set()
    while True:
        unresolved_ids = referenced_kpi_ids - resolved_kpi_ids
        if not unresolved_ids:
            break
            
        missing_ids = unresolved_ids - set(kpis_by_id.keys())
        if missing_ids:
            add_result = await db.execute(
                select(KPI)
                .where(KPI.id.in_(missing_ids))
                .options(
                    selectinload(KPI.fields).selectinload(KPIField.sub_fields),
                    noload(KPI.organization),
                    noload(KPI.domain)
                )
            )
            for k in add_result.scalars().unique().all():
                kpis_by_id[k.id] = k
                
        new_deps = set()
        for kid in unresolved_ids:
            kpi = kpis_by_id.get(kid)
            if not kpi:
                continue
            for f in kpi.fields or []:
                ft = f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type)
                # Direct formula field
                if ft == "formula" and f.formula_expression:
                    for dep_id in _parse_kpi_formula_dependencies(f.formula_expression):
                        new_deps.add(dep_id)
                # Subfields of multi-line field
                elif ft == "multi_line_items":
                    sub_fields = getattr(f, "sub_fields", []) or []
                    for sf in sub_fields:
                        sft = sf.field_type.value if hasattr(sf.field_type, "value") else str(sf.field_type)
                        sf_cfg = getattr(sf, "config", None) or {}
                        expr = sf_cfg.get("formula_expression")
                        if (sft == "formula" or expr) and expr:
                            for dep_id in _parse_kpi_formula_dependencies(expr):
                                new_deps.add(dep_id)
                                
        resolved_kpi_ids.update(unresolved_ids)
        referenced_kpi_ids.update(new_deps)

    # 3. Load Org Config
    org = await db.get(Organization, org_id)
    org_td = TimeDimension(getattr(org, "time_dimension", None) or "yearly") if org else TimeDimension.YEARLY

    # Batch load all entries for all referenced KPIs at once to avoid loop database hits
    all_entries_filters = [
        KPIEntry.organization_id == org_id,
        KPIEntry.kpi_id.in_(referenced_kpi_ids),
    ]
    if not date_range:
        all_entries_filters.append(KPIEntry.year == yr)
    if not include_drafts:
        all_entries_filters.append(KPIEntry.is_draft == False)

    entries_result = await db.execute(
        select(KPIEntry)
        .where(*all_entries_filters)
        .options(selectinload(KPIEntry.field_values))
    )
    all_entries_list = list(entries_result.scalars().all())

    # Fetch counts of multi-line rows for all these entries to prioritize populated entries
    entry_ids = [e.id for e in all_entries_list]
    mli_counts = {}
    if entry_ids:
        from sqlalchemy import func
        counts_res = await db.execute(
            select(KpiMultiLineRow.entry_id, func.count(KpiMultiLineRow.id))
            .where(KpiMultiLineRow.entry_id.in_(entry_ids))
            .group_by(KpiMultiLineRow.entry_id)
        )
        mli_counts = {r[0]: r[1] for r in counts_res.all()}

    # Group entries by kpi_id
    entries_by_kpi = {}
    for entry in all_entries_list:
        entries_by_kpi.setdefault(entry.kpi_id, []).append(entry)

    kpi_evaluated_data = {}

    if on_progress:
        on_progress(30)

    # 4. Fetch entries and evaluate formulas for each referenced KPI (topologically sorted)
    sorted_kpis_list = _topological_sort_report_kpis(list(kpis_by_id.values()))
    total_kpis = len(sorted_kpis_list)

    recalculated_kpi_values: dict[tuple[int, str], float | int] = {}
    recalculated_kpi_mli_data: dict[tuple[int, str], list[dict]] = {}

    # ------------------------------------------------------------------
    # Pre-fetch all numeric KPI field values for this org+year in ONE
    # query, eliminating N separate _load_other_kpi_values calls inside
    # the per-KPI loop below.
    # ------------------------------------------------------------------
    base_other_kpi_values = await bulk_load_org_kpi_values(db, yr, org_id)

    # -------------------------------------------------------------------
    # Filter-Aware MLI Processing flag
    # When True, MLI Advance Filters are applied BEFORE scalar formula
    # evaluation so that formula results reflect only the visible rows.
    # -------------------------------------------------------------------
    apply_filter_processing: bool = bool(getattr(custom_report, "apply_further_processing_based_on_mli_filter", False))

    # Build a lookup: kpi_field_id -> CustomReportField.config
    # (only for fields that have actual advance-filter conditions)
    mf_config_map: dict[int, dict] = {}
    if apply_filter_processing:
        for _sec in custom_report.sections:
            for _f in _sec.fields:
                _cfg = getattr(_f, "config", None) or {}
                _filters = _cfg.get("filters") or {}
                if _filters.get("conditions"):
                    mf_config_map[_f.kpi_field_id] = _cfg

    for idx, kpi in enumerate(sorted_kpis_list):
        kid = kpi.id
        fields_to_include = sorted(list(kpi.fields or []), key=lambda f: (f.sort_order, f.id))
        kpi_td_raw = getattr(kpi, "time_dimension", None)
        kpi_td = TimeDimension(kpi_td_raw) if kpi_td_raw else None
        effective_td = effective_kpi_time_dimension(kpi_td, org_td)

        all_entries = entries_by_kpi.get(kpi.id, [])
        if date_range:
            real_entry = None
            if all_entries:
                def _score_entry(e):
                    # Check if entry has any populated multi-line rows
                    rows_count = mli_counts.get(e.id, 0)
                    # Check if entry has any populated scalar field values
                    has_scalar_data = False
                    for fv in (e.field_values or []):
                        if (fv.value_number is not None and fv.value_number != 0) or \
                           (fv.value_text is not None and fv.value_text != "") or \
                           fv.value_boolean is not None or \
                           fv.value_date is not None or \
                           fv.value_json is not None:
                            has_scalar_data = True
                            break
                    
                    data_score = 1 if (rows_count > 0 or has_scalar_data) else 0
                    
                    if e.year == yr:
                        year_score = 3
                    elif entry_start_year is not None and e.year == entry_start_year:
                        year_score = 2
                    else:
                        year_score = 1 - abs(e.year - yr) * 0.1
                        
                    return (data_score, year_score, e.year)

                all_entries_sorted = sorted(
                    all_entries,
                    key=_score_entry,
                    reverse=True
                )
                real_entry = all_entries_sorted[0]
            if real_entry:
                entries_sorted = [real_entry]
            else:
                entries_sorted = []
        else:
            entries_sorted = sorted(
                all_entries,
                key=lambda e: period_key_sort_order(getattr(e, "period_key", "") or "", effective_td),
            )
            if len(entries_sorted) > 1:
                entries_sorted = [entries_sorted[-1]]

        need_cross_kpi = _formulas_need_other_kpi_values(fields_to_include)
        # Build per-KPI other_kpi_values from the pre-fetched base dict.
        # Overlay any values already recalculated earlier in this loop so
        # that formula dependencies resolved in the correct topological order
        # always see the freshly-computed value, not the stored DB value.
        if need_cross_kpi:
            other_kpi_values = dict(base_other_kpi_values)
            for (r_kid, r_fkey), rval in recalculated_kpi_values.items():
                other_kpi_values[(r_kid, r_fkey)] = rval
        else:
            # Still overlay recalculated values for formulas that only depend
            # on sibling fields within the same KPI batch.
            other_kpi_values = dict(recalculated_kpi_values)

        entry_ids_sorted = [e.id for e in entries_sorted]

        # Load multi-line rows
        ml_fields = [f for f in fields_to_include if f.field_type == FieldType.multi_line_items]
        ml_rows_by_field_id = {}
        for mf in ml_fields:
            if mf.id in attachment_kpi_field_ids:
                limit_val = 50
            else:
                limit_val = None
                
            mf_date_range = None
            if date_range:
                config = getattr(custom_report, "date_fetching_config", None) or {}
                date_col_key = get_widget_date_col_key(config, kpi.id, mf.key, mf)
                if date_col_key:
                    start_date, end_date = date_range
                    mf_date_range = (start_date, end_date, str(date_col_key))

            target_entry_ids = [e.id for e in all_entries if e.id] if (date_range and mf_date_range) else entry_ids_sorted
            batch_res = await _load_multi_line_items_rows_batch(
                db, entry_ids=target_entry_ids, field=mf, limit=limit_val, date_range=mf_date_range, custom_report_id=id, current_user=current_user
            )
            # Re-evaluate any formula subfields on MLI rows
            sub_fields_orm = getattr(mf, "sub_fields", []) or []
            formula_sfs = []
            for sf in sub_fields_orm:
                cfg = getattr(sf, "config", None) or {}
                expr = cfg.get("formula_expression")
                cond_logic = cfg.get("conditional_logic")
                sft = getattr(sf.field_type, "value", str(sf.field_type))
                if sft == "formula" or expr:
                    formula_sfs.append((sf.key, expr, cond_logic, cfg))

            recalculated_batch = {}
            for eid, rows_list in batch_res.items():
                if formula_sfs:
                    sorted_formula_sfs = _sort_formula_subfields(formula_sfs)
                    current_rows = [dict(r) for r in rows_list]
                    for sf_key, expr, cond_logic, cfg in sorted_formula_sfs:
                        if not expr:
                            continue
                        # Build the multi-line context ONCE per formula subfield
                        # (not once per row) — avoids rebuilding a potentially
                        # large list reference on every iteration.
                        ml_context = {mf.key: current_rows}
                        for r_copy in current_rows:
                            new_val = evaluate_formula(
                                expr,
                                {},
                                ml_context,
                                other_kpi_values,
                                current_row=r_copy,
                                other_kpi_multi_line_data=recalculated_kpi_mli_data,
                            )
                            if cond_logic and isinstance(cond_logic, dict) and cond_logic.get("enabled"):
                                new_val = apply_conditional_logic(new_val, cond_logic)
                            
                            if new_val is not None:
                                dec_places = cfg.get("decimal_places") if isinstance(cfg, dict) else None
                                if dec_places is None:
                                    dec_places = 2
                                if dec_places is not None and str(dec_places).lower() != "auto":
                                    try:
                                        dp = int(dec_places)
                                        if isinstance(new_val, (float, int)) and not isinstance(new_val, bool):
                                            new_val = round(float(new_val), dp)
                                            if dp == 0:
                                                new_val = int(new_val)
                                    except (ValueError, TypeError):
                                        pass

                            r_copy[sf_key] = new_val if new_val is not None else 0
                    recalculated_batch[eid] = current_rows
                else:
                    recalculated_batch[eid] = rows_list

            # ----------------------------------------------------------
            # Filter-Aware Pre-Evaluation Filtering
            # If enabled, apply advance filters NOW (before formula eval)
            # so scalar formulas operate on the same filtered dataset.
            # ----------------------------------------------------------
            if apply_filter_processing and mf.id in mf_config_map:
                _field_cfg = mf_config_map[mf.id]
                _raw_filters = _field_cfg.get("filters") or {}
                if _raw_filters and _raw_filters.get("conditions"):
                    try:
                        from app.entries.multi_item_filters import row_passes_filters
                        from app.entries.reference_filter_resolve import build_reference_resolution_map
                        _conds = _raw_filters.get("conditions")
                        # Collect all rows across entries for resolution-map building
                        _all_rows_flat: list[dict] = []
                        for _rl in recalculated_batch.values():
                            if isinstance(_rl, list):
                                _all_rows_flat.extend(_rl)
                        _resolution_maps = await build_reference_resolution_map(
                            db, org_id, yr, mf, _conds, _all_rows_flat
                        )
                        _ref_ftypes = {
                            sf.key: (sf.field_type.value if hasattr(sf.field_type, "value") else sf.field_type)
                            for sf in (getattr(mf, "sub_fields", []) or [])
                        }
                        _filtered_batch: dict = {}
                        for _eid, _rl in recalculated_batch.items():
                            _filtered_batch[_eid] = [
                                r for r in _rl
                                if row_passes_filters(r, _raw_filters, resolution_maps=_resolution_maps, reference_field_types=_ref_ftypes)
                            ]
                        recalculated_batch = _filtered_batch
                    except Exception:
                        pass  # Fallback: keep unfiltered rows if filter logic fails

            if date_range:
                target_eid = entries_sorted[0].id if entries_sorted else 0
                combined_rows = []
                for rows_list in recalculated_batch.values():
                    combined_rows.extend(rows_list)
                combined_rows = _deduplicate_rows(combined_rows)
                ml_rows_by_field_id[mf.id] = {target_eid: combined_rows}
                recalculated_kpi_mli_data[(kpi.id, mf.key)] = combined_rows
            else:
                ml_rows_by_field_id[mf.id] = recalculated_batch
                all_rows = []
                for rlist in recalculated_batch.values():
                    if isinstance(rlist, list):
                        all_rows.extend(rlist)
                all_rows = _deduplicate_rows(all_rows)
                recalculated_kpi_mli_data[(kpi.id, mf.key)] = all_rows

        if on_progress:
            on_progress(int(30 + 50 * (idx + 1) / total_kpis))

        evaluated_fields = {}

        if not entries_sorted:
            # Placeholder for no data
            _NO_DATA_PLACEHOLDER = "No data entered"
            for f in fields_to_include:
                field_payload = {
                    "field_key": f.key,
                    "field_name": f.name,
                    "value": _NO_DATA_PLACEHOLDER,
                    "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                }
                if f.field_type == FieldType.multi_line_items:
                    sub_fields_orm = getattr(f, "sub_fields") or []
                    field_payload["value_items"] = []
                    field_payload["sub_field_keys"] = [sf.key for sf in sub_fields_orm]
                    field_payload["sub_fields"] = [{"key": sf.key, "name": getattr(sf, "name", sf.key)} for sf in sub_fields_orm]
                evaluated_fields[f.key] = field_payload
        else:
            entry = entries_sorted[0]
            if kpi and getattr(kpi, "is_joined", False):
                from app.entries.load_joined import load_joined_scalar_values
                entry_fvs = await load_joined_scalar_values(
                    db, joined_kpi=kpi, entry_id=entry.id, current_user_id=current_user.id if current_user else None
                )
                fv_by_field = {fv.field_id: fv for fv in entry_fvs}
            else:
                fv_by_field = {fv.field_id: fv for fv in (entry.field_values or [])}
            value_by_key = {}
            multi_line_items_data = {}

            for f in fields_to_include:
                if f.field_type == FieldType.formula:
                    continue
                fv = fv_by_field.get(f.id)
                val = None
                
                if f.field_type == FieldType.multi_line_items:
                    rows_items = ml_rows_by_field_id.get(f.id, {}).get(entry.id, [])
                    multi_line_items_data[f.key] = rows_items
                    val = rows_items
                elif fv:
                    if fv.value_date is not None:
                        val = fv.value_date.isoformat() if hasattr(fv.value_date, "isoformat") else str(fv.value_date)
                    elif fv.value_text is not None:
                        val = fv.value_text
                    elif fv.value_number is not None:
                        val = fv.value_number
                    elif fv.value_json is not None:
                        val = fv.value_json
                    elif fv.value_boolean is not None:
                        val = fv.value_boolean

                    if f.field_type == FieldType.number and fv.value_number is not None:
                        value_by_key[f.key] = fv.value_number
                        try:
                            recalculated_kpi_values[(kpi.id, f.key)] = float(fv.value_number)
                        except (TypeError, ValueError):
                            pass

                field_payload = {
                    "field_key": f.key,
                    "field_name": f.name,
                    "value": val,
                    "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                }
                if f.field_type == FieldType.multi_line_items:
                    sub_fields_orm = getattr(f, "sub_fields") or []
                    field_payload["sub_field_keys"] = [sf.key for sf in sub_fields_orm]
                    field_payload["sub_fields"] = [{"key": sf.key, "name": getattr(sf, "name", sf.key)} for sf in sub_fields_orm]
                    field_payload["value_items"] = val if isinstance(val, list) else []
                evaluated_fields[f.key] = field_payload
                if val is not None and f.field_type == FieldType.number:
                    value_by_key[f.key] = val
                    try:
                        recalculated_kpi_values[(kpi.id, f.key)] = float(val)
                    except (TypeError, ValueError):
                        pass

            # Seed formula values
            for f in fields_to_include:
                if f.field_type != FieldType.formula:
                    continue
                fv_formula = fv_by_field.get(f.id)
                if not fv_formula or fv_formula.value_number is None:
                    continue
                try:
                    value_by_key[f.key] = float(fv_formula.value_number)
                except (TypeError, ValueError):
                    continue

            # Evaluate formula fields
            for f in fields_to_include:
                if f.field_type == FieldType.formula and f.formula_expression:
                    computed = evaluate_formula(
                        f.formula_expression,
                        value_by_key,
                        multi_line_items_data,
                        other_kpi_values,
                        other_kpi_multi_line_data=recalculated_kpi_mli_data,
                    )
                    if computed is None:
                        fv_formula = fv_by_field.get(f.id)
                        if fv_formula and fv_formula.value_number is not None:
                            computed = fv_formula.value_number
                    field_payload = {
                        "field_key": f.key,
                        "field_name": f.name,
                        "value": computed,
                        "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                    }
                    evaluated_fields[f.key] = field_payload
                    if computed is not None:
                        value_by_key[f.key] = computed
                        try:
                            recalculated_kpi_values[(kpi.id, f.key)] = float(computed)
                        except (TypeError, ValueError):
                            pass

        kpi_evaluated_data[kpi.id] = evaluated_fields

    # 5. Build structured layout output with hierarchical numbering
    sections_out = []
    for s_idx, sec in enumerate(custom_report.sections):
        sec_num = str(s_idx + 1)
        sec_header = sec.custom_header or (sec.kpi.name if sec.kpi else f"Section {sec_num}")

        fields_out = []
        for f_idx, f in enumerate(sec.fields):
            f_num = f"{sec_num}.{f_idx + 1}"
            kfield = f.kpi_field

            # Lookup evaluated value
            kpi_data = kpi_evaluated_data.get(kfield.kpi_id, {})
            field_eval = kpi_data.get(kfield.key, {})

            cfg = getattr(f, "config", None) or {}
            custom_name = cfg.get("custom_name") or kfield.name

            field_payload = {
                "id": f.id,
                "kpi_field_id": f.kpi_field_id,
                "field_key": kfield.key,
                "field_name": custom_name,
                "field_type": kfield.field_type.value if hasattr(kfield.field_type, "value") else str(kfield.field_type),
                "number": f_num,
                "value": field_eval.get("value"),
            }
            if kfield.field_type == FieldType.multi_line_items:
                cfg = getattr(f, "config", None) or {}
                field_payload["config"] = cfg
                
                # Fetch all sub_fields and values
                all_sub_fields = field_eval.get("sub_fields") or []
                all_value_items = field_eval.get("value_items") or []
                
                # Column selection filtering
                visible_keys = cfg.get("selected_columns")
                if visible_keys is None:
                    # Default to first 5 columns if not explicitly configured to prevent ReportLab crash on export
                    visible_keys = [sf["key"] for sf in all_sub_fields][:5]

                sf_map = {sf["key"]: sf for sf in all_sub_fields}
                custom_sub_labels = cfg.get("custom_sub_field_labels") or {}
                filtered_sub_fields = []
                for k in visible_keys:
                    if k in sf_map:
                        sf_copy = dict(sf_map[k])
                        if k in custom_sub_labels and custom_sub_labels[k]:
                            sf_copy["name"] = custom_sub_labels[k]
                        filtered_sub_fields.append(sf_copy)
                filtered_sub_field_keys = [sf["key"] for sf in filtered_sub_fields]
                
                # Row filtering
                # When apply_filter_processing=True, rows were already filtered during
                # the MLI loading phase (before formula eval), so we skip re-filtering here
                # to avoid redundant async work and potential double-reduction.
                raw_filters = cfg.get("filters") or {}
                filtered_value_items = []
                if all_value_items:
                    if not apply_filter_processing and raw_filters and raw_filters.get("conditions"):
                        from app.entries.multi_item_filters import row_passes_filters
                        from app.entries.reference_filter_resolve import build_reference_resolution_map
                        conds = raw_filters.get("conditions")
                        resolution_maps = await build_reference_resolution_map(
                            db, org_id, yr, kfield, conds, all_value_items
                        )
                        reference_field_types = {sf.key: sf.field_type.value if hasattr(sf.field_type, "value") else sf.field_type for sf in kfield.sub_fields}
                        for r in all_value_items:
                            if row_passes_filters(r, raw_filters, resolution_maps=resolution_maps, reference_field_types=reference_field_types):
                                filtered_value_items.append(r)
                    else:
                        filtered_value_items = all_value_items

                # Sort rows in memory if specified
                sort_col = cfg.get("sort_column")
                sort_dir = cfg.get("sort_direction") or "asc"
                if sort_col and filtered_value_items:
                    reverse = sort_dir == "desc"
                    def sort_key(row: dict):
                        v = row.get(sort_col)
                        try:
                            return float(v)
                        except (TypeError, ValueError):
                            return str(v) if v is not None else ""
                    try:
                        filtered_value_items = sorted(filtered_value_items, key=sort_key, reverse=reverse)
                    except Exception:
                        pass

                
                # Query total count of rows
                if date_range or (raw_filters and raw_filters.get("conditions")):
                    total_cnt = len(filtered_value_items)
                else:
                    from sqlalchemy import func
                    total_cnt = 0
                    if entries_sorted:
                        entry = entries_sorted[0]
                        total_cnt = (
                            await db.execute(
                                select(func.count(KpiMultiLineRow.id)).where(
                                    KpiMultiLineRow.entry_id == entry.id,
                                    KpiMultiLineRow.field_id == kfield.id,
                                )
                            )
                        ).scalar_one() or 0

                if preview:
                    filtered_value_items = filtered_value_items[:50]
                
                field_payload["sub_fields"] = filtered_sub_fields
                field_payload["sub_field_keys"] = filtered_sub_field_keys
                field_payload["value_items"] = filtered_value_items
                field_payload["total_count"] = total_cnt

                footer_cfg = (field_payload.get("config") or {}).get("footer_config")
                evaluated_footer = evaluate_report_table_footer_rows(footer_cfg, filtered_sub_fields, filtered_value_items)
                if evaluated_footer:
                    field_payload["evaluated_footer_rows"] = evaluated_footer
            else:
                field_payload["config"] = getattr(f, "config", None) or {}

            fields_out.append(field_payload)

        sections_out.append({
            "id": sec.id,
            "kpi_id": sec.kpi_id,
            "custom_header": sec_header,
            "number": sec_num,
            "fields": fields_out,
        })

    attachments_out = []
    if getattr(custom_report, "attachments", None):
        for att in custom_report.attachments:
            attachments_out.append({
                "id": att.id,
                "title": att.title,
                "kpi_name": att.kpi.name if att.kpi else "KPI",
                "field_name": att.kpi_field.name if att.kpi_field else "Field",
            })

    return {
        "custom_report_id": custom_report.id,
        "custom_report_name": custom_report.name,
        "template_id": custom_report.id,
        "template_name": custom_report.name,
        "custom_report_description": custom_report.description,
        "organization_id": custom_report.organization_id,
        "report_header_id": custom_report.report_header_id,
        "show_report_name": custom_report.show_report_name,
        "branding_title": custom_report.branding_title,
        "show_odoo_button": custom_report.show_odoo_button,
        "odoo_sync_kpi_ids": custom_report.odoo_sync_kpi_ids or [],

        "scalar_bold": custom_report.scalar_bold,
        "scalar_font_size": custom_report.scalar_font_size,
        "mli_font_size": custom_report.mli_font_size,
        "year": yr_display,
        "period_type": resolved_period_type,
        "period_info": period_info,
        "sections": sections_out,
        "attachments": attachments_out,
        "unique_kpi_count": unique_kpi_count,
    }


async def render_custom_report_html(
    db: AsyncSession,
    id: int,
    org_id: int,
    year: str | int | None = None,
    include_drafts: bool = False,
    report_data: dict | None = None,
) -> str | None:
    if report_data is not None:
        data = report_data
    else:
        data = await generate_custom_report_data(
            db, id, org_id, year=year, include_drafts=include_drafts
        )
    if not data:
        return None

    period_info = data.get("period_info") or f"Data Entry : {data.get('year') or year or ''}"

    # Load custom report header if present
    custom_header_model = None
    if data.get("report_header_id"):
        from app.core.models import CustomReportHeader
        custom_header_model = await db.get(CustomReportHeader, data["report_header_id"])

    # Compile the sections and fields into styled HTML matching Simple Reports styling
    out = []
    
    font_family = "Helvetica, Arial, sans-serif"
    if custom_header_model and custom_header_model.font_family:
        font_family = f"'{custom_header_model.font_family}', Helvetica, Arial, sans-serif"

    out.append(f'<div class="custom-report" style="color: #111; font-family: {font_family};">')
    out.append('''<style>
      th, table th, .report-kpi-table th, .report-simple-table th {
        word-break: normal !important;
        overflow-wrap: normal !important;
        word-wrap: normal !important;
        white-space: normal !important;
        hyphens: none !important;
        -webkit-hyphens: none !important;
        padding: 6px 5px !important;
        vertical-align: middle !important;
        box-sizing: border-box !important;
      }
      .custom-report table tbody tr td:nth-child(n+3),
      .custom-report table thead tr th:nth-child(n+3) {
        text-align: center;
      }
      .custom-report table tbody tr td:nth-child(1),
      .custom-report table thead tr th:nth-child(1) {
        text-align: center;
      }
      .custom-report table tbody tr td:nth-child(2),
      .custom-report table thead tr th:nth-child(2) {
        text-align: left;
      }
    </style>''')

    if not custom_header_model:
        out.append(f'<div style="text-align: right; font-size: 0.85rem; font-weight: bold; color: #475569; margin-bottom: 1.5rem;">{html.escape(period_info)}</div>')

    # Render Header at the top
    if custom_header_model:
        from app.storage.service import get_file_stream as storage_get_file_stream
        import base64
        import mimetypes

        logo1_src = None
        if custom_header_model.logo_path:
            try:
                lbytes = await storage_get_file_stream(db, custom_header_model.organization_id, custom_header_model.logo_path)
                ctype, _ = mimetypes.guess_type(custom_header_model.logo_path)
                ctype = ctype or "image/png"
                b64 = base64.b64encode(lbytes).decode("utf-8")
                logo1_src = f"data:{ctype};base64,{b64}"
            except Exception:
                logo1_src = f"/api/reports/headers/{custom_header_model.id}/logo"

        logo2_src = None
        if custom_header_model.logo_path_2:
            try:
                lbytes2 = await storage_get_file_stream(db, custom_header_model.organization_id, custom_header_model.logo_path_2)
                ctype2, _ = mimetypes.guess_type(custom_header_model.logo_path_2)
                ctype2 = ctype2 or "image/png"
                b64_2 = base64.b64encode(lbytes2).decode("utf-8")
                logo2_src = f"data:{ctype2};base64,{b64_2}"
            except Exception:
                logo2_src = f"/api/reports/headers/{custom_header_model.id}/logo2"

        out.append('<div class="report-header-container" style="display: flex; justify-content: space-between; align-items: center; width: 100%; margin-bottom: 2.25rem; border-bottom: 2px solid #e5e7eb; padding-bottom: 1rem; position: relative;">')
        
        # Left Logo Slot (Pinned Left)
        out.append('<div style="flex: 0 0 auto; display: flex; justify-content: flex-start; align-items: center;">')
        if logo1_src:
            out.append(f'<img src="{logo1_src}" style="max-height: 70px; max-width: 140px; object-fit: contain;" alt="Logo 1" />')
        out.append('</div>')

        # Middle Heading Text Slot
        align = (custom_header_model.text_align or "center").lower()
        desired_fs = custom_header_model.font_size or 18
        main_fs_pt = calc_auto_header_font_size(custom_header_model.main_heading, 500.0, desired_fs)
        main_fs_px = round(main_fs_pt * 1.25, 1)
        main_color = custom_header_model.text_color or "#1e3a8a"

        out.append(f'<div style="flex: 1 1 auto; text-align: {align}; padding: 0 0.5rem;">')
        out.append(f'<h1 style="margin: 0; font-size: {main_fs_px}px; color: {main_color}; font-weight: bold; font-family: {font_family}; text-align: {align}; white-space: nowrap;">{html.escape(custom_header_model.main_heading)}</h1>')
        if custom_header_model.sub_heading:
            sub_fs = custom_header_model.sub_font_size or 11
            sub_color = custom_header_model.sub_text_color or "#4b5563"
            sub_align = (custom_header_model.sub_text_align or align).lower()
            sub_ff = f"'{custom_header_model.sub_font_family}', Helvetica, Arial, sans-serif" if custom_header_model.sub_font_family else font_family
            out.append(f'<div style="margin-top: 0.25rem; font-size: {sub_fs}px; color: {sub_color}; text-align: {sub_align}; font-family: {sub_ff}; font-style: italic;">{html.escape(custom_header_model.sub_heading)}</div>')
        out.append('</div>')

        # Right Logo Slot (Pinned Right) with Period Metadata below Logo 2
        out.append('<div style="flex: 0 0 auto; display: flex; flex-direction: column; align-items: flex-end; justify-content: center;">')
        if logo2_src:
            out.append(f'<img src="{logo2_src}" style="max-height: 70px; max-width: 140px; object-fit: contain; margin-bottom: 4px;" alt="Logo 2" />')
        out.append(f'<div style="font-size: 0.85rem; font-weight: bold; color: #475569; white-space: nowrap; margin-top: 4px;">{html.escape(period_info)}</div>')
        out.append('</div>')

        out.append('</div>')

    # Configurable report name just below header (only if more than 1 KPI)
    if data.get("unique_kpi_count", 0) > 1 and data.get("show_report_name", True):
        rname_color = custom_header_model.kpi_name_color if (custom_header_model and custom_header_model.kpi_name_color) else "#1e3a8a"
        out.append(f'<div style="text-align: center; margin-bottom: 1.5rem;">')
        out.append(f'<h2 style="margin: 0; font-size: 1.4rem; color: {rname_color}; font-weight: bold;">{data.get("custom_report_name", "Custom Report")}</h2>')
        if year:
            out.append(f'<div style="font-size: 0.9rem; color: #4b5563; margin-top: 0.25rem; font-style: italic;">Academic Year: {year}</div>')
        out.append('</div>')

    scalar_bold = data.get("scalar_bold", True)
    if scalar_bold is None:
        scalar_bold = True
    scalar_font_size = data.get("scalar_font_size", 11) or 11
    mli_font_size = data.get("mli_font_size", 10) or 10

    for sec in data["sections"]:
        h1_color = custom_header_model.kpi_name_color if (custom_header_model and custom_header_model.kpi_name_color) else "#1e3a8a"
        out.append('<section style="margin-bottom: 1.5rem;">')
        out.append(
            f'<h2 style="font-size: 1.2rem; margin-bottom: 0.5rem; color: {h1_color}; border-bottom: 2px solid {h1_color}; padding-bottom: 0.25rem; font-weight: bold;">'
            f'{sec["number"]}. {sec["custom_header"]}'
            f'</h2>'
        )
        out.append('<div style="margin-left: 0.5rem; margin-bottom: 0.75rem;">')
        for f in sec["fields"]:
            h2_color = "#374151"
            if f["field_type"] != "multi_line_items":
                val = clean_numeric_value_string(f["value"])
                out.append('<div class="report-field-block" style="display: flex; justify-content: space-between; align-items: center; border-bottom: 1px dashed #e5e7eb; padding-bottom: 0.25rem; margin-bottom: 0.75rem; margin-right: 2rem;">')
                out.append(
                    f'<h3 style="font-size: 1.0rem; margin: 0; color: {h2_color}; font-weight: 600;">'
                    f'{f["number"]}. {f["field_name"]}'
                    f'</h3>'
                )
                bold_style = "font-weight: bold;" if scalar_bold else "font-weight: normal;"
                out.append(f'<span style="color: #111827; {bold_style} font-size: {scalar_font_size}pt;">{val}</span>')
                out.append('</div>')
            else:
                out.append('<div class="report-field-block" style="margin-bottom: 1.25rem;">')
                out.append(
                    f'<h3 style="font-size: 1.0rem; margin-top: 0.75rem; margin-bottom: 0.35rem; color: {h2_color}; font-weight: 600;">'
                    f'{f["number"]}. {f["field_name"]}'
                    f'</h3>'
                )
                out.append('<div style="margin-left: 0.5rem;">')
                if f.get("value_items") or f.get("evaluated_footer_rows"):
                    total_cnt = f.get("total_count", len(f["value_items"]))
                    if total_cnt > len(f["value_items"]):
                        out.append(f'<div style="margin-bottom: 0.5rem; font-style: italic; color: #4b5563; font-size: 0.85rem; font-weight: 500;">Showing first {len(f["value_items"])} rows of {total_cnt} records.</div>')
                    out.append('<table style="border-collapse: collapse; width: 100%; border: 1px solid #d1d5db; margin-top: 0.25rem; margin-bottom: 0.5rem;">')
                    out.append('<thead>')
                    out.append(f'<tr style="background-color: {h1_color}; color: #ffffff; border-bottom: 2px solid {h1_color}; font-size: {mli_font_size}pt;">')
                    out.append(f'<th style="border: 1px solid #d1d5db; padding: 6px 5px; text-align: center; font-weight: 600; color: #ffffff; word-break: normal; overflow-wrap: normal; white-space: normal; hyphens: none; vertical-align: middle;">S.No</th>')
                    col_alignments = (f.get("config") or {}).get("column_alignments") or {}
                    for s_idx, sub in enumerate(f["sub_fields"]):
                        align_css = col_alignments.get(sub["key"])
                        if not align_css:
                            align_css = "left" if s_idx == 0 else "center"
                        out.append(f'<th style="border: 1px solid #d1d5db; padding: 6px 5px; text-align: {align_css}; font-weight: 600; color: #ffffff; word-break: normal; overflow-wrap: normal; white-space: normal; hyphens: none; vertical-align: middle;">{sub["name"]}</th>')
                    out.append('</tr>')
                    out.append('</thead>')
                    out.append('<tbody>')
                    for r_idx, row in enumerate(f["value_items"]):
                        bg = ' style="background-color: #f9fafb;"' if r_idx % 2 == 1 else ''
                        out.append(f'<tr{bg} style="font-size: {mli_font_size}pt;">')
                        out.append(f'<td style="border: 1px solid #d1d5db; padding: 8px; color: #4b5563; text-align: center;">{r_idx + 1}</td>')
                        for s_idx, sub in enumerate(f["sub_fields"]):
                            rval = clean_numeric_value_string(row.get(sub["key"]))
                            align_css = col_alignments.get(sub["key"])
                            if not align_css:
                                align_css = "left" if s_idx == 0 else "center"
                            out.append(f'<td style="border: 1px solid #d1d5db; padding: 8px; color: #111827; text-align: {align_css};">{rval}</td>')
                        out.append('</tr>')
                    out.append('</tbody>')
                    if f.get("evaluated_footer_rows"):
                        out.append('<tfoot>')
                        for f_row in f["evaluated_footer_rows"]:
                            out.append(f'<tr style="font-size: {mli_font_size}pt; background-color: #f8fafc;">')
                            cells = f_row.get("cells", [])
                            total_sub_cols = len(f["sub_fields"])
                            sum_colspan = sum(c.get("colspan", 1) for c in cells)
                            first_cell_extra = 1 if sum_colspan == total_sub_cols else 0

                            for c_idx, cell in enumerate(cells):
                                c_span = cell.get("colspan", 1) + (first_cell_extra if c_idx == 0 else 0)
                                c_val = cell.get("value", "")
                                c_align = cell.get("align", "left")
                                c_bold = "font-weight: bold;" if cell.get("bold", True) else "font-weight: normal;"
                                out.append(f'<td colspan="{c_span}" style="border: 1px solid #d1d5db; padding: 8px; color: #111827; text-align: {c_align}; {c_bold}">{c_val}</td>')
                            out.append('</tr>')
                        out.append('</tfoot>')
                    out.append('</table>')
                else:
                    out.append('<span style="color: #9ca3af; font-style: italic; font-size: 0.9rem;">No data entered</span>')
                out.append('</div>')
                out.append('</div>')
        out.append('</div>')
        out.append('</section>')
    out.append('</div>')
    return "\n".join(out)
    return "\n".join(out)


async def export_custom_report_file(
    db: AsyncSession,
    custom_report_id: int,
    org_id: int,
    year: str | int | None,
    format: str,
    by_default: bool = False,
    period_type: str | None = None,
    current_user: User | None = None,
) -> tuple[bytes, str, str]:
    """Export custom report as PDF, DOCX, or XLSX bytes, with name and content-type."""
    import re
    import io

    # Generate custom report data
    data = await generate_custom_report_data(
        db, custom_report_id, org_id, year=year, include_drafts=False, by_default=by_default, period_type=period_type, current_user=current_user
    )
    if not data:
        raise ValueError("Report data generation failed")
    report_name = data.get("template_name", "Custom Report")
    period_info = data.get("period_info") or f"Data Entry : {data.get('year') or year or ''}"
    # Clean filename
    clean_report_name = re.sub(r'[^\w\s-]', '', report_name).strip().replace(' ', '_')

    # Determine page orientation based on MLI columns
    use_landscape = False
    for sec in data.get("sections", []):
        for f in sec.get("fields", []):
            if f.get("field_type") == "multi_line_items":
                sub_fields = f.get("sub_fields", [])
                if len(sub_fields) > 8:
                    use_landscape = True
                    break
        if use_landscape:
            break

    if format == "xlsx":
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()
        wb.remove(wb.active) # Remove default sheet
        
        # Styles
        title_font = Font(name="Calibri", size=15, bold=True, color="1E3A8A")
        section_font = Font(name="Calibri", size=12, bold=True, color="1F2937")
        header_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
        label_font = Font(name="Calibri", size=10, bold=True)
        normal_font = Font(name="Calibri", size=10)
        border_side = Side(border_style="thin", color="E5E7EB")
        thin_border = Border(left=border_side, right=border_side, top=border_side, bottom=border_side)

        for s_idx, sec in enumerate(data.get("sections", [])):
            sec_name = sec.get("custom_header") or sec.get("kpi_name", f"Section {s_idx+1}")
            sheet_title = re.sub(r'[\\*?:/\[\]]', '', sec_name)[:30].strip() or f"Sheet {s_idx+1}"
            ws = wb.create_sheet(title=sheet_title)
            ws.views.sheetView[0].showGridLines = True

            # Title Block
            ws.merge_cells("A1:D1")
            ws["A1"] = clean_excel_value(f"{report_name} - {sec_name}")
            ws["A1"].font = title_font
            ws["A1"].alignment = Alignment(vertical="center")
            ws.row_dimensions[1].height = 28

            ws["A2"] = period_info
            ws["A2"].font = Font(name="Calibri", size=10, bold=True, color="475569")
            ws.row_dimensions[2].height = 18

            row_num = 4
            scalars = [f for f in sec.get("fields", []) if f.get("field_type") != "multi_line_items"]
            mlis = [f for f in sec.get("fields", []) if f.get("field_type") == "multi_line_items"]

            if scalars:
                # Column Headers
                ws.cell(row=row_num, column=1, value="Field Label").font = header_font
                ws.cell(row=row_num, column=1).fill = header_fill
                ws.cell(row=row_num, column=2, value="Value").font = header_font
                ws.cell(row=row_num, column=2).fill = header_fill
                ws.row_dimensions[row_num].height = 20
                row_num += 1

                for f in scalars:
                    ws.cell(row=row_num, column=1, value=clean_excel_value(f.get("field_name"))).font = label_font
                    ws.cell(row=row_num, column=1).border = thin_border
                    ws.cell(row=row_num, column=2, value=clean_excel_value(f.get("value") if f.get("value") is not None else "—")).font = normal_font
                    ws.cell(row=row_num, column=2).border = thin_border
                    ws.row_dimensions[row_num].height = 18
                    row_num += 1
                row_num += 2 # spacer

            for f in mlis:
                ws.cell(row=row_num, column=1, value=f.get("field_name")).font = section_font
                row_num += 1

                sub_fields = f.get("sub_fields", [])
                value_items = f.get("value_items", [])

                if sub_fields:
                    merged_headers = f.get("config", {}).get("merged_headers") or []
                    
                    def style_header_cell(cell):
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                        cell.border = thin_border

                    if merged_headers:
                        cols_keys = [sf.get("key") for sf in sub_fields]
                        col_indices = {key: idx + 2 for idx, key in enumerate(cols_keys)}
                        covered_cols = set()

                        # Pre-style all header cells in row_num and row_num+1
                        for r_offset in [0, 1]:
                            for c_idx in range(1, len(sub_fields) + 2):
                                style_header_cell(ws.cell(row=row_num + r_offset, column=c_idx))

                        # Vertical-merge Sr. No.
                        ws.cell(row=row_num, column=1, value="Sr. No.")
                        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num+1, end_column=1)

                        # Write grouped merged headers
                        for group in merged_headers:
                            title = group.get("title")
                            start_key = group.get("start_key")
                            end_key = group.get("end_key")
                            if start_key in col_indices and end_key in col_indices:
                                start_col = col_indices[start_key]
                                end_col = col_indices[end_key]
                                ws.cell(row=row_num, column=start_col, value=title)
                                ws.merge_cells(start_row=row_num, start_column=start_col, end_row=row_num, end_column=end_col)
                                for c in range(start_col, end_col + 1):
                                    covered_cols.add(c)

                        # Write column headers and vertical merge ungrouped columns
                        for idx, sf in enumerate(sub_fields):
                            c_idx = idx + 2
                            col_name = sf.get("name") or sf.get("key")
                            if c_idx not in covered_cols:
                                ws.cell(row=row_num, column=c_idx, value=col_name)
                                ws.merge_cells(start_row=row_num, start_column=c_idx, end_row=row_num+1, end_column=c_idx)
                            else:
                                ws.cell(row=row_num+1, column=c_idx, value=col_name)

                        ws.row_dimensions[row_num].height = 20
                        ws.row_dimensions[row_num+1].height = 20
                        row_num += 2
                    else:
                        # Normal Headers
                        sr_c = ws.cell(row=row_num, column=1, value="Sr. No.")
                        style_header_cell(sr_c)

                        for col_idx, sf in enumerate(sub_fields):
                            c = ws.cell(row=row_num, column=col_idx+2, value=sf.get("name") or sf.get("key"))
                            style_header_cell(c)
                        ws.row_dimensions[row_num].height = 20
                        row_num += 1

                    # Body
                    for item_idx, item in enumerate(value_items):
                        sr_val_c = ws.cell(row=row_num, column=1, value=item_idx + 1)
                        sr_val_c.font = normal_font
                        sr_val_c.border = thin_border
                        sr_val_c.alignment = Alignment(horizontal="center", vertical="center")

                        col_alignments = (f.get("config") or {}).get("column_alignments") or {}
                        for col_idx, sf in enumerate(sub_fields):
                            val = item.get(sf.get("key"))
                            c = ws.cell(row=row_num, column=col_idx+2, value=clean_excel_value(val if val is not None else "—"))
                            c.font = normal_font
                            c.border = thin_border
                            col_align = col_alignments.get(sf.get("key"))
                            if col_align not in ["left", "center", "right"]:
                                col_align = "left" if col_idx == 0 else "center"
                            c.alignment = Alignment(horizontal=col_align, vertical="center")
                        ws.row_dimensions[row_num].height = 18
                        row_num += 1

                    # Footer Rows
                    evaluated_footer = f.get("evaluated_footer_rows")
                    if evaluated_footer:
                        for f_row in evaluated_footer:
                            cells = f_row.get("cells", [])
                            total_sub_cols = len(sub_fields)
                            sum_colspan = sum(c.get("colspan", 1) for c in cells)
                            first_cell_extra = 1 if sum_colspan == total_sub_cols else 0

                            col_cursor = 1
                            for c_idx, cell in enumerate(cells):
                                c_span = cell.get("colspan", 1) + (first_cell_extra if c_idx == 0 else 0)
                                c_val = cell.get("value", "")
                                c_align = cell.get("align", "left")
                                c_bold = cell.get("bold", True)

                                target_cell = ws.cell(row=row_num, column=col_cursor, value=c_val)
                                target_cell.font = Font(name="Calibri", size=10, bold=c_bold)
                                target_cell.border = thin_border
                                target_cell.alignment = Alignment(horizontal=c_align, vertical="center")

                                if c_span > 1:
                                    end_col = col_cursor + c_span - 1
                                    ws.merge_cells(start_row=row_num, start_column=col_cursor, end_row=row_num, end_column=end_col)
                                    for c_col in range(col_cursor, end_col + 1):
                                        m_c = ws.cell(row=row_num, column=c_col)
                                        m_c.border = thin_border
                                        m_c.font = Font(name="Calibri", size=10, bold=c_bold)
                                    col_cursor = end_col + 1
                                else:
                                    col_cursor += 1

                            ws.row_dimensions[row_num].height = 20
                            row_num += 1

                row_num += 2 # spacer

            # Column widths formatting (respect custom_widths if set)
            custom_widths = f.get("config", {}).get("column_widths") or {} if mlis else {}
            for col in ws.columns:
                col_letter = get_column_letter(col[0].column)
                c_idx = col[0].column
                max_len = 0
                for cell in col:
                    if cell.value is not None:
                        max_len = max(max_len, len(str(cell.value)))
                
                # If custom width is defined
                cw = None
                if c_idx == 1:
                    cw = custom_widths.get("S.No")
                elif c_idx - 2 < len(sub_fields):
                    sf_obj = sub_fields[c_idx - 2]
                    cw = custom_widths.get(sf_obj.get("key")) or custom_widths.get(sf_obj.get("name"))
                
                if cw:
                    ws.column_dimensions[col_letter].width = max(8.0, float(cw) / 7.5)
                else:
                    ws.column_dimensions[col_letter].width = max(max_len + 3, 14)

        out_io = io.BytesIO()
        wb.save(out_io)
        return out_io.getvalue(), f"{clean_report_name}_{year}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    elif format == "docx":
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from app.core.models import Organization, OrganizationBranding

        scalar_bold = data.get("scalar_bold", True)
        if scalar_bold is None:
            scalar_bold = True
        scalar_font_size = data.get("scalar_font_size", 11) or 11
        mli_font_size = data.get("mli_font_size", 10) or 10

        # Fetch custom header model
        custom_header = None
        if data.get("report_header_id"):
            from app.core.models import CustomReportHeader
            custom_header = await db.get(CustomReportHeader, data["report_header_id"])

        doc = Document()
        for section in doc.sections:
            if use_landscape:
                section.page_width = Inches(11)
                section.page_height = Inches(8.5)
                from docx.enum.section import WD_ORIENT
                section.orientation = WD_ORIENT.LANDSCAPE
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Set docx font family based on custom header
        docx_font = "Calibri"
        if custom_header and custom_header.font_family:
            font_family_lower = custom_header.font_family.lower()
            if font_family_lower in ("times-roman", "times new roman", "times"):
                docx_font = "Times New Roman"
            elif font_family_lower in ("courier", "courier new"):
                docx_font = "Courier New"
            elif font_family_lower == "arial":
                docx_font = "Arial"
            elif font_family_lower == "georgia":
                docx_font = "Georgia"
            elif font_family_lower == "verdana":
                docx_font = "Verdana"
            elif font_family_lower == "calibri":
                docx_font = "Calibri"
            elif font_family_lower == "garamond":
                docx_font = "Garamond"

        # Footer branding title logic
        org = await db.get(Organization, org_id)
        org_branding = (await db.execute(
            select(OrganizationBranding).where(OrganizationBranding.organization_id == org_id)
        )).scalar_one_or_none()

        footer_label = "Confidential Document"
        if org_branding and org_branding.footer_label:
            footer_label = org_branding.footer_label
        else:
            if org and getattr(org, "name", None):
                footer_label += f" | {org.name}"

        # Apply footer to docx
        section = doc.sections[0]
        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_ftr = p_ftr.add_run(footer_label)
        run_ftr.font.size = Pt(8)
        run_ftr.font.name = docx_font
        run_ftr.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

        # Render Header logo + headings if present
        if custom_header:
            from app.storage.service import get_file_stream as storage_get_file_stream
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            if use_landscape:
                heading_len = len(custom_header.main_heading)
                sub_heading_len = len(custom_header.sub_heading) if custom_header.sub_heading else 0
                max_len = max(heading_len, sub_heading_len)
                font_sz = custom_header.font_size or 16
                char_w = (font_sz * 0.45) / 72.0
                est_w = max_len * char_w + 0.1
                max_allowed = 7.2
                mid_w_val = min(max_allowed, max(1.5, est_w))
                
                header_table = doc.add_table(rows=1, cols=3)
                header_table.autofit = False
                header_table.columns[0].width = Inches(1.1)
                header_table.columns[1].width = Inches(mid_w_val)
                header_table.columns[2].width = Inches(1.1)
                
                from docx.enum.table import WD_TABLE_ALIGNMENT
                header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                header_table = doc.add_table(rows=1, cols=3)
                header_table.autofit = False
                header_table.columns[0].width = Inches(1.1)
                header_table.columns[1].width = Inches(4.3)
                header_table.columns[2].width = Inches(1.1)
            cell_logo = header_table.cell(0, 0)
            cell_text = header_table.cell(0, 1)
            cell_logo2 = header_table.cell(0, 2)

            # Logo 1 (Aligned Right against text)
            p_logo = cell_logo.paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                logo_bytes = await storage_get_file_stream(db, org_id, custom_header.logo_path)
                p_logo.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.1))
            except Exception:
                pass

            # Logo 2 (Aligned Left against text)
            if cell_logo2:
                p_logo2 = cell_logo2.paragraphs[0]
                p_logo2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                try:
                    logo_bytes2 = await storage_get_file_stream(db, org_id, custom_header.logo_path_2)
                    p_logo2.add_run().add_picture(io.BytesIO(logo_bytes2), width=Inches(1.1))
                except Exception:
                    pass

                # Add period info to the right cell, right-aligned, below the logo if logo exists
                p_period = cell_logo2.add_paragraph()
                p_period.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_period = p_period.add_run(period_info)
                run_period.font.size = Pt(8.5)
                run_period.font.bold = True
                run_period.font.name = docx_font
                run_period.font.color.rgb = RGBColor(71, 85, 105)

            header_align_str = (custom_header.text_align or "center").lower()
            if header_align_str == "left":
                align_wd = WD_ALIGN_PARAGRAPH.LEFT
            elif header_align_str == "right":
                align_wd = WD_ALIGN_PARAGRAPH.RIGHT
            elif header_align_str == "justify":
                align_wd = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                align_wd = WD_ALIGN_PARAGRAPH.CENTER

            p_head = cell_text.paragraphs[0]
            p_head.alignment = align_wd
            p_head.paragraph_format.space_before = Pt(4)
            run_main = p_head.add_run(custom_header.main_heading)
            run_main.bold = True
            run_main.font.name = docx_font
            run_main.font.size = Pt(custom_header.font_size or 16)
            
            # Apply color
            color_hex = (custom_header.text_color or "#1e3a8a").lstrip("#")
            try:
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
            except Exception:
                r, g, b = 0x1e, 0x3a, 0x8a
            run_main.font.color.rgb = RGBColor(r, g, b)

            # Apply font mapping to oxml
            r_elem = run_main._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), docx_font)
            rFonts.set(qn('w:hAnsi'), docx_font)
            r_elem.append(rFonts)

            if custom_header.sub_heading:
                p_sub = cell_text.add_paragraph()
                sub_align_str = (custom_header.sub_text_align or custom_header.text_align or "center").lower()
                if sub_align_str == "left":
                    sub_align_wd = WD_ALIGN_PARAGRAPH.LEFT
                elif sub_align_str == "right":
                    sub_align_wd = WD_ALIGN_PARAGRAPH.RIGHT
                elif sub_align_str == "justify":
                    sub_align_wd = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    sub_align_wd = WD_ALIGN_PARAGRAPH.CENTER

                p_sub.alignment = sub_align_wd
                p_sub.paragraph_format.space_before = Pt(2)
                run_sub = p_sub.add_run(custom_header.sub_heading)
                run_sub.italic = True
                run_sub.font.name = docx_font
                run_sub.font.size = Pt(custom_header.sub_font_size or 11)

                sub_color_hex = (custom_header.sub_text_color or "#4b5563").lstrip("#")
                try:
                    sr = int(sub_color_hex[0:2], 16)
                    sg = int(sub_color_hex[2:4], 16)
                    sb = int(sub_color_hex[4:6], 16)
                except Exception:
                    sr, sg, sb = 0x4b, 0x55, 0x63
                run_sub.font.color.rgb = RGBColor(sr, sg, sb)

                r_elem_sub = run_sub._r.get_or_add_rPr()
                rFonts_sub = OxmlElement('w:rFonts')
                rFonts_sub.set(qn('w:ascii'), docx_font)
                rFonts_sub.set(qn('w:hAnsi'), docx_font)
                r_elem_sub.append(rFonts_sub)

            # Spacer
            doc.add_paragraph()
        else:
            p_period = doc.add_paragraph()
            p_period.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_period = p_period.add_run(period_info)
            run_period.font.size = Pt(9.5)
            run_period.font.bold = True
            run_period.font.name = docx_font
            run_period.font.color.rgb = RGBColor(71, 85, 105)
            doc.add_paragraph()

        # Configurable report name just below header (only if more than 1 KPI)
        if data.get("unique_kpi_count", 0) > 1 and data.get("show_report_name", True):
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run(report_name)
            r_title.font.size = Pt(20)
            r_title.font.name = docx_font
            r_title.font.bold = True
            
            kn_color_hex = (custom_header.kpi_name_color if (custom_header and custom_header.kpi_name_color) else "#1e3a8a").lstrip("#")
            try:
                kr = int(kn_color_hex[0:2], 16)
                kg = int(kn_color_hex[2:4], 16)
                kb = int(kn_color_hex[4:6], 16)
            except Exception:
                kr, kg, kb = 0x1e, 0x3a, 0x8a
            r_title.font.color.rgb = RGBColor(kr, kg, kb)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run(f"Academic Year: {year}")
            r_sub.font.size = Pt(11)
            r_sub.font.name = docx_font
            r_sub.font.italic = True

        for s_idx, sec in enumerate(data.get("sections", [])):
            sec_name = sec.get("custom_header") or sec.get("kpi_name", f"Section {s_idx+1}")
            h = doc.add_heading(level=1)
            r_h = h.add_run(f"{sec['number']}. {sec_name}")
            r_h.font.name = docx_font
            r_h.font.size = Pt(13)
            r_h.font.bold = True
            
            # Apply heading 1 color
            h1_color_hex = (custom_header.kpi_name_color if (custom_header and custom_header.kpi_name_color) else "#1e3a8a").lstrip("#")
            try:
                hr = int(h1_color_hex[0:2], 16)
                hg = int(h1_color_hex[2:4], 16)
                hb = int(h1_color_hex[4:6], 16)
            except Exception:
                hr, hg, hb = 0x1e, 0x3a, 0x8a
            r_h.font.color.rgb = RGBColor(hr, hg, hb)

            for f in sec.get("fields", []):
                if f.get("field_type") != "multi_line_items":
                    h2 = doc.add_heading(level=2)
                    from docx.enum.text import WD_TAB_ALIGNMENT
                    from docx.shared import Inches
                    h2.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=WD_TAB_ALIGNMENT.RIGHT)
                    
                    r_f = h2.add_run(f"{f['number']}. {f.get('field_name')}\t")
                    r_f.font.name = docx_font
                    r_f.font.size = Pt(11)
                    r_f.font.bold = True
                    r_f.font.color.rgb = RGBColor(0x37, 0x41, 0x55)
                    
                    val_str = clean_numeric_value_string(f.get("value"))
                    r_val = h2.add_run(val_str)
                    r_val.font.name = docx_font
                    r_val.font.size = Pt(scalar_font_size)
                    r_val.font.bold = scalar_bold
                    r_val.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
                else:
                    h2 = doc.add_heading(level=2)
                    r_f = h2.add_run(f"{f['number']}. {f.get('field_name')}")
                    r_f.font.name = docx_font
                    r_f.font.size = Pt(11)
                    r_f.font.bold = True
                    r_f.font.color.rgb = RGBColor(0x37, 0x41, 0x55)
                    
                    sub_fields = f.get("sub_fields", [])
                    value_items = f.get("value_items", [])

                    if sub_fields:
                        merged_headers = f.get("config", {}).get("merged_headers") or []

                        def style_cell_run(cell):
                            if cell.paragraphs and cell.paragraphs[0].runs:
                                run = cell.paragraphs[0].runs[0]
                                run.font.size = Pt(mli_font_size)
                                run.font.name = docx_font
                                run.font.bold = True
                            if cell.paragraphs:
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        if merged_headers:
                            table = doc.add_table(rows=2, cols=len(sub_fields) + 1)
                            table.style = 'Light Shading Accent 1'
                            hdr_rows = table.rows
                            
                            cols_keys = [sf.get("key") for sf in sub_fields]
                            col_indices = {key: idx + 1 for idx, key in enumerate(cols_keys)}
                            covered_cols = set()

                            # Vertical-merge Sr. No.
                            c_sr = hdr_rows[0].cells[0]
                            c_sr.text = "Sr. No."
                            style_cell_run(c_sr)
                            c_sr.merge(hdr_rows[1].cells[0])

                            # Write grouped headers
                            for group in merged_headers:
                                title = group.get("title")
                                start_key = group.get("start_key")
                                end_key = group.get("end_key")
                                if start_key in col_indices and end_key in col_indices:
                                    start_col = col_indices[start_key]
                                    end_col = col_indices[end_key]
                                    cell_g = hdr_rows[0].cells[start_col]
                                    cell_g.text = title
                                    style_cell_run(cell_g)
                                    cell_g.merge(hdr_rows[0].cells[end_col])
                                    for c in range(start_col, end_col + 1):
                                        covered_cols.add(c)

                            # Vertical-merge ungrouped columns and set row 1 headers for grouped columns
                            for idx, sf in enumerate(sub_fields):
                                c_idx = idx + 1
                                col_name = sf.get("name") or sf.get("key")
                                if c_idx not in covered_cols:
                                    cell_c = hdr_rows[0].cells[c_idx]
                                    cell_c.text = col_name
                                    style_cell_run(cell_c)
                                    cell_c.merge(hdr_rows[1].cells[c_idx])
                                else:
                                    cell_c = hdr_rows[1].cells[c_idx]
                                    cell_c.text = col_name
                                    style_cell_run(cell_c)
                        else:
                            table = doc.add_table(rows=1, cols=len(sub_fields) + 1)
                            table.style = 'Light Shading Accent 1'
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = "Sr. No."
                            style_cell_run(hdr_cells[0])

                            for col_idx, sf in enumerate(sub_fields):
                                hdr_cells[col_idx + 1].text = sf.get("name") or sf.get("key")
                                style_cell_run(hdr_cells[col_idx + 1])

                        col_alignments = (f.get("config") or {}).get("column_alignments") or {}
                        for item_idx, item in enumerate(value_items):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(item_idx + 1)
                            row_cells[0].paragraphs[0].runs[0].font.size = Pt(mli_font_size)
                            row_cells[0].paragraphs[0].runs[0].font.name = docx_font

                            for col_idx, sf in enumerate(sub_fields):
                                row_cells[col_idx + 1].text = clean_numeric_value_string(item.get(sf.get("key")))
                                if row_cells[col_idx + 1].paragraphs:
                                    p = row_cells[col_idx + 1].paragraphs[0]
                                    col_align = col_alignments.get(sf.get("key"))
                                    if col_align == "left":
                                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    elif col_align == "right":
                                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    elif col_align == "center":
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                                    if p.runs:
                                        run = p.runs[0]
                                        run.font.size = Pt(mli_font_size)
                                        run.font.name = docx_font

                        evaluated_footer = f.get("evaluated_footer_rows")
                        if evaluated_footer:
                            for f_row in evaluated_footer:
                                f_cells = f_row.get("cells", [])
                                total_sub_cols = len(sub_fields)
                                sum_colspan = sum(c.get("colspan", 1) for c in f_cells)
                                first_cell_extra = 1 if sum_colspan == total_sub_cols else 0

                                row_cells = table.add_row().cells
                                col_cursor = 0
                                for c_idx, cell in enumerate(f_cells):
                                    c_span = cell.get("colspan", 1) + (first_cell_extra if c_idx == 0 else 0)
                                    c_val = cell.get("value", "")
                                    c_align = cell.get("align", "left")
                                    c_bold = cell.get("bold", True)

                                    start_tc = row_cells[col_cursor]
                                    start_tc.text = c_val
                                    if start_tc.paragraphs:
                                        p = start_tc.paragraphs[0]
                                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_align == "right" else (WD_ALIGN_PARAGRAPH.CENTER if c_align == "center" else WD_ALIGN_PARAGRAPH.LEFT)
                                        if p.runs:
                                            p.runs[0].font.size = Pt(mli_font_size)
                                            p.runs[0].font.name = docx_font
                                            p.runs[0].font.bold = c_bold

                                    if c_span > 1 and (col_cursor + c_span - 1) < len(row_cells):
                                        end_tc = row_cells[col_cursor + c_span - 1]
                                        start_tc.merge(end_tc)
                                        col_cursor += c_span
                                    else:
                                        col_cursor += 1

        out_io = io.BytesIO()
        doc.save(out_io)
        return out_io.getvalue(), f"{clean_report_name}_{year}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif format == "pdf":
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from app.reports.service import NumberedCanvas
        from app.core.models import Organization, OrganizationBranding

        available_width = 756.0 if use_landscape else 576.0

        scalar_bold = data.get("scalar_bold", True)
        if scalar_bold is None:
            scalar_bold = True
        scalar_font_size = data.get("scalar_font_size", 11) or 11
        mli_font_size = data.get("mli_font_size", 10) or 10

        # Fetch custom header model
        custom_header = None
        if data.get("report_header_id"):
            from app.core.models import CustomReportHeader
            custom_header = await db.get(CustomReportHeader, data["report_header_id"])

        out_io = io.BytesIO()
        pdf_doc = SimpleDocTemplate(out_io, pagesize=landscape(letter) if use_landscape else letter, rightMargin=18, leftMargin=18, topMargin=36, bottomMargin=54)
        story = []

        # Setup custom fonts and styles
        styles = getSampleStyleSheet()
        
        font_name_bold = "Helvetica-Bold"
        font_name_regular = "Helvetica"
        font_name_italic = "Helvetica-Oblique"
        if custom_header:
            font_family = custom_header.font_family or "Helvetica"
            font_family_lower = font_family.lower()
            if font_family_lower in ("times-roman", "times new roman", "times"):
                font_name_bold = "Times-Bold"
                font_name_regular = "Times-Roman"
                font_name_italic = "Times-Italic"
            elif font_family_lower in ("courier", "courier new"):
                font_name_bold = "Courier-Bold"
                font_name_regular = "Courier"
                font_name_italic = "Courier-Oblique"
            elif font_family_lower == "helvetica":
                font_name_bold = "Helvetica-Bold"
                font_name_regular = "Helvetica"
                font_name_italic = "Helvetica-Oblique"
            else:
                from app.reports.service import register_extra_fonts
                register_extra_fonts()
                from reportlab.pdfbase import pdfmetrics
                registered = pdfmetrics.getRegisteredFontNames()
                chosen_font = "Helvetica"
                if font_family in registered:
                    chosen_font = font_family
                elif font_family_lower == "arial" and "Arial" in registered:
                    chosen_font = "Arial"
                elif font_family_lower == "georgia" and "Georgia" in registered:
                    chosen_font = "Georgia"
                elif font_family_lower == "verdana" and "Verdana" in registered:
                    chosen_font = "Verdana"
                elif font_family_lower == "calibri" and "Calibri" in registered:
                    chosen_font = "Calibri"
                elif font_family_lower == "garamond" and "Garamond" in registered:
                    chosen_font = "Garamond"
                
                font_name_regular = chosen_font
                font_name_bold = chosen_font
                if f"{chosen_font}-Bold" in registered:
                    font_name_bold = f"{chosen_font}-Bold"
                else:
                    font_name_bold = "Helvetica-Bold"
                
                font_name_italic = chosen_font
                if f"{chosen_font}-Italic" in registered:
                    font_name_italic = f"{chosen_font}-Italic"
                else:
                    font_name_italic = "Helvetica-Oblique"

        h1_color_hex = (custom_header.kpi_name_color if (custom_header and custom_header.kpi_name_color) else "#1e3a8a")

        title_style = ParagraphStyle(
            "CustomTitleStyle",
            parent=styles["Title"],
            fontName=font_name_bold,
            fontSize=20,
            textColor=colors.HexColor(h1_color_hex),
            spaceAfter=10
        )
        subtitle_style = ParagraphStyle(
            "CustomSubtitleStyle",
            parent=styles["Normal"],
            fontName=font_name_italic,
            fontSize=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=18
        )
        h1_style = ParagraphStyle(
            "CustomHeading1Style",
            parent=styles["Heading1"],
            fontName=font_name_bold,
            fontSize=13,
            textColor=colors.HexColor(h1_color_hex),
            spaceBefore=12,
            spaceAfter=6
        )
        h2_style = ParagraphStyle(
            "CustomHeading2Style",
            parent=styles["Heading2"],
            fontName=font_name_bold,
            fontSize=11,
            textColor=colors.HexColor("#374151"),
            spaceBefore=8,
            spaceAfter=4
        )
        h2_right_style = ParagraphStyle(
            "CustomHeading2RightStyle",
            parent=h2_style,
            fontName=font_name_bold if scalar_bold else font_name_regular,
            fontSize=scalar_font_size,
            alignment=2
        )
        body_style = ParagraphStyle(
            "CustomBodyStyle",
            parent=styles["Normal"],
            fontName=font_name_regular,
            fontSize=9.5,
            textColor=colors.HexColor("#111111"),
            leading=13
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        table_hdr_style = ParagraphStyle(
            "CustomTableHdrStyle",
            parent=styles["Normal"],
            fontName=font_name_bold,
            fontSize=mli_font_size,
            leading=mli_font_size + 3,
            textColor=colors.white,
            alignment=TA_CENTER
        )
        table_hdr_style_left = ParagraphStyle(
            "CustomTableHdrStyleLeft",
            parent=table_hdr_style,
            alignment=TA_LEFT
        )
        table_body_style_center = ParagraphStyle(
            "CustomTableBodyStyleCenter",
            parent=styles["Normal"],
            fontName=font_name_regular,
            fontSize=mli_font_size,
            leading=mli_font_size + 3,
            alignment=TA_CENTER
        )
        table_body_style_left = ParagraphStyle(
            "CustomTableBodyStyleLeft",
            parent=styles["Normal"],
            fontName=font_name_regular,
            fontSize=mli_font_size,
            leading=mli_font_size + 3,
            alignment=TA_LEFT
        )
        table_body_style = table_body_style_center

        # Render Header logo + headings if present
        if custom_header:
            from app.storage.service import get_file_stream as storage_get_file_stream
            from reportlab.platypus import Image
            from reportlab.lib.units import inch
            
            img = ""
            logo_w = 0
            try:
                logo_bytes = await storage_get_file_stream(db, org_id, custom_header.logo_path)
                from reportlab.lib.utils import ImageReader
                reader = ImageReader(io.BytesIO(logo_bytes))
                w, h = reader.getSize()
                aspect = float(h) / float(w) if w else 1.0
                logo_w = 1.2 * inch
                logo_h = logo_w * aspect
                if logo_h > 0.6 * inch:
                    logo_h = 0.6 * inch
                    logo_w = logo_h / aspect
                img = Image(io.BytesIO(logo_bytes), width=logo_w, height=logo_h)
            except Exception as ex:
                import logging
                logging.getLogger("custom_service").error(f"Failed to load logo 1: {ex}")
                pass

            img2 = ""
            logo_w2 = 0
            if custom_header.logo_path_2:
                try:
                    logo_bytes2 = await storage_get_file_stream(db, org_id, custom_header.logo_path_2)
                    from reportlab.lib.utils import ImageReader
                    reader2 = ImageReader(io.BytesIO(logo_bytes2))
                    w2, h2 = reader2.getSize()
                    aspect2 = float(h2) / float(w2) if w2 else 1.0
                    logo_w2 = 1.2 * inch
                    logo_h2 = logo_w2 * aspect2
                    if logo_h2 > 0.6 * inch:
                        logo_h2 = 0.6 * inch
                        logo_w2 = logo_h2 / aspect2
                    img2 = Image(io.BytesIO(logo_bytes2), width=logo_w2, height=logo_h2)
                except Exception as ex:
                    import logging
                    logging.getLogger("custom_service").error(f"Failed to load logo 2: {ex}")
                    pass

            header_align_str = (custom_header.text_align or "center").lower()
            if header_align_str == "left":
                align_code = TA_LEFT
            elif header_align_str == "right":
                align_code = TA_RIGHT
            elif header_align_str == "justify":
                align_code = TA_JUSTIFY
            else:
                align_code = TA_CENTER

            desired_p_fs = custom_header.font_size or 16
            w1 = (logo_w + 4) if img else 0
            w2 = (logo_w2 + 4) if img2 else 0
            mid_w = available_width - w1 - w2
            pdf_main_fs = calc_auto_header_font_size(custom_header.main_heading, mid_w, desired_p_fs)

            heading_paragraph_style = ParagraphStyle(
                "CustomHeaderHeadings",
                parent=styles["Normal"],
                leading=int(pdf_main_fs * 1.25),
                alignment=align_code
            )

            sub_align_str = (custom_header.sub_text_align or custom_header.text_align or "center").lower()
            if sub_align_str == "left":
                sub_align_code = TA_LEFT
            elif sub_align_str == "right":
                sub_align_code = TA_RIGHT
            elif sub_align_str == "justify":
                sub_align_code = TA_JUSTIFY
            else:
                sub_align_code = TA_CENTER

            sub_font_family = custom_header.sub_font_family or "Helvetica"
            sub_font_lower = sub_font_family.lower()
            sub_font_name_italic = "Helvetica-Oblique"
            if sub_font_lower in ("times-roman", "times new roman", "times"):
                sub_font_name_italic = "Times-Italic"
            elif sub_font_lower in ("courier", "courier new"):
                sub_font_name_italic = "Courier-Oblique"
            elif sub_font_lower == "helvetica":
                sub_font_name_italic = "Helvetica-Oblique"
            else:
                from reportlab.pdfbase import pdfmetrics
                registered = pdfmetrics.getRegisteredFontNames()
                chosen_sub = "Helvetica"
                if sub_font_family in registered:
                    chosen_sub = sub_font_family
                elif sub_font_lower == "arial" and "Arial" in registered:
                    chosen_sub = "Arial"
                sub_font_name_italic = chosen_sub
                if f"{chosen_sub}-Italic" in registered:
                    sub_font_name_italic = f"{chosen_sub}-Italic"
                else:
                    sub_font_name_italic = "Helvetica-Oblique"

            sub_font_size = custom_header.sub_font_size or max(9, int(pdf_main_fs * 0.65))
            sub_color_hex = custom_header.sub_text_color or "#4b5563"

            sub_heading_paragraph_style = ParagraphStyle(
                "CustomHeaderSubHeadings",
                parent=styles["Normal"],
                leading=int(sub_font_size * 1.25),
                alignment=sub_align_code,
                spaceBefore=2
            )

            main_p = Paragraph(f"<font face='{font_name_bold}' size='{pdf_main_fs}' color='{custom_header.text_color or '#1e3a8a'}'><b>{html.escape(custom_header.main_heading)}</b></font>", heading_paragraph_style)
            
            header_elements = [main_p]
            if custom_header.sub_heading:
                sub_p = Paragraph(f"<font face='{sub_font_name_italic}' size='{sub_font_size}' color='{sub_color_hex}'><i>{html.escape(custom_header.sub_heading)}</i></font>", sub_heading_paragraph_style)
                header_elements.append(sub_p)

            if img and img2:
                if use_landscape:
                    from reportlab.pdfbase import pdfmetrics
                    text_w = pdfmetrics.stringWidth(custom_header.main_heading, font_name_bold, pdf_main_fs)
                    sub_w = pdfmetrics.stringWidth(custom_header.sub_heading, sub_font_name_italic, sub_font_size) if custom_header.sub_heading else 0
                    mid_w = min(available_width - w1 - w2, max(text_w, sub_w) + 6)
                header_table = Table(
                    [[img, header_elements, img2]],
                    colWidths=[w1, mid_w, w2],
                    hAlign='CENTER'
                )
                header_table.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (0, 0), (0, 0), "LEFT"),
                    ("ALIGN", (2, 0), (2, 0), "RIGHT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 1),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 1),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]))
            elif img:
                w1 = logo_w + 8
                mid_w = available_width - w1
                if use_landscape:
                    from reportlab.pdfbase import pdfmetrics
                    text_w = pdfmetrics.stringWidth(custom_header.main_heading, font_name_bold, pdf_main_fs)
                    sub_w = pdfmetrics.stringWidth(custom_header.sub_heading, sub_font_name_italic, sub_font_size) if custom_header.sub_heading else 0
                    mid_w = min(available_width - w1, max(text_w, sub_w) + 6)
                header_table = Table(
                    [[img, header_elements]],
                    colWidths=[w1, mid_w],
                    hAlign='CENTER'
                )
                header_table.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (0, 0), (0, 0), "RIGHT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]))
            elif img2:
                w2 = logo_w2 + 8
                mid_w = available_width - w2
                if use_landscape:
                    from reportlab.pdfbase import pdfmetrics
                    text_w = pdfmetrics.stringWidth(custom_header.main_heading, font_name_bold, pdf_main_fs)
                    sub_w = pdfmetrics.stringWidth(custom_header.sub_heading, sub_font_name_italic, sub_font_size) if custom_header.sub_heading else 0
                    mid_w = min(available_width - w2, max(text_w, sub_w) + 6)
                header_table = Table(
                    [[header_elements, img2]],
                    colWidths=[mid_w, w2],
                    hAlign='CENTER'
                )
                header_table.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (1, 0), (1, 0), "LEFT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]))
            else:
                header_table = Table(
                    [[header_elements]],
                    colWidths=[available_width],
                    hAlign='CENTER'
                )
            story.append(header_table)
            period_p = Paragraph(f"<font face='{font_name_bold}' size='9' color='#475569'><b>{html.escape(period_info)}</b></font>", h2_right_style)
            story.append(period_p)
            story.append(Spacer(1, 10))
        else:
            period_p = Paragraph(f"<font face='{font_name_bold}' size='9' color='#475569'><b>{html.escape(period_info)}</b></font>", h2_right_style)
            story.append(period_p)
            story.append(Spacer(1, 10))

        # Configurable report name just below header (only if more than 1 KPI)
        if data.get("unique_kpi_count", 0) > 1 and data.get("show_report_name", True):
            story.append(Paragraph(report_name, title_style))
            story.append(Paragraph(f"Academic Year: {year}", subtitle_style))

        for s_idx, sec in enumerate(data.get("sections", [])):
            sec_name = sec.get("custom_header") or sec.get("kpi_name", f"Section {s_idx+1}")
            story.append(Paragraph(f"{sec['number']}. {sec_name}", h1_style))
            story.append(Spacer(1, 4))

            for f in sec.get("fields", []):
                if f.get("field_type") != "multi_line_items":
                    val = f.get("value")
                    val_str = clean_numeric_value_string(val)
                    from reportlab.platypus import Table, TableStyle
                    t = Table(
                        [[Paragraph(f"{f['number']}. {f.get('field_name')}", h2_style), Paragraph(f"<b>{val_str}</b>", h2_right_style)]],
                        colWidths=[450, 90]
                    )
                    t.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                        ('ALIGN', (0,0), (0,0), 'LEFT'),
                        ('ALIGN', (1,0), (1,0), 'RIGHT'),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                        ('TOPPADDING', (0,0), (-1,-1), 1),
                        ('LEFTPADDING', (0,0), (-1,-1), 0),
                        ('RIGHTPADDING', (0,0), (-1,-1), 0),
                        ('RIGHTPADDING', (1,0), (1,0), 25),
                        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor("#e5e7eb")),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 6))
                else:
                    # Fields as Heading 2s
                    story.append(Paragraph(f"{f['number']}. {f.get('field_name')}", h2_style))
                    story.append(Spacer(1, 2))
                    sub_fields = f.get("sub_fields", [])
                    value_items = f.get("value_items", [])

                    if sub_fields:
                        col_count = len(sub_fields)
                        if col_count > 14:
                            tbl_font_size = 7.0
                            tbl_padding = 1.5
                            sr_no_width = 24.0
                        elif col_count > 10:
                            tbl_font_size = 8.0
                            tbl_padding = 2.0
                            sr_no_width = 28.0
                        else:
                            tbl_font_size = float(mli_font_size)
                            tbl_padding = 3.0
                            sr_no_width = 40.0

                        cur_hdr_style = ParagraphStyle(
                            f"CustomTblHdrStyle_{col_count}",
                            parent=styles["Normal"],
                            fontName=font_name_bold,
                            fontSize=tbl_font_size,
                            leading=tbl_font_size + 1.5,
                            textColor=colors.white,
                            alignment=TA_CENTER
                        )
                        cur_hdr_style_left = ParagraphStyle(
                            f"CustomTblHdrStyleLeft_{col_count}",
                            parent=cur_hdr_style,
                            alignment=TA_LEFT
                        )
                        cur_hdr_style_right = ParagraphStyle(
                            f"CustomTblHdrStyleRight_{col_count}",
                            parent=cur_hdr_style,
                            alignment=TA_RIGHT
                        )
                        cur_body_style_center = ParagraphStyle(
                            f"CustomTblBodyStyleCenter_{col_count}",
                            parent=styles["Normal"],
                            fontName=font_name_regular,
                            fontSize=tbl_font_size,
                            leading=tbl_font_size + 1.5,
                            alignment=TA_CENTER
                        )
                        cur_body_style_left = ParagraphStyle(
                            f"CustomTblBodyStyleLeft_{col_count}",
                            parent=styles["Normal"],
                            fontName=font_name_regular,
                            fontSize=tbl_font_size,
                            leading=tbl_font_size + 1.5,
                            alignment=TA_LEFT
                        )
                        cur_body_style_right = ParagraphStyle(
                            f"CustomTblBodyStyleRight_{col_count}",
                            parent=styles["Normal"],
                            fontName=font_name_regular,
                            fontSize=tbl_font_size,
                            leading=tbl_font_size + 1.5,
                            alignment=TA_RIGHT
                        )

                        # Expand printable table width into side margins for wide tables (> 10 columns)
                        max_table_printable_w = float(available_width)
                        available_table_w = max_table_printable_w - sr_no_width
                        min_col_width = 18.0 if col_count > 14 else (28.0 if col_count > 8 else 40.0)
                        
                        col_chars = []
                        for s_idx, sf in enumerate(sub_fields):
                            name = sf.get("name") or sf.get("key")
                            k_lower = sf.get("key", "").lower()
                            if s_idx == 0 or "department" in k_lower or "name" in k_lower or "title" in k_lower:
                                w = 28 if col_count > 10 else 35
                            elif col_count > 8:
                                if "faculty_who_submitted" in k_lower or "faculty_who" in k_lower:
                                    w = 18
                                else:
                                    w = max(10, min(len(name), 16))
                            else:
                                w = min(len(name), 25)
                            col_chars.append(w)

                        for item in value_items:
                            for idx, sf in enumerate(sub_fields):
                                val = item.get(sf.get("key"))
                                val_str = clean_numeric_value_string(val)
                                max_val_len = 15 if col_count > 8 else 60
                                col_chars[idx] = max(col_chars[idx], min(len(val_str), max_val_len))
                                
                        custom_widths = f.get("config", {}).get("column_widths") or {}
                        if custom_widths:
                            configured_sr = float(custom_widths.get("S.No", 60)) * 0.75
                            configured_cols = []
                            for sf in sub_fields:
                                cw = custom_widths.get(sf.get("key")) or custom_widths.get(sf.get("name"))
                                if cw:
                                    configured_cols.append(float(cw) * 0.75)
                                else:
                                    configured_cols.append(26.0 if col_count > 14 else 45.0)
                            raw_widths = [configured_sr] + configured_cols
                            sum_raw = sum(raw_widths)
                            scale = max_table_printable_w / sum_raw if sum_raw > 0 else 1.0
                            col_widths = [w * scale for w in raw_widths]
                        else:
                            if col_count > 14:
                                dept_w = min(120.0, max(85.0, available_table_w * 0.17))
                                remain_w = available_table_w - dept_w
                                per_col_w = remain_w / (col_count - 1)
                                col_widths = [sr_no_width, dept_w] + [per_col_w] * (col_count - 1)
                            else:
                                total_chars = sum(col_chars) or col_count
                                raw_widths = [max(min_col_width, available_table_w * (c / total_chars)) for c in col_chars]
                                scale = available_table_w / sum(raw_widths) if sum(raw_widths) > 0 else 1.0
                                col_widths = [sr_no_width] + [w * scale for w in raw_widths]

                        merged_headers = f.get("config", {}).get("merged_headers") or []
                        col_alignments = (f.get("config") or {}).get("column_alignments") or {}

                        def get_pdf_hdr_style(col_key, idx):
                            align = col_alignments.get(col_key)
                            if align == "left":
                                return cur_hdr_style_left
                            elif align == "right":
                                return cur_hdr_style_right
                            elif align == "center":
                                return cur_hdr_style
                            else:
                                return cur_hdr_style_left if idx == 0 else cur_hdr_style

                        def get_pdf_body_style(col_key, idx):
                            align = col_alignments.get(col_key)
                            if align == "left":
                                return cur_body_style_left
                            elif align == "right":
                                return cur_body_style_right
                            elif align == "center":
                                return cur_body_style_center
                            else:
                                return cur_body_style_left if idx == 0 else cur_body_style_center
                        
                        t_style_cmds = [
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E5E7EB")),
                            ("TOPPADDING", (0, 0), (-1, -1), tbl_padding),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), tbl_padding),
                            ("LEFTPADDING", (0, 0), (-1, -1), tbl_padding),
                            ("RIGHTPADDING", (0, 0), (-1, -1), tbl_padding),
                        ]

                        if merged_headers:
                            cols_keys = [sf.get("key") for sf in sub_fields]
                            col_indices = {key: idx + 1 for idx, key in enumerate(cols_keys)}
                            covered_cols = set()

                            row_0 = [Paragraph("", cur_hdr_style) for _ in range(len(sub_fields) + 1)]
                            row_1 = [Paragraph("Sr. No.", cur_hdr_style)] + [Paragraph(sf.get("name") or sf.get("key"), get_pdf_hdr_style(sf.get("key"), s_idx)) for s_idx, sf in enumerate(sub_fields)]
                            
                            t_style_cmds.append(("BACKGROUND", (0, 0), (-1, 1), colors.HexColor(h1_color_hex)))
                            
                            # Vertical merge Sr. No.
                            row_0[0] = Paragraph("Sr. No.", cur_hdr_style)
                            t_style_cmds.append(("SPAN", (0, 0), (0, 1)))

                            # Write grouped headers
                            for group in merged_headers:
                                title = group.get("title")
                                start_key = group.get("start_key")
                                end_key = group.get("end_key")
                                if start_key in col_indices and end_key in col_indices:
                                    start_col = col_indices[start_key]
                                    end_col = col_indices[end_key]
                                    row_0[start_col] = Paragraph(f"<b>{title}</b>", cur_hdr_style)
                                    t_style_cmds.append(("SPAN", (start_col, 0), (end_col, 0)))
                                    t_style_cmds.append(("ALIGN", (start_col, 0), (end_col, 0), "CENTER"))
                                    for c in range(start_col, end_col + 1):
                                        covered_cols.add(c)

                            # Vertical merge ungrouped columns
                            for idx, sf in enumerate(sub_fields):
                                c_idx = idx + 1
                                if c_idx not in covered_cols:
                                    hdr_st = get_pdf_hdr_style(sf.get("key"), idx)
                                    row_0[c_idx] = Paragraph(sf.get("name") or sf.get("key"), hdr_st)
                                    t_style_cmds.append(("SPAN", (c_idx, 0), (c_idx, 1)))

                            pdf_table_data = [row_0, row_1]
                        else:
                            t_style_cmds.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(h1_color_hex)))
                            hdr_row = [Paragraph("Sr. No.", cur_hdr_style)] + [Paragraph(sf.get("name") or sf.get("key"), get_pdf_hdr_style(sf.get("key"), s_idx)) for s_idx, sf in enumerate(sub_fields)]
                            pdf_table_data = [hdr_row]

                        for item_idx, item in enumerate(value_items):
                            row = [Paragraph(str(item_idx + 1), cur_body_style_center)]
                            for s_idx, sf in enumerate(sub_fields):
                                val = item.get(sf.get("key"))
                                val_str = clean_numeric_value_string(val)
                                cell_st = get_pdf_body_style(sf.get("key"), s_idx)
                                row.append(Paragraph(val_str, cell_st))
                            pdf_table_data.append(row)

                        evaluated_footer = f.get("evaluated_footer_rows")
                        if evaluated_footer:
                            for f_row in evaluated_footer:
                                cur_row_idx = len(pdf_table_data)
                                cells = f_row.get("cells", [])
                                total_sub_cols = len(sub_fields)
                                sum_colspan = sum(c.get("colspan", 1) for c in cells)
                                first_cell_extra = 1 if sum_colspan == total_sub_cols else 0

                                pdf_f_row = []
                                col_cursor = 0
                                for c_idx, cell in enumerate(cells):
                                    c_span = cell.get("colspan", 1) + (first_cell_extra if c_idx == 0 else 0)
                                    c_val = str(cell.get("value", ""))
                                    c_align = cell.get("align", "left")
                                    c_bold = cell.get("bold", True)

                                    align_code = TA_RIGHT if c_align == "right" else (TA_CENTER if c_align == "center" else TA_LEFT)
                                    footer_style = ParagraphStyle(
                                        f"PDFFooter_{cur_row_idx}_{c_idx}",
                                        parent=styles["Normal"],
                                        fontName=font_name_bold if c_bold else font_name_regular,
                                        fontSize=mli_font_size,
                                        leading=mli_font_size + 3,
                                        alignment=align_code
                                    )

                                    pdf_f_row.append(Paragraph(c_val, footer_style))

                                    if c_span > 1:
                                        start_c = col_cursor
                                        end_c = min(col_cursor + c_span - 1, total_sub_cols)
                                        t_style_cmds.append(("SPAN", (start_c, cur_row_idx), (end_c, cur_row_idx)))
                                        for _ in range(c_span - 1):
                                            pdf_f_row.append(Paragraph("", footer_style))
                                        col_cursor = end_c + 1
                                    else:
                                        col_cursor += 1

                                pdf_table_data.append(pdf_f_row)
                                t_style_cmds.append(("BACKGROUND", (0, cur_row_idx), (-1, cur_row_idx), colors.HexColor("#f8fafc")))

                        t_mli = Table(pdf_table_data, colWidths=col_widths, hAlign='CENTER')
                        t_mli.setStyle(TableStyle(t_style_cmds))
                        story.append(t_mli)
                        story.append(Spacer(1, 8))

        # Fetch organization and custom branding
        org = await db.get(Organization, org_id)
        org_branding = (await db.execute(
            select(OrganizationBranding).where(OrganizationBranding.organization_id == org_id)
        )).scalar_one_or_none()

        footer_confidentiality = "Confidential Document"
        is_custom_brand = False
        if org_branding and org_branding.footer_label:
            footer_confidentiality = org_branding.footer_label
            is_custom_brand = True

        def make_canvas(*args, **kwargs):
            c = NumberedCanvas(*args, **kwargs)
            c.organization_name = org.name if org else ""
            c.include_date = True
            c.confidentiality_text = footer_confidentiality
            c.is_custom_branding = is_custom_brand
            return c

        pdf_doc.build(story, canvasmaker=make_canvas)
        return out_io.getvalue(), f"{clean_report_name}_{year}.pdf", "application/pdf"

    raise ValueError("Invalid format: " + format)


# ----------------------------------------------------
# Optimized Caching & Background Task Generation
# ----------------------------------------------------

class CustomReportCache:
    def __init__(self):
        self._cache = {}

    def get(self, key):
        return self._cache.get(key)

    def set(self, key, value):
        # Prevent unbounded growth
        if len(self._cache) > 256:
            self._cache.clear()
        self._cache[key] = value

    def invalidate_report(self, report_id: int):
        keys_to_del = [k for k in self._cache if isinstance(k, tuple) and k[0] == report_id]
        for k in keys_to_del:
            self._cache.pop(k, None)

    def invalidate_user(self, user_id: int):
        keys_to_del = [k for k in self._cache if isinstance(k, tuple) and len(k) == 9 and k[-1] == user_id]
        for k in keys_to_del:
            self._cache.pop(k, None)

    def invalidate_all(self):
        self._cache.clear()

CUSTOM_REPORT_CACHE = CustomReportCache()

# Thread-safe/process-safe in-memory task registry
REPORT_TASKS = {} # task_id -> {"status": "processing"/"completed"/"failed", "progress": 0, "result": None, "error": None}


async def run_background_generation_task(task_id: str, id: int, org_id: int, year: int | None):
    REPORT_TASKS[task_id] = {
        "status": "processing",
        "progress": 0,
        "result": None,
        "error": None
    }
    try:
        from app.core.database import AsyncSessionLocal
        async with AsyncSessionLocal() as db:
            def update_progress(p):
                REPORT_TASKS[task_id]["progress"] = p

            # Run full generation in background (preview=False)
            data = await generate_custom_report_data(
                db, id, org_id, year=year, include_drafts=False, preview=False, on_progress=update_progress
            )
            if not data:
                REPORT_TASKS[task_id]["status"] = "failed"
                REPORT_TASKS[task_id]["error"] = "Report not found or empty"
                return

            html = await render_custom_report_html(
                db, id, org_id, year=year, include_drafts=False, report_data=data
            )
            if html is not None:
                data["rendered_html"] = html

            REPORT_TASKS[task_id]["status"] = "completed"
            REPORT_TASKS[task_id]["progress"] = 100
            REPORT_TASKS[task_id]["result"] = data
    except Exception as e:
        logger.exception("Error in background generation task: %s", task_id)
        REPORT_TASKS[task_id]["status"] = "failed"
        REPORT_TASKS[task_id]["error"] = str(e)


async def stream_custom_report_data(
    db: AsyncSession,
    id: int,
    org_id: int,
    year: str | int | None = None,
    by_default: bool = False,
    period_type: str | None = None,
):
    custom_report = await get_custom_report(db, id, org_id)
    if not custom_report:
        yield {"type": "error", "message": "Report not found"}
        return

    selected_period = year
    if custom_report and getattr(custom_report, "fetch_data_with_date", False):
        if not selected_period or by_default or selected_period == "by_default":
            selected_period = "2025/26"
            by_default = False

    date_range = None
    entry_start_year: int | None = None  # calendar start year of the period (e.g. 2026 for "2026/27")
    if custom_report and getattr(custom_report, "fetch_data_with_date", False) and selected_period and not by_default:
        org = (await db.execute(select(Organization).where(Organization.id == org_id))).scalar_one_or_none()
        if org:
            try:
                from app.widget_data.service import get_widget_date_col_key, resolve_date_range_for_period
                start_date, end_date, entry_year = resolve_date_range_for_period(org, str(selected_period), period_type)
                entry_start_year = start_date.year  # e.g. 2026 for July-2026 to June-2027
                year = entry_year  # e.g. 2027 (the year entries are stored under)
                date_range = (start_date, end_date)
            except Exception:
                pass

    yr = _parse_year_int(year)
    yr_display = selected_period if (custom_report and getattr(custom_report, "fetch_data_with_date", False) and selected_period) else yr

    # Yield metadata
    yield {
        "type": "metadata",
        "custom_report_id": custom_report.id,
        "custom_report_name": custom_report.name,
        "custom_report_description": custom_report.description,
        "year": yr_display,
    }

    # 1. Identify all referenced KPIs
    referenced_kpi_ids = set()
    for sec in custom_report.sections:
        referenced_kpi_ids.add(sec.kpi_id)
        for f in sec.fields:
            referenced_kpi_ids.add(f.kpi_field.kpi_id)

    if not referenced_kpi_ids:
        yield {"type": "done"}
        return

    # Load KPIs, Org Config, and Entries
    kpis_result = await db.execute(
        select(KPI)
        .where(KPI.id.in_(referenced_kpi_ids))
        .options(
            selectinload(KPI.fields).selectinload(KPIField.sub_fields),
            noload(KPI.organization),
            noload(KPI.domain)
        )
    )
    kpis_by_id = {k.id: k for k in kpis_result.scalars().unique().all()}

    org = await db.get(Organization, org_id)
    org_td = TimeDimension(getattr(org, "time_dimension", None) or "yearly") if org else TimeDimension.YEARLY

    all_entries_filters = [
        KPIEntry.organization_id == org_id,
        KPIEntry.kpi_id.in_(referenced_kpi_ids),
        KPIEntry.is_draft == False,
    ]
    if not date_range:
        all_entries_filters.append(KPIEntry.year == yr)
    entries_result = await db.execute(
        select(KPIEntry)
        .where(*all_entries_filters)
        .options(selectinload(KPIEntry.field_values))
    )
    all_entries_list = list(entries_result.scalars().all())

    # Fetch counts of multi-line rows for all these entries to prioritize populated entries
    entry_ids = [e.id for e in all_entries_list]
    mli_counts = {}
    if entry_ids:
        from sqlalchemy import func
        counts_res = await db.execute(
            select(KpiMultiLineRow.entry_id, func.count(KpiMultiLineRow.id))
            .where(KpiMultiLineRow.entry_id.in_(entry_ids))
            .group_by(KpiMultiLineRow.entry_id)
        )
        mli_counts = {r[0]: r[1] for r in counts_res.all()}

    entries_by_kpi = {}
    for entry in all_entries_list:
        entries_by_kpi.setdefault(entry.kpi_id, []).append(entry)

    # Evaluate scalar and formula fields
    kpi_evaluated_data = {}
    for kid, kpi in kpis_by_id.items():
        fields_to_include = sorted(list(kpi.fields or []), key=lambda f: (f.sort_order, f.id))
        kpi_td_raw = getattr(kpi, "time_dimension", None)
        kpi_td = TimeDimension(kpi_td_raw) if kpi_td_raw else None
        effective_td = effective_kpi_time_dimension(kpi_td, org_td)

        all_entries = entries_by_kpi.get(kpi.id, [])
        if date_range:
            real_entry = None
            if all_entries:
                def _score_entry(e):
                    # Check if entry has any populated multi-line rows
                    rows_count = mli_counts.get(e.id, 0)
                    # Check if entry has any populated scalar field values
                    has_scalar_data = False
                    for fv in (e.field_values or []):
                        if (fv.value_number is not None and fv.value_number != 0) or \
                           (fv.value_text is not None and fv.value_text != "") or \
                           fv.value_boolean is not None or \
                           fv.value_date is not None or \
                           fv.value_json is not None:
                            has_scalar_data = True
                            break
                    
                    data_score = 1 if (rows_count > 0 or has_scalar_data) else 0
                    
                    if e.year == yr:
                        year_score = 3
                    elif entry_start_year is not None and e.year == entry_start_year:
                        year_score = 2
                    else:
                        year_score = 1 - abs(e.year - yr) * 0.1
                        
                    return (data_score, year_score, e.year)

                all_entries_sorted = sorted(
                    all_entries,
                    key=_score_entry,
                    reverse=True
                )
                real_entry = all_entries_sorted[0]
            if real_entry:
                entries_sorted = [real_entry]
            else:
                entries_sorted = []
        else:
            entries_sorted = sorted(
                all_entries,
                key=lambda e: period_key_sort_order(getattr(e, "period_key", "") or "", effective_td),
            )
            if len(entries_sorted) > 1:
                entries_sorted = [entries_sorted[-1]]

        need_cross_kpi = _formulas_need_other_kpi_values(fields_to_include)
        cross_kpi_year = yr
        if date_range and entries_sorted:
            cross_kpi_year = entries_sorted[0].year
        other_kpi_values = (
            await _load_other_kpi_values(db, cross_kpi_year, org_id, kpi.id)
            if entries_sorted and need_cross_kpi
            else {}
        )

        evaluated_fields = {}
        if not entries_sorted:
            _NO_DATA_PLACEHOLDER = "No data entered"
            for f in fields_to_include:
                if f.field_type == FieldType.multi_line_items:
                    continue
                evaluated_fields[f.key] = {
                    "field_key": f.key,
                    "field_name": f.name,
                    "value": _NO_DATA_PLACEHOLDER,
                    "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                }
        else:
            entry = entries_sorted[0]
            fv_by_field = {fv.field_id: fv for fv in entry.field_values}
            value_by_key = {}
            multi_line_items_data = {}

            # Load multi-line rows for all multi-line fields to evaluate formulas
            ml_fields = [f for f in fields_to_include if f.field_type == FieldType.multi_line_items]
            for mf in ml_fields:
                mf_date_range = None
                if date_range:
                    config = getattr(custom_report, "date_fetching_config", None) or {}
                    date_col_key = get_widget_date_col_key(config, kpi.id, mf.key, mf)
                    if date_col_key:
                        start_date, end_date = date_range
                        mf_date_range = (start_date, end_date, str(date_col_key))

                target_entry_ids = [e.id for e in all_entries if e.id] if (date_range and mf_date_range) else [entry.id]
                ml_rows = await _load_multi_line_items_rows_batch(
                    db, entry_ids=target_entry_ids, field=mf, date_range=mf_date_range, custom_report_id=id, current_user=current_user
                )
                if date_range and mf_date_range:
                    combined_rows = []
                    for rows_list in ml_rows.values():
                        combined_rows.extend(rows_list)
                    combined_rows = _deduplicate_rows(combined_rows)
                    multi_line_items_data[mf.key] = combined_rows
                else:
                    multi_line_items_data[mf.key] = ml_rows.get(entry.id, [])

            for f in fields_to_include:
                if f.field_type == FieldType.formula or f.field_type == FieldType.multi_line_items:
                    continue
                fv = fv_by_field.get(f.id)
                val = None
                if fv:
                    if fv.value_date is not None:
                        val = fv.value_date.isoformat() if hasattr(fv.value_date, "isoformat") else str(fv.value_date)
                    elif fv.value_text is not None:
                        val = fv.value_text
                    elif fv.value_number is not None:
                        val = fv.value_number
                    elif fv.value_json is not None:
                        val = fv.value_json
                    elif fv.value_boolean is not None:
                        val = fv.value_boolean
                    if f.field_type == FieldType.number and fv.value_number is not None:
                        value_by_key[f.key] = fv.value_number
                evaluated_fields[f.key] = {
                    "field_key": f.key,
                    "field_name": f.name,
                    "value": val,
                    "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                }
                if val is not None and f.field_type == FieldType.number:
                    value_by_key[f.key] = val

            for f in fields_to_include:
                if f.field_type == FieldType.formula:
                    fv_formula = fv_by_field.get(f.id)
                    if fv_formula and fv_formula.value_number is not None:
                        try:
                            value_by_key[f.key] = float(fv_formula.value_number)
                        except (TypeError, ValueError):
                            pass

            for f in fields_to_include:
                if f.field_type == FieldType.formula and f.formula_expression:
                    computed = evaluate_formula(
                        f.formula_expression,
                        value_by_key,
                        multi_line_items_data,
                        other_kpi_values,
                    )
                    if computed is None:
                        fv_formula = fv_by_field.get(f.id)
                        if fv_formula and fv_formula.value_number is not None:
                            computed = fv_formula.value_number
                    evaluated_fields[f.key] = {
                        "field_key": f.key,
                        "field_name": f.name,
                        "value": computed,
                        "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                    }
        kpi_evaluated_data[kpi.id] = evaluated_fields

    # Yield structure
    sections_structure = []
    for s_idx, sec in enumerate(custom_report.sections):
        sec_num = str(s_idx + 1)
        sec_header = sec.custom_header or (sec.kpi.name if sec.kpi else f"Section {sec_num}")

        fields_struct = []
        for f_idx, f in enumerate(sec.fields):
            f_num = f"{sec_num}.{f_idx + 1}"
            kfield = f.kpi_field

            cfg = getattr(f, "config", None) or {}
            custom_name = cfg.get("custom_name") or kfield.name

            field_payload = {
                "id": f.id,
                "kpi_field_id": f.kpi_field_id,
                "field_key": kfield.key,
                "field_name": custom_name,
                "field_type": kfield.field_type.value if hasattr(kfield.field_type, "value") else str(kfield.field_type),
                "number": f_num,
                "config": cfg,
            }
            if kfield.field_type == FieldType.multi_line_items:
                sub_fields_orm = getattr(kfield, "sub_fields") or []
                cfg = getattr(f, "config", None) or {}
                visible_keys = cfg.get("selected_columns")
                if visible_keys is None:
                    # Default to first 5 columns if not explicitly configured to prevent ReportLab crash on export
                    visible_keys = [sf.key for sf in sub_fields_orm][:5]

                all_subs = [{"key": sf.key, "name": getattr(sf, "name", sf.key), "field_type": sf.field_type.value if hasattr(sf.field_type, "value") else str(sf.field_type)} for sf in sub_fields_orm]
                key_to_sf = {sf["key"]: sf for sf in all_subs}

                field_payload["sub_fields"] = [key_to_sf[k] for k in visible_keys if k in key_to_sf]
                field_payload["sub_field_keys"] = [k for k in visible_keys]
            else:
                kpi_data = kpi_evaluated_data.get(kfield.kpi_id, {})
                field_eval = kpi_data.get(kfield.key, {})
                field_payload["value"] = field_eval.get("value")

            fields_struct.append(field_payload)

        sections_structure.append({
            "id": sec.id,
            "kpi_id": sec.kpi_id,
            "custom_header": sec_header,
            "number": sec_num,
            "fields": fields_struct
        })

    yield {
        "type": "structure",
        "sections": sections_structure
    }

    # Stream table rows in chunks
    for s_idx, sec in enumerate(custom_report.sections):
        for f_idx, f in enumerate(sec.fields):
            kfield = f.kpi_field
            if kfield.field_type != FieldType.multi_line_items:
                continue

            all_entries = entries_by_kpi.get(kfield.kpi_id, [])
            if not all_entries:
                yield {
                    "type": "table_rows",
                    "field_id": f.id,
                    "value_items": [],
                    "total_count": 0,
                    "offset": 0,
                    "done": True
                }
                continue

            kpi_td_raw = getattr(kfield.kpi, "time_dimension", None)
            kpi_td = TimeDimension(kpi_td_raw) if kpi_td_raw else None
            effective_td = effective_kpi_time_dimension(kpi_td, org_td)
            entries_sorted = sorted(
                all_entries,
                key=lambda e: period_key_sort_order(getattr(e, "period_key", "") or "", effective_td),
            )
            if len(entries_sorted) > 1:
                entries_sorted = [entries_sorted[-1]]
            entry = entries_sorted[0]

            mf_date_range = None
            if date_range:
                config = getattr(custom_report, "date_fetching_config", None) or {}
                date_col_key = get_widget_date_col_key(config, kfield.kpi_id, kfield.key, kfield)
                if date_col_key:
                    start_date, end_date = date_range
                    mf_date_range = (start_date, end_date, str(date_col_key))

            if date_range and mf_date_range:
                target_entry_ids = [e.id for e in all_entries if e.id]
                batch_res = await _load_multi_line_items_rows_batch(
                    db, entry_ids=target_entry_ids, field=kfield, date_range=mf_date_range
                )
                combined_rows = []
                for rows_list in batch_res.values():
                    combined_rows.extend(rows_list)
                combined_rows = _deduplicate_rows(combined_rows)
                total_count = len(combined_rows)
                
                yield {
                    "type": "table_meta",
                    "field_id": f.id,
                    "total_count": total_count
                }
                
                chunk_size = 1000
                for offset in range(0, total_count, chunk_size):
                    chunk_rows = combined_rows[offset : offset + chunk_size]
            else:
                count_stmt = (
                    select(func.count(KpiMultiLineRow.id))
                    .where(
                        KpiMultiLineRow.entry_id == entry.id,
                        KpiMultiLineRow.field_id == kfield.id,
                      )
                )
                total_count = (await db.execute(count_stmt)).scalar_one() or 0

                yield {
                    "type": "table_meta",
                    "field_id": f.id,
                    "total_count": total_count
                }

                if total_count == 0:
                    continue

                chunk_size = 1000
                for offset in range(0, total_count, chunk_size):
                    rows_stmt = (
                        select(KpiMultiLineRow.id, KpiMultiLineRow.row_index)
                        .where(
                            KpiMultiLineRow.entry_id == entry.id,
                            KpiMultiLineRow.field_id == kfield.id,
                        )
                        .order_by(KpiMultiLineRow.row_index)
                        .limit(chunk_size)
                        .offset(offset)
                    )
                    rows_res = await db.execute(rows_stmt)
                    rows_list = rows_res.all()
                    if not rows_list:
                        break

                    row_ids = [r[0] for r in rows_list]

                    cells_res = await db.execute(
                        select(
                            KpiMultiLineCell.row_id,
                            KpiMultiLineCell.value_text,
                            KpiMultiLineCell.value_number,
                            KpiMultiLineCell.value_boolean,
                            KpiMultiLineCell.value_date,
                            KpiMultiLineCell.value_json,
                            KPIFieldSubField.key
                        )
                        .join(KPIFieldSubField, KPIFieldSubField.id == KpiMultiLineCell.sub_field_id)
                        .where(KpiMultiLineCell.row_id.in_(row_ids))
                    )
                    cells_list = cells_res.all()

                    cells_by_row = defaultdict(dict)
                    for row_id, vt, vn, vb, vd, vj, sf_key in cells_list:
                        raw_val = None
                        if vj is not None:
                            raw_val = vj
                        elif vt is not None:
                            raw_val = vt
                        elif vn is not None:
                            raw_val = vn
                        elif vb is not None:
                            raw_val = vb
                        elif vd is not None:
                            try:
                                raw_val = vd.isoformat()
                            except Exception:
                                raw_val = str(vd)
                        cells_by_row[row_id][str(sf_key)] = raw_val

                    chunk_rows = []
                    for rid, r_idx in rows_list:
                        row_dict = cells_by_row.get(rid, {})
                        chunk_rows.append(row_dict)

                cfg = getattr(f, "config", None) or {}
                raw_filters = cfg.get("filters") or {}
                filtered_chunk_rows = []

                if chunk_rows:
                    if raw_filters and raw_filters.get("conditions"):
                        from app.entries.multi_item_filters import row_passes_filters
                        from app.entries.reference_filter_resolve import build_reference_resolution_map
                        conds = raw_filters.get("conditions")
                        resolution_maps = await build_reference_resolution_map(
                            db, org_id, yr, kfield, conds, chunk_rows
                        )
                        reference_field_types = {sf.key: sf.field_type.value if hasattr(sf.field_type, "value") else sf.field_type for sf in kfield.sub_fields}
                        for r in chunk_rows:
                            if row_passes_filters(r, raw_filters, resolution_maps=resolution_maps, reference_field_types=reference_field_types):
                                filtered_chunk_rows.append(r)
                    else:
                        filtered_chunk_rows = chunk_rows

                # Sort rows in memory if specified
                sort_col = cfg.get("sort_column")
                sort_dir = cfg.get("sort_direction") or "asc"
                if sort_col and filtered_chunk_rows:
                    reverse = sort_dir == "desc"
                    def sort_key(row: dict):
                        v = row.get(sort_col)
                        try:
                            return float(v)
                        except (TypeError, ValueError):
                            return str(v) if v is not None else ""
                    try:
                        filtered_chunk_rows = sorted(filtered_chunk_rows, key=sort_key, reverse=reverse)
                    except Exception:
                        pass


                yield {
                    "type": "table_rows",
                    "field_id": f.id,
                    "value_items": filtered_chunk_rows,
                    "offset": offset,
                    "done": (offset + chunk_size >= total_count)
                }

    if getattr(custom_report, "attachments", None):
        for att in custom_report.attachments:
            if getattr(att, "kpi_field", None):
                sub_fields_orm = getattr(att.kpi_field, "sub_fields", [])
                yield {
                    "type": "attachment_meta",
                    "attachment_id": att.id,
                    "title": att.title,
                    "kpi_name": getattr(att.kpi, "name", ""),
                    "field_name": getattr(att.kpi_field, "name", ""),
                    "sub_fields": [{"key": sf.key, "name": sf.name or sf.key} for sf in sub_fields_orm],
                    "total_count": 0,
                }

    yield {"type": "done"}


async def export_custom_report_attachments(
    db, custom_report_id: int, org_id: int, year: int | str | None, format: str, attachment_ids: list[int], current_user: User | None = None
) -> tuple[bytes, str, str]:
    import io
    import zipfile
    import datetime
    import re
    from sqlalchemy import select
    from app.core.models import CustomReport, KPI, KPIEntry, KpiMultiLineRow, KpiMultiLineCell, KPIFieldSubField, User
    from sqlalchemy.orm import selectinload, noload

    report = await get_custom_report(db, custom_report_id, org_id)
    if not report:
        raise ValueError("Report not found")
        
    attachments = [a for a in (report.attachments or []) if a.id in attachment_ids]
    if not attachments:
        raise ValueError("No matching attachments found")

    # Generate file for each attachment
    files = [] # list of (filename, bytes)
    yr = _parse_year_int(year)

    for att in attachments:
        # Load entries for this KPI
        kfield = att.kpi_field
        kpi = att.kpi
        
        entries_res = await db.execute(
            select(KPIEntry)
            .where(
                KPIEntry.organization_id == org_id,
                KPIEntry.kpi_id == kpi.id,
                KPIEntry.year == yr,
                KPIEntry.is_draft == False
            )
        )
        entries = entries_res.scalars().all()
        if not entries:
            continue
            
        entry = entries[-1] # Simplification, typically we sort by period_key

        rows_stmt = (
            select(KpiMultiLineRow.id, KpiMultiLineRow.row_index)
            .where(
                KpiMultiLineRow.entry_id == entry.id,
                KpiMultiLineRow.field_id == kfield.id,
            )
            .order_by(KpiMultiLineRow.row_index)
        )
        u_key = None
        if current_user is not None:
            if "unique_user_key" in current_user.__dict__:
                u_key = current_user.unique_user_key
            else:
                from app.core.models import User
                u_key = (await db.execute(select(User.unique_user_key).where(User.id == current_user.id))).scalar_one_or_none()

        if current_user is not None and u_key is not None:
            from app.core.models import ReportUserFilterConfiguration
            from sqlalchemy import and_, or_
            filter_config_res = await db.execute(
                select(ReportUserFilterConfiguration)
                .where(
                    ReportUserFilterConfiguration.report_id == custom_report_id,
                    ReportUserFilterConfiguration.enabled == True,
                    ReportUserFilterConfiguration.mli_id == kfield.id
                )
            )
            filter_config = filter_config_res.scalar_one_or_none()
            if filter_config and filter_config.field_id is not None:
                rows_stmt = rows_stmt.join(
                    KpiMultiLineCell,
                    and_(
                        KpiMultiLineCell.row_id == KpiMultiLineRow.id,
                        KpiMultiLineCell.sub_field_id == filter_config.field_id
                    )
                )
                val_str = u_key
                conditions = [KpiMultiLineCell.value_text == val_str]
                try:
                    val_float = float(val_str)
                    conditions.append(KpiMultiLineCell.value_number == val_float)
                except ValueError:
                    pass
                val_bool = val_str.lower() in ("true", "1", "yes", "y")
                if val_bool or val_str.lower() in ("false", "0", "no", "n"):
                    conditions.append(KpiMultiLineCell.value_boolean == val_bool)
                rows_stmt = rows_stmt.where(or_(*conditions))

        rows_list = (await db.execute(rows_stmt)).all()
        
        chunk_rows = []
        if rows_list:
            row_ids = [r[0] for r in rows_list]
            cells_res = await db.execute(
                select(
                    KpiMultiLineCell.row_id, KpiMultiLineCell.value_text, KpiMultiLineCell.value_number,
                    KpiMultiLineCell.value_boolean, KpiMultiLineCell.value_date, KpiMultiLineCell.value_json,
                    KPIFieldSubField.key, KPIFieldSubField.name
                )
                .join(KPIFieldSubField, KPIFieldSubField.id == KpiMultiLineCell.sub_field_id)
                .where(KpiMultiLineCell.row_id.in_(row_ids))
            )
            cells_list = cells_res.all()
            
            cells_by_row = {}
            for row_id, vt, vn, vb, vd, vj, sf_key, sf_name in cells_list:
                if row_id not in cells_by_row:
                    cells_by_row[row_id] = {}
                raw_val = vj if vj is not None else vt if vt is not None else vn if vn is not None else vb if vb is not None else (vd.isoformat() if vd else None)
                cells_by_row[row_id][sf_key] = raw_val

            for rid, r_idx in rows_list:
                chunk_rows.append(cells_by_row.get(rid, {}))

        # Build Document
        # Column selection filtering for attachment
        if att.selected_columns:
            sf_map = {sf.key: sf for sf in getattr(kfield, "sub_fields", [])}
            sub_fields = [{"key": k, "name": (sf_map[k].name if sf_map[k].name else k)} for k in att.selected_columns if k in sf_map]
        else:
            sub_fields = [{"key": sf.key, "name": sf.name or sf.key} for sf in getattr(kfield, "sub_fields", [])]

        # Row filtering for attachment
        raw_filters = att.filters or {}
        filtered_rows = []
        if chunk_rows:
            if raw_filters and raw_filters.get("conditions"):
                from app.entries.multi_item_filters import row_passes_filters
                from app.entries.reference_filter_resolve import build_reference_resolution_map
                conds = raw_filters.get("conditions")
                resolution_maps = await build_reference_resolution_map(
                    db, org_id, yr, kfield, conds, chunk_rows
                )
                reference_field_types = {sf.key: sf.field_type.value if hasattr(sf.field_type, "value") else sf.field_type for sf in kfield.sub_fields}
                for r in chunk_rows:
                    if row_passes_filters(r, raw_filters, resolution_maps=resolution_maps, reference_field_types=reference_field_types):
                        filtered_rows.append(r)
            else:
                filtered_rows = chunk_rows

        # Row sorting for attachment
        sort_col = raw_filters.get("sort_column")
        sort_dir = raw_filters.get("sort_direction") or "asc"
        if sort_col and filtered_rows:
            reverse = sort_dir == "desc"
            def sort_key(row: dict):
                v = row.get(sort_col)
                try:
                    return float(v)
                except (TypeError, ValueError):
                    return str(v) if v is not None else ""
            try:
                filtered_rows = sorted(filtered_rows, key=sort_key, reverse=reverse)
            except Exception:
                pass
        chunk_rows = filtered_rows

        clean_title = re.sub(r'[^\w\s-]', '', att.title).strip().replace(' ', '_')
        
        if format == "xlsx":
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = clean_title[:30]
            
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
            
            # Headers
            sr_c = ws.cell(row=1, column=1, value="Sr. No.")
            sr_c.font = header_font
            sr_c.fill = header_fill
            for col_idx, sf in enumerate(sub_fields):
                c = ws.cell(row=1, column=col_idx+2, value=sf["name"])
                c.font = header_font
                c.fill = header_fill
            
            for r_idx, item in enumerate(chunk_rows):
                ws.cell(row=r_idx+2, column=1, value=r_idx + 1)
                for col_idx, sf in enumerate(sub_fields):
                    ws.cell(row=r_idx+2, column=col_idx+2, value=clean_excel_value(str(item.get(sf["key"], ""))))
            
            out_io = io.BytesIO()
            wb.save(out_io)
            files.append((f"{clean_title}.xlsx", out_io.getvalue()))
            
        elif format == "pdf":
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            
            out_io = io.BytesIO()
            num_cols = (len(sub_fields) + 1) if sub_fields else 1
            
            page_width = max(612, num_cols * 65)
            pagesize = (page_width, 792)
            col_width = max(35, (page_width - 54 - 35) / max(1, len(sub_fields)))
            col_widths = [35] + [col_width] * len(sub_fields)

            pdf_doc = SimpleDocTemplate(out_io, pagesize=pagesize, leftMargin=27, rightMargin=27, topMargin=27, bottomMargin=27)
            story = []
            styles = getSampleStyleSheet()
            
            cell_header_style = ParagraphStyle(
                "CellHeaderStyle",
                parent=styles["Normal"],
                fontName="Helvetica-Bold",
                fontSize=7,
                leading=8,
                textColor=colors.whitesmoke,
            )
            cell_body_style = ParagraphStyle(
                "CellBodyStyle",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=7,
                leading=8,
                textColor=colors.HexColor("#1F2937"),
            )

            story.append(Paragraph(att.title, styles["Heading1"]))
            story.append(Spacer(1, 12))
            
            if sub_fields and chunk_rows:
                pdf_rows = chunk_rows
                if len(chunk_rows) > 3000:
                    pdf_rows = chunk_rows[:3000]
                    story.append(Paragraph(f"<i>Note: PDF attachment limited to first 3,000 rows of {len(chunk_rows)} total records. For full raw data, export as Excel (.xlsx).</i>", styles["Italic"]))
                    story.append(Spacer(1, 8))

                table_data = [[Paragraph("Sr. No.", cell_header_style)] + [Paragraph(sf["name"], cell_header_style) for sf in sub_fields]]
                for r_idx, item in enumerate(pdf_rows):
                    row = [Paragraph(str(r_idx + 1), cell_body_style)] + [Paragraph(str(item.get(sf["key"], "") or ""), cell_body_style) for sf in sub_fields]
                    table_data.append(row)
                
                t = Table(table_data, colWidths=col_widths)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]))
                story.append(t)
            else:
                story.append(Paragraph("No data available.", styles["Normal"]))
                
            pdf_doc.build(story)
            files.append((f"{clean_title}.pdf", out_io.getvalue()))
            
        elif format == "docx":
            import docx
            doc = docx.Document()
            doc.add_heading(att.title, level=1)
            
            if sub_fields and chunk_rows:
                table = doc.add_table(rows=1, cols=len(sub_fields) + 1)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Sr. No."
                for col_idx, sf in enumerate(sub_fields):
                    hdr_cells[col_idx + 1].text = sf["name"]
                
                for r_idx, item in enumerate(chunk_rows):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(r_idx + 1)
                    for col_idx, sf in enumerate(sub_fields):
                        row_cells[col_idx + 1].text = str(item.get(sf["key"], ""))
            else:
                doc.add_paragraph("No data available.")
                
            out_io = io.BytesIO()
            doc.save(out_io)
            files.append((f"{clean_title}.docx", out_io.getvalue()))

    if len(files) == 1:
        fname, fbytes = files[0]
        content_type = "application/pdf" if format == "pdf" else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" if format == "xlsx" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return fbytes, fname, content_type
    elif len(files) > 1:
        zip_io = io.BytesIO()
        with zipfile.ZipFile(zip_io, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, fbytes in files:
                zf.writestr(fname, fbytes)
        return zip_io.getvalue(), "Attachments.zip", "application/zip"
    else:
        raise ValueError("No attachments could be generated")
