"""Report template CRUD, KPI/field selection, access control, and report generation."""

import datetime
import re
from contextvars import ContextVar
import os
import time
from collections import defaultdict
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from sqlalchemy.orm import selectinload

from app.core.models import (
    ReportTemplate,
    ReportTemplateKPI,
    ReportTemplateField,
    ReportAccessPermission,
    KPI,
    KPIField,
    KPIEntry,
    KPIFieldValue,
    KpiMultiLineRow,
    KpiMultiLineCell,
    Domain,
    ReportTemplateTextBlock,
    Category,
    KPICategory,
    Organization,
    TimeDimension,
    User,
    effective_kpi_time_dimension,
    period_key_sort_order,
)
from app.reports.schemas import ReportTemplateCreate, ReportTemplateUpdate, ReportTemplateKPIAdd
from app.fields.schemas import KPIFieldResponse
from app.formula_engine.evaluator import evaluate_formula
from app.core.models import FieldType
from app.entries.service import _load_other_kpi_values, _is_subfield_satisfied_for_row
from app.entries.reference_filter_resolve import (
    build_reference_resolution_map,
    _extract_ref_label,
    _multi_raw_pieces,
    _normalize_reference_value,
)
from app.entries.multi_item_filters import row_passes_filters
from jinja2 import Environment, BaseLoader, select_autoescape
from markupsafe import escape as html_escape


# Cache within a single request/preview render to avoid repeating the same expensive
# reference-resolution work across blocks/entries. Key format is opaque to callers.
_report_preview_cache: ContextVar[dict[tuple, object] | None] = ContextVar(
    "_report_preview_cache", default=None
)


def _get_report_preview_cache() -> dict[tuple, object]:
    cache = _report_preview_cache.get()
    if cache is None:
        cache = {}
        _report_preview_cache.set(cache)
    return cache


def _report_profile_enabled() -> bool:
    return os.environ.get("REPORT_PREVIEW_PROFILE", "").strip().lower() in ("1", "true", "yes")


def _prof(msg: str) -> None:
    if _report_profile_enabled():
        print(f"[report-profile] {msg}", flush=True)


def _report_data_cache_ttl_s() -> float:
    """Optional short TTL cache for preview/generate hot reload loops."""
    raw = os.environ.get("REPORT_DATA_CACHE_SECONDS", "").strip()
    if not raw:
        return 0.0
    try:
        v = float(raw)
        return v if v > 0 else 0.0
    except (TypeError, ValueError):
        return 0.0


_report_data_cache: dict[tuple, tuple[float, dict]] = {}


def _cache_get(key: tuple) -> dict | None:
    ttl = _report_data_cache_ttl_s()
    if ttl <= 0:
        return None
    hit = _report_data_cache.get(key)
    if not hit:
        return None
    ts, val = hit
    if (time.time() - ts) > ttl:
        _report_data_cache.pop(key, None)
        return None
    return val


def _cache_set(key: tuple, val: dict) -> None:
    ttl = _report_data_cache_ttl_s()
    if ttl <= 0:
        return
    # Best-effort bounded cache (avoid unbounded growth in dev).
    if len(_report_data_cache) > 64:
        _report_data_cache.clear()
    _report_data_cache[key] = (time.time(), val)


def _ml_cell_raw(c: KpiMultiLineCell):
    if getattr(c, "value_json", None) is not None:
        return c.value_json
    if getattr(c, "value_text", None) is not None:
        return c.value_text
    if getattr(c, "value_number", None) is not None:
        return c.value_number
    if getattr(c, "value_boolean", None) is not None:
        return c.value_boolean
    if getattr(c, "value_date", None) is not None:
        try:
            return c.value_date.isoformat()
        except Exception:
            return str(c.value_date)
    return None


def _kpi_multi_line_orm_row_to_dict(r: KpiMultiLineRow) -> dict:
    d: dict = {}
    for c in getattr(r, "cells", None) or []:
        sf = getattr(c, "sub_field", None)
        key = getattr(sf, "key", None) if sf is not None else None
        if not key:
            continue
        d[str(key)] = _ml_cell_raw(c)
    return d


async def _load_multi_line_items_rows_batch(
    db: AsyncSession, *, entry_ids: list[int], field: KPIField
) -> dict[int, list[dict]]:
    """Load multi-line rows for many entries in one query (report preview was N entries × M fields round-trips)."""
    if not entry_ids:
        return {}
    res = await db.execute(
        select(KpiMultiLineRow)
        .where(
            KpiMultiLineRow.entry_id.in_(entry_ids),
            KpiMultiLineRow.field_id == field.id,
        )
        .order_by(KpiMultiLineRow.entry_id, KpiMultiLineRow.row_index)
        .options(selectinload(KpiMultiLineRow.cells).selectinload(KpiMultiLineCell.sub_field))
    )
    
    sub_fields = getattr(field, "sub_fields", None) or []
    subfields_dict = {}
    for sf in sub_fields:
        subfields_dict[sf.key] = sf
        if getattr(sf, "id", None) is not None:
            subfields_dict[int(sf.id)] = sf

    from app.fields.conditional import is_subfield_visible
    by_entry: dict[int, list[dict]] = defaultdict(list)
    for r in res.scalars().all():
        eid = getattr(r, "entry_id", None)
        if eid is None:
            continue
        row_dict = _kpi_multi_line_orm_row_to_dict(r)
        
        cleaned_r = {}
        for k, v in row_dict.items():
            sf = subfields_dict.get(k)
            if sf and is_subfield_visible(sf, subfields_dict, row_dict):
                cleaned_r[k] = v
            elif not sf:
                cleaned_r[k] = v
                
        by_entry[int(eid)].append(cleaned_r)
    return dict(by_entry)


async def _load_multi_line_items_rows(db: AsyncSession, *, entry_id: int, field: KPIField) -> list[dict]:
    m = await _load_multi_line_items_rows_batch(db, entry_ids=[entry_id], field=field)
    return m.get(entry_id, [])


async def create_report_template(
    db: AsyncSession, org_id: int, data: ReportTemplateCreate
) -> ReportTemplate:
    """Create report template."""
    rt = ReportTemplate(
        organization_id=org_id,
        name=data.name,
        description=data.description,
        body_template=data.body_template,
    )
    db.add(rt)
    await db.flush()
    return rt


async def get_report_template(
    db: AsyncSession, template_id: int, org_id: int
) -> ReportTemplate | None:
    """Get report template by id within org."""
    result = await db.execute(
        select(ReportTemplate).where(
            ReportTemplate.id == template_id,
            ReportTemplate.organization_id == org_id,
        )
    )
    return result.scalar_one_or_none()


async def get_report_template_detail(
    db: AsyncSession, template_id: int, org_id: int
) -> dict | None:
    """Get template with body_template, kpis_from_domains, and fields_by_kpi_id (all KPIs in org; read-only)."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None

    # Used by the designer as the default year for previewing KPI values.
    # Without this, Super Admin often lands on the current year, which may have no submitted entries yet.
    latest_year_row = await db.execute(
        select(func.max(KPIEntry.year)).where(KPIEntry.organization_id == org_id)
    )
    default_year = latest_year_row.scalar_one_or_none()
    try:
        default_year_int = int(default_year) if default_year is not None else None
    except (TypeError, ValueError):
        default_year_int = None

    # All KPIs in the same organization (no domain attachment required)
    kpis_from_domains = []
    org_row = await db.execute(select(Organization).where(Organization.id == org_id))
    org = org_row.scalar_one_or_none()
    org_td = TimeDimension(getattr(org, "time_dimension", None) or "yearly") if org else TimeDimension.YEARLY
    result = await db.execute(
        select(KPI)
        .where(KPI.organization_id == org_id)
        .order_by(KPI.sort_order, KPI.name)
        .options(
            selectinload(KPI.fields).selectinload(KPIField.options),
            selectinload(KPI.fields).selectinload(KPIField.sub_fields),
        )
    )
    fields_by_kpi_id: dict[str, list] = {}
    for kpi in result.unique().scalars().all():
        field_count = len(kpi.fields) if kpi.fields else 0
        kpi_year = getattr(kpi, "year", None)
        kpi_td_raw = getattr(kpi, "time_dimension", None)
        try:
            kpi_td = TimeDimension(kpi_td_raw) if kpi_td_raw else None
        except (ValueError, TypeError):
            kpi_td = None
        effective_td = effective_kpi_time_dimension(kpi_td, org_td)
        kpis_from_domains.append({
            "kpi_id": kpi.id,
            "kpi_name": kpi.name,
            "kpi_year": kpi_year,
            "fields_count": field_count,
            "time_dimension": effective_td.value,
        })
        raw_fields = list(kpi.fields) if kpi.fields else []
        raw_fields.sort(key=lambda f: (f.sort_order, f.id))
        fields_by_kpi_id[str(kpi.id)] = [
            KPIFieldResponse.model_validate(f).model_dump(mode="json") for f in raw_fields
        ]

    return {
        "id": rt.id,
        "organization_id": rt.organization_id,
        "name": rt.name,
        "description": rt.description,
        "default_year": default_year_int,
        "template_mode": getattr(rt, "template_mode", "designer"),
        "body_template": rt.body_template,
        "body_blocks": getattr(rt, "body_blocks", None),
        "attached_domains": [],
        "kpis_from_domains": kpis_from_domains,
        "fields_by_kpi_id": fields_by_kpi_id,
    }


async def list_report_templates(db: AsyncSession, org_id: int) -> list[ReportTemplate]:
    """List report templates in organization (templates are general; year is passed at generate time)."""
    q = select(ReportTemplate).where(ReportTemplate.organization_id == org_id).order_by(ReportTemplate.name)
    result = await db.execute(q)
    return list(result.scalars().all())


async def list_all_report_templates(db: AsyncSession) -> list[ReportTemplate]:
    """List all report templates across organizations (for Super Admin when no org is specified)."""
    q = select(ReportTemplate).order_by(ReportTemplate.name)
    result = await db.execute(q)
    return list(result.scalars().all())


async def add_text_block(
    db: AsyncSession, template_id: int, org_id: int, title: str | None, content: str, sort_order: int = 0
) -> ReportTemplateTextBlock | None:
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None
    tb = ReportTemplateTextBlock(report_template_id=template_id, title=title, content=content or "", sort_order=sort_order)
    db.add(tb)
    await db.flush()
    return tb


async def delete_text_block(
    db: AsyncSession, template_id: int, org_id: int, text_block_id: int
) -> bool:
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return False
    tb = (
        await db.execute(
            select(ReportTemplateTextBlock).where(
                ReportTemplateTextBlock.id == text_block_id,
                ReportTemplateTextBlock.report_template_id == template_id,
            )
        )
    ).scalar_one_or_none()
    if not tb:
        return False
    await db.delete(tb)
    await db.flush()
    return True


async def update_report_template(
    db: AsyncSession, template_id: int, org_id: int, data: ReportTemplateUpdate
) -> ReportTemplate | None:
    """Update report template."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None
    if data.name is not None:
        rt.name = data.name
    if data.description is not None:
        rt.description = data.description
    if getattr(data, "template_mode", None) is not None:
        rt.template_mode = data.template_mode
    if data.body_template is not None:
        rt.body_template = data.body_template
    if data.body_blocks is not None:
        rt.body_blocks = data.body_blocks
    await db.flush()
    return rt


async def delete_report_template(
    db: AsyncSession, template_id: int, org_id: int
) -> bool:
    """Delete report template and its related data (Super Admin only). Cascades to KPIs, text blocks, access permissions."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return False
    await db.delete(rt)
    await db.flush()
    return True


async def add_kpi_to_template(
    db: AsyncSession, template_id: int, org_id: int, data: ReportTemplateKPIAdd
) -> ReportTemplateKPI | None:
    """Add KPI to report template with optional field selection."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None
    # Verify KPI belongs to org (KPI has organization_id; domain is optional)
    result = await db.execute(
        select(KPI).where(KPI.id == data.kpi_id, KPI.organization_id == org_id)
    )
    if result.scalar_one_or_none() is None:
        return None
    rtk = ReportTemplateKPI(
        report_template_id=template_id,
        kpi_id=data.kpi_id,
        include_all_fields=data.include_all_fields,
        sort_order=data.sort_order,
    )
    db.add(rtk)
    await db.flush()
    if not data.include_all_fields and data.field_ids:
        for i, fid in enumerate(data.field_ids):
            db.add(
                ReportTemplateField(
                    report_template_kpi_id=rtk.id,
                    kpi_field_id=fid,
                    sort_order=i,
                )
            )
    await db.flush()
    return rtk


async def remove_kpi_from_template(
    db: AsyncSession, template_id: int, org_id: int, report_template_kpi_id: int
) -> bool:
    """Remove KPI from report template."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return False
    result = await db.execute(
        select(ReportTemplateKPI).where(
            ReportTemplateKPI.id == report_template_kpi_id,
            ReportTemplateKPI.report_template_id == template_id,
        )
    )
    rtk = result.scalar_one_or_none()
    if not rtk:
        return False
    await db.delete(rtk)
    await db.flush()
    return True


async def assign_report_to_user(
    db: AsyncSession,
    template_id: int,
    org_id: int,
    user_id: int,
    can_view: bool = True,
    can_print: bool = True,
    can_export: bool = True,
) -> ReportAccessPermission | None:
    """Assign report template to user (upsert). Template and user must be in same org."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None
    user_row = (await db.execute(select(User).where(User.id == user_id))).scalar_one_or_none()
    if not user_row or user_row.organization_id != org_id:
        return None
    existing = (
        await db.execute(
            select(ReportAccessPermission).where(
                ReportAccessPermission.report_template_id == template_id,
                ReportAccessPermission.user_id == user_id,
            )
        )
    ).scalar_one_or_none()
    if existing:
        existing.can_view = can_view
        existing.can_print = can_print
        existing.can_export = can_export
        await db.flush()
        return existing
    perm = ReportAccessPermission(
        report_template_id=template_id,
        user_id=user_id,
        can_view=can_view,
        can_print=can_print,
        can_export=can_export,
    )
    db.add(perm)
    await db.flush()
    return perm


async def unassign_report_from_user(
    db: AsyncSession, template_id: int, org_id: int, user_id: int
) -> bool:
    """Remove report template assignment from user."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return False
    perm = (
        await db.execute(
            select(ReportAccessPermission).where(
                ReportAccessPermission.report_template_id == template_id,
                ReportAccessPermission.user_id == user_id,
            )
        )
    ).scalar_one_or_none()
    if not perm:
        return False
    await db.delete(perm)
    await db.flush()
    return True


async def list_template_assignments(
    db: AsyncSession, template_id: int, org_id: int
) -> list[dict]:
    """List users assigned to a report template (with user info). Returns list of dicts with user_id, email, full_name, can_view, can_print, can_export."""
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return []
    result = await db.execute(
        select(ReportAccessPermission, User)
        .join(User, ReportAccessPermission.user_id == User.id)
        .where(
            ReportAccessPermission.report_template_id == template_id,
        )
    )
    rows = result.all()
    return [
        {
            "user_id": perm.user_id,
            "email": user.email,
            "full_name": user.full_name,
            "can_view": perm.can_view,
            "can_print": perm.can_print,
            "can_export": perm.can_export,
        }
        for perm, user in rows
    ]


async def user_can_access_report(
    db: AsyncSession, user_id: int, template_id: int, action: str = "view"
) -> bool:
    """Check if user can view/print/export report.

    Rules:
    - SUPER_ADMIN: can access any template.
    - ORG_ADMIN: can access any template in their organization.
    - Other roles: must be explicitly assigned (ReportAccessPermission).
    """
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()
    if not user:
        return False
    if user.role.value == "SUPER_ADMIN":
        result = await db.execute(select(ReportTemplate).where(ReportTemplate.id == template_id))
        if result.scalar_one_or_none():
            return True
    if user.role.value == "ORG_ADMIN" and user.organization_id:
        result = await db.execute(
            select(ReportTemplate).where(
                ReportTemplate.id == template_id,
                ReportTemplate.organization_id == user.organization_id,
            )
        )
        if result.scalar_one_or_none():
            return True
    result = await db.execute(
        select(ReportAccessPermission).where(
            ReportAccessPermission.report_template_id == template_id,
            ReportAccessPermission.user_id == user_id,
        )
    )
    perm = result.scalar_one_or_none()
    if not perm:
        return False
    if action == "view":
        return perm.can_view
    if action == "print":
        return perm.can_print
    if action == "export":
        return perm.can_export
    return False


def _get_kpi_field_value(kpis: list, kpi_id: int, field_key: str, sub_field_key: str | None = None, entry_index: int = 0):
    """
    Jinja-accessible helper: get value for a KPI field (optionally a sub-field of multi_line_items).
    Returns the value from the first entry by default (entry_index=0); used for placeholder rendering.
    """
    if not kpis:
        return ""
    kpi = next((k for k in kpis if k.get("kpi_id") == kpi_id), None)
    if not kpi:
        return ""
    entries = kpi.get("entries") or []
    if entry_index >= len(entries):
        return ""
    entry = entries[entry_index]
    fields = entry.get("fields") or []
    field = next((f for f in fields if f.get("field_key") == field_key), None)
    if not field:
        return ""
    val = field.get("value")
    if sub_field_key and isinstance(val, list):
        # multi_line_items: val is list of dicts; extract sub_field_key from each item
        parts = []
        for item in val:
            if isinstance(item, dict) and sub_field_key in item:
                parts.append(item[sub_field_key])
        return ", ".join(str(p) for p in parts) if parts else ""
    if val is None:
        return ""
    return val


_jinja_env = Environment(
    loader=BaseLoader(),
    autoescape=True,
)
_jinja_env.globals["get_kpi_field_value"] = _get_kpi_field_value


def _get_multi_line_field(kpis: list, kpi_id: int, field_key: str, entry_index: int = 0) -> dict | None:
    """
    Jinja-accessible helper: get the multi_line_items field dict for a given KPI and field key.
    Returns a dict with value_items (list of row dicts) and sub_field_keys (list of column keys),
    or None if not found / not multi_line. Use once then loop: {% set ml = get_multi_line_field(...) %}.
    """
    if not kpis:
        return None
    kpi = next((k for k in kpis if k.get("kpi_id") == kpi_id), None)
    if not kpi:
        return None
    entries = kpi.get("entries") or []
    if entry_index >= len(entries):
        return None
    entry = entries[entry_index]
    fields = entry.get("fields") or []
    field = next((f for f in fields if f.get("field_key") == field_key), None)
    if not field or field.get("field_type") != "multi_line_items":
        return None
    value_items = field.get("value_items")
    if not isinstance(value_items, list):
        return None
    sub_field_keys = field.get("sub_field_keys") or []
    sub_fields = field.get("sub_fields")
    if not sub_fields and sub_field_keys:
        sub_fields = [{"key": k, "name": k} for k in sub_field_keys]
    return {"value_items": value_items, "sub_field_keys": sub_field_keys, "sub_fields": sub_fields or [], "field_name": field.get("field_name", field_key)}


_jinja_env.globals["get_multi_line_field"] = _get_multi_line_field


def _apply_formula(value, formula: str):
    """
    Jinja-accessible helper: apply a formula expression to a single value.
    Formula can use variable 'value' (e.g. "value * 1.1", "round(value, 2)").
    Returns the computed result or the original value if not numeric / formula invalid.
    """
    if not formula or not str(formula).strip():
        return value
    try:
        num = float(value) if value is not None else 0
    except (TypeError, ValueError):
        return value
    result = evaluate_formula(str(formula).strip(), {"value": num}, None, None)
    return result if result is not None else value


_jinja_env.globals["apply_formula"] = _apply_formula


def _build_formula_context_from_report(kpis: list, kpi_id: int, entry_index: int):
    """
    Build (value_by_key, multi_line_items_data, other_kpi_values) from report kpis payload
    for the given kpi_id and entry_index. Used by evaluate_report_formula.
    """
    value_by_key: dict[str, float] = {}
    multi_line_items_data: dict[str, list] = {}
    other_kpi_values: dict[tuple[int, str], float] = {}

    if not kpis:
        return value_by_key, multi_line_items_data, other_kpi_values

    # Current KPI entry
    kpi_payload = next((k for k in kpis if k.get("kpi_id") == kpi_id), None)
    if kpi_payload:
        entries = kpi_payload.get("entries") or []
        entry = entries[entry_index] if entry_index < len(entries) else (entries[0] if entries else None)
        if entry:
            for f in entry.get("fields") or []:
                fkey = f.get("field_key")
                if not fkey:
                    continue
                ft = f.get("field_type") or ""
                val = f.get("value")
                if ft in ("number", "formula"):
                    try:
                        value_by_key[fkey] = float(val) if val is not None else 0.0
                    except (TypeError, ValueError):
                        value_by_key[fkey] = 0.0
                elif ft == "multi_line_items":
                    items = f.get("value_items")
                    if isinstance(items, list):
                        multi_line_items_data[fkey] = items

    # Other KPIs' numeric values (same entry index)
    for k in kpis:
        other_id = k.get("kpi_id")
        if other_id is None or other_id == kpi_id:
            continue
        entries = k.get("entries") or []
        other_entry = entries[entry_index] if entry_index < len(entries) else (entries[0] if entries else None)
        if not other_entry:
            continue
        for f in other_entry.get("fields") or []:
            ft = f.get("field_type") or ""
            if ft not in ("number", "formula"):
                continue
            fkey = f.get("field_key")
            if not fkey:
                continue
            val = f.get("value")
            try:
                other_kpi_values[(other_id, fkey)] = float(val) if val is not None else 0.0
            except (TypeError, ValueError):
                other_kpi_values[(other_id, fkey)] = 0.0

    return value_by_key, multi_line_items_data, other_kpi_values


def _normalize_report_formula_expression(expression: str) -> str:
    """
    Collapse accidental duplicate paste: the same expression concatenated back-to-back
    (e.g. COUNT_ITEMS(a,b)COUNT_ITEMS(a,b)) which is invalid Python and breaks simpleeval.
    Repeatedly halve while the string is two identical halves.
    """
    s = (expression or "").strip()
    while len(s) >= 4 and len(s) % 2 == 0:
        mid = len(s) // 2
        left, right = s[:mid], s[mid:]
        if left == right:
            s = left.strip()
        else:
            break
    return s


def _evaluate_report_formula(kpis: list, expression: str, kpi_id: int, entry_index: int = 0):
    """
    Jinja-accessible helper: evaluate a full formula expression in report context.
    Uses the same expression language as KPI formula fields (field refs, SUM_ITEMS, KPI_FIELD, etc.).
    """
    if not expression or not str(expression).strip():
        return ""
    expression = _normalize_report_formula_expression(str(expression))
    value_by_key, multi_line_items_data, other_kpi_values = _build_formula_context_from_report(
        kpis, kpi_id, entry_index
    )
    result = evaluate_formula(expression, value_by_key, multi_line_items_data, other_kpi_values)
    if result is None:
        return ""
    return result


_jinja_env.globals["evaluate_report_formula"] = _evaluate_report_formula


def _filter_entries_by_period(entries: list, period_key: str | None = None, all_periods: bool = False) -> list:
    """
    Jinja filter: filter entry list by time dimension.
    - If all_periods True: return all entries.
    - If period_key set (e.g. 'Q1'): return entries with that period_key only.
    - Else (latest): return the last entry only (entries assumed sorted by period).
    """
    if not entries:
        return []
    if all_periods:
        return entries
    if period_key is not None and str(period_key).strip():
        pk = str(period_key).strip()
        return [e for e in entries if (e.get("period_key") or "") == pk]
    return [entries[-1]] if entries else []


_jinja_env.filters["filter_entries_by_period"] = _filter_entries_by_period


def _block_time_dimension_vars(b: dict) -> tuple[str, bool]:
    """
    Return (jinja_prefix, use_filter) for block time dimension.
    prefix sets _td_period and _td_all for use in filter_entries_by_period.
    use_filter True => caller should inject _td_entries and replace kpi.entries in content.
    """
    mode = (b.get("timeDimensionMode") or b.get("time_dimension_mode") or "latest").strip().lower()
    period_key = b.get("periodKey") or b.get("period_key") or ""
    if isinstance(period_key, str):
        period_key = period_key.strip()
    else:
        period_key = str(period_key).strip()
    all_periods = mode == "all_periods" or mode == "all"
    if all_periods:
        return "{% set _td_period = none %}{% set _td_all = true %}", True
    if mode == "single_period" and period_key:
        # Escape for Jinja string
        pk_esc = str(period_key).replace("\\", "\\\\").replace("'", "\\'")
        return f"{{% set _td_period = '{pk_esc}' %}}{{% set _td_all = false %}}", True
    return "{% set _td_period = none %}{% set _td_all = false %}", True


def _inject_time_dimension_filter(content: str, td_prefix: str) -> str:
    """Prepend td_prefix and replace kpi.entries with _td_entries (with filter)."""
    if not content.strip():
        return content
    inject = "{% set _td_entries = __KPI_ENTRIES__ | filter_entries_by_period(_td_period, _td_all) %}"
    content = td_prefix + content
    content = content.replace("{% for kpi in kpis %}", "{% for kpi in kpis %}" + inject, 1)
    content = content.replace("kpi.entries", "_td_entries")
    content = content.replace("__KPI_ENTRIES__", "kpi.entries")
    # Safe access when _td_entries is empty (e.g. latest with no data, or single_period with no match)
    content = content.replace("_td_entries[0].fields", "((_td_entries|first)|default({})).fields|default([])")
    content = content.replace("_td_entries[0]", "(_td_entries|first)")
    return content


def _blocks_to_jinja(blocks: list[dict]) -> str:
    """
    Convert visual builder block list to Jinja2 HTML template.
    Block types: title, section_heading, spacer, text, domain_list, domain_categories,
    domain_kpis, kpi_table, kpi_grid, kpi_list, single_value.
    """
    out: list[str] = []
    for bi, b in enumerate(blocks):
        block_type = (b.get("type") or "").strip()
        if not block_type:
            continue
        if block_type == "title":
            use_name = b.get("useTemplateName", True)
            custom = (b.get("customText") or "").strip()
            if custom:
                out.append(f'<h1 class="report-title">{custom}</h1>')
            elif use_name:
                out.append('<h1 class="report-title">{{ template_name }}</h1>')
            out.append('<p class="report-year">Year: {{ year }}</p>')
        elif block_type == "section_heading":
            text = (b.get("text") or "").strip() or "Section"
            level = min(4, max(1, int(b.get("level") or 2)))
            out.append(f"<h{level} class=\"report-section\">{text}</h{level}>")
        elif block_type == "spacer":
            size = b.get("size") or "medium"
            height = {"small": "16px", "medium": "24px", "large": "40px"}.get(size, "24px")
            out.append(f'<div class="report-spacer" style="height: {height}"></div>')
        elif block_type == "text":
            content = (b.get("content") or "").strip()
            if content:
                out.append(f'<div class="report-text-block">{content}</div>')
        elif block_type == "domain_list":
            domain_ids = b.get("domainIds") or []
            if domain_ids:
                ids_str = ", ".join(str(i) for i in domain_ids)
                out.append(
                    "{% for domain in domains %}"
                    f"{{% if domain.id in [{ids_str}] %}}"
                    '<div class="report-domain"><h3>{{ domain.name }}</h3></div>'
                    "{% endif %}{% endfor %}"
                )
            else:
                out.append(
                    '{% for domain in domains %}'
                    '<div class="report-domain"><h3>{{ domain.name }}</h3></div>'
                    '{% endfor %}'
                )
        elif block_type == "domain_categories":
            domain_ids = b.get("domainIds") or []
            if domain_ids:
                ids_str = ", ".join(str(i) for i in domain_ids)
                out.append(
                    "{% for domain in domains %}"
                    f"{{% if domain.id in [{ids_str}] %}}"
                    '<div class="report-domain"><h3>{{ domain.name }}</h3>'
                    '<ul>{% for cat in domain.categories %}<li>{{ cat.name }}</li>{% endfor %}</ul></div>'
                    "{% endif %}{% endfor %}"
                )
            else:
                out.append(
                    "{% for domain in domains %}"
                    '<div class="report-domain"><h3>{{ domain.name }}</h3>'
                    '<ul>{% for cat in domain.categories %}<li>{{ cat.name }}</li>{% endfor %}</ul></div>'
                    "{% endfor %}"
                )
        elif block_type == "domain_kpis":
            domain_ids = b.get("domainIds") or []
            if domain_ids:
                ids_str = ", ".join(str(i) for i in domain_ids)
                out.append(
                    "{% for domain in domains %}"
                    f"{{% if domain.id in [{ids_str}] %}}"
                    '<div class="report-domain"><h3>{{ domain.name }}</h3>'
                    '<ul>{% for cat in domain.categories %}<li>{{ cat.name }}'
                    '<ul>{% for kpi in cat.kpis %}<li>{{ kpi.kpi_name }}</li>{% endfor %}</ul>'
                    '</li>{% endfor %}</ul></div>'
                    "{% endif %}{% endfor %}"
                )
            else:
                out.append(
                    "{% for domain in domains %}"
                    '<div class="report-domain"><h3>{{ domain.name }}</h3>'
                    '<ul>{% for cat in domain.categories %}<li>{{ cat.name }}'
                    '<ul>{% for kpi in cat.kpis %}<li>{{ kpi.kpi_name }}</li>{% endfor %}</ul>'
                    '</li>{% endfor %}</ul></div>'
                    "{% endfor %}"
                )
        elif block_type == "single_value":
            kpi_id = int(b.get("kpiId") or 0)
            field_key = (b.get("fieldKey") or "").strip()
            sub_key = (b.get("subFieldKey") or "").strip() or None
            entry_idx = int(b.get("entryIndex") or 0)
            if not field_key:
                continue
            sub_arg = f", '{sub_key}'" if sub_key else ", none"
            out.append(
                f'<span class="report-single-value">'
                f"{{{{ get_kpi_field_value(kpis, {kpi_id}, '{field_key}'{sub_arg}, {entry_idx}) }}}}"
                f"</span>"
            )
        elif block_type == "kpi_table":
            block_uid = str(b.get("id") or f"__idx_{bi}__")
            _buid = _jinja_quote_block_uid(block_uid)
            _ml_prefix_f = (
                "{% set _eid = (entry.entry_id|string) if entry.entry_id is not none else '__none__' %}"
                "{% set _by_ent = ((multi_line_block_rows.get('" + _buid + "', {}) | default({})).get((kpi.kpi_id|string), {}) | default({})).get(_eid, {}) | default({}) %}"
                "{% set _f_rows = _by_ent.get(f.field_key, none) %}"
                "{% set _vi = _f_rows if _f_rows is not none else f.value_items %}"
            )
            _ml_prefix_ef = (
                "{% set _eid = (entry.entry_id|string) if entry.entry_id is not none else '__none__' %}"
                "{% set _by_ent = ((multi_line_block_rows.get('" + _buid + "', {}) | default({})).get((kpi.kpi_id|string), {}) | default({})).get(_eid, {}) | default({}) %}"
                "{% set _ef_rows = _by_ent.get(ef.field_key, none) %}"
                "{% set _vi = _ef_rows if _ef_rows is not none else ef.value_items %}"
            )
            kpi_ids = b.get("kpiIds") or []
            field_keys = b.get("fieldKeys") or []
            one_per_kpi = b.get("oneTablePerKpi", True)
            fields_layout = b.get("fieldsLayout") or b.get("fields_layout") or "columns"
            # Explicit False → hide KPI heading; missing or True → show (support both camelCase and snake_case)
            _sth = b.get("showTableHeading") if "showTableHeading" in b else b.get("show_table_heading")
            show_table_heading = _sth is not False
            # For KPI tables, always render multi-line items as standalone tables (not nested in a cell)
            # to keep the main KPI table readable.
            show_multi_as_table = True
            # When False, hide the parent multi-line field name; inner table stays
            _sml = b.get("showMultiLineFieldLabel") if "showMultiLineFieldLabel" in b else b.get("show_multi_line_field_label")
            show_multi_line_field_label = _sml is not False
            _field_display = b.get("fieldDisplayNames") or b.get("field_display_names") or {}
            _display_parts = [
                f"'{str(k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': '{str(v).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}'"
                for k, v in _field_display.items()
                if v is not None and str(v).strip()
            ]
            _display_prefix = "{% set field_display_names = {" + ", ".join(_display_parts) + "} %}"
            _label_f = "{{ (field_display_names.get(f.field_key) or f.field_name) | default(f.field_name) }}"
            _label_key = "{{ (field_display_names.get(key) or kpi.field_names.get(key, key)) | default(key) }}"
            _sub_field_display = b.get("subFieldDisplayNames") or b.get("sub_field_display_names") or {}
            _sub_display_outer = []
            for _fk, _inner in (_sub_field_display or {}).items():
                if not _inner or not isinstance(_inner, dict):
                    continue
                _inner_parts = [
                    f"'{str(_k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': '{str(_v).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}'"
                    for _k, _v in _inner.items()
                    if _v is not None and str(_v).strip()
                ]
                if _inner_parts:
                    _sub_display_outer.append(f"'{str(_fk).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': {{{', '.join(_inner_parts)}}}")
            _sub_display_prefix = "{% set sub_field_display_names = {" + ", ".join(_sub_display_outer) + "} %}" if _sub_display_outer else "{% set sub_field_display_names = {} %}"
            _sub_label_sf_f = "{{ ((sub_field_display_names.get(f.field_key) or {}) | default({})).get(sf.key) or sf.name | default(sf.name) }}"
            _sub_label_sf_ef = "{{ ((sub_field_display_names.get(ef.field_key) or {}) | default({})).get(sf.key) or sf.name | default(sf.name) }}"
            _ml_sub = b.get("multiLineSubFieldKeys") or b.get("multi_line_sub_field_keys") or {}
            _sub_keys_parts = [
                f"'{str(k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': [{', '.join(repr(str(s)) for s in (v or []))}]"
                for k, v in _ml_sub.items()
            ]
            _sub_keys_prefix = "{% set field_sub_field_keys = {" + ", ".join(_sub_keys_parts) + "} %}"
            _show_ml_label_prefix = "{% set show_multi_line_field_label = " + ("true" if show_multi_line_field_label else "false") + " %}"
            _column_align_raw = b.get("columnAlign") or b.get("column_align") or {}
            _align_map = {k: "left" for k in (field_keys or [])}
            for _k, _v in _column_align_raw.items():
                if _v in ("left", "center", "right", "justify"):
                    _align_map[_k] = _v
            _column_align_parts = [f"'{str(k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': '{v}'" for k, v in _align_map.items()]
            _column_align_prefix = "{% set column_align = {" + ", ".join(_column_align_parts) + "} %}" if _column_align_parts else "{% set column_align = {} %}"
            _th_style_key = ' style="text-align: {{ column_align.get(key, \'left\') }}"'
            _td_style_key = ' style="text-align: {{ column_align.get(key, \'left\') }}"'
            _td_style_f = ' style="text-align: {{ column_align.get(f.field_key, \'left\') }}"'
            _td_style_ef = ' style="text-align: {{ column_align.get(ef.field_key, \'left\') }}"'
            _label_f_cond = "{% if show_multi_line_field_label or f.field_type != 'multi_line_items' %}" + _label_f + "{% endif %}"
            _label_key_cond = "{% set _fl = (kpi.entries[0].fields | default([]) | selectattr('field_key', 'equalto', key) | list) %}{% if show_multi_line_field_label or (_fl | length == 0) or (((_fl|first)|default({})).field_type != 'multi_line_items') %}" + _label_key + "{% endif %}"
            # Scalar cell (main KPI table). Multi-line fields are excluded from the main table.
            _cell_scalar = "{{ f.value }}"
            _cell_scalar_ef = "{{ ef.value }}"

            # Standalone multi-line table snippet (rendered beneath the main KPI table).
            _multi_table_f = (
                _ml_prefix_f
                + "{% set show_sub_keys = field_sub_field_keys.get(f.field_key, []) | default([]) %}"
                "{% if f.field_type == 'multi_line_items' and _vi %}"
                "<table border=\"1\" cellpadding=\"4\" style=\"border-collapse: collapse; width: 100%;\">"
                "<tr>"
                "{% if show_sub_keys and (show_sub_keys|length) > 0 %}"
                "{% for _sk in show_sub_keys %}{% for sf in (f.sub_fields | default([])) %}{% if sf.key == _sk %}<th>"
                + _sub_label_sf_f
                + "</th>{% endif %}{% endfor %}{% endfor %}"
                "{% else %}"
                "{% for sf in (f.sub_fields | default([])) %}<th>" + _sub_label_sf_f + "</th>{% endfor %}"
                "{% endif %}"
                "</tr>"
                "{% for item in _vi %}<tr>"
                "{% if show_sub_keys and (show_sub_keys|length) > 0 %}"
                "{% for _sk in show_sub_keys %}{% for sf in (f.sub_fields | default([])) %}{% if sf.key == _sk %}<td>{{ item[sf.key] }}</td>{% endif %}{% endfor %}{% endfor %}"
                "{% else %}"
                "{% for sf in (f.sub_fields | default([])) %}<td>{{ item[sf.key] }}</td>{% endfor %}"
                "{% endif %}"
                "</tr>{% endfor %}"
                "</table>{% else %}<p style=\"margin:0; color:#666;\">No rows.</p>{% endif %}"
            )
            _multi_table_ef = (
                _ml_prefix_ef
                + "{% set show_sub_keys = field_sub_field_keys.get(ef.field_key, []) | default([]) %}"
                "{% if ef.field_type == 'multi_line_items' and _vi %}"
                "<table border=\"1\" cellpadding=\"4\" style=\"border-collapse: collapse; width: 100%;\">"
                "<tr>"
                "{% if show_sub_keys and (show_sub_keys|length) > 0 %}"
                "{% for _sk in show_sub_keys %}{% for sf in (ef.sub_fields | default([])) %}{% if sf.key == _sk %}<th>"
                + _sub_label_sf_ef
                + "</th>{% endif %}{% endfor %}{% endfor %}"
                "{% else %}"
                "{% for sf in (ef.sub_fields | default([])) %}<th>" + _sub_label_sf_ef + "</th>{% endfor %}"
                "{% endif %}"
                "</tr>"
                "{% for item in _vi %}<tr>"
                "{% if show_sub_keys and (show_sub_keys|length) > 0 %}"
                "{% for _sk in show_sub_keys %}{% for sf in (ef.sub_fields | default([])) %}{% if sf.key == _sk %}<td>{{ item[sf.key] }}</td>{% endif %}{% endfor %}{% endfor %}"
                "{% else %}"
                "{% for sf in (ef.sub_fields | default([])) %}<td>{{ item[sf.key] }}</td>{% endfor %}"
                "{% endif %}"
                "</tr>{% endfor %}"
                "</table>{% else %}<p style=\"margin:0; color:#666;\">No rows.</p>{% endif %}"
            )
            heading_html = '<h4>{{ kpi.kpi_name }}</h4>' if show_table_heading else ""
            _td_prefix, _ = _block_time_dimension_vars(b)
            if fields_layout == "rows":
                if not kpi_ids and not field_keys:
                    _multi_section = (
                        "{% for entry in kpi.entries %}"
                        "{% for f in entry.fields %}"
                        "{% if f.field_type == 'multi_line_items' %}"
                        "{% if show_multi_line_field_label %}<div style=\"margin-top: 0.75rem; font-weight: 600;\">"
                        + _label_f
                        + "</div>{% endif %}"
                        + _multi_table_f
                        + "{% endif %}"
                        + "{% endfor %}"
                        + "{% endfor %}"
                    )
                    _content = (
                        _display_prefix
                        + _sub_display_prefix
                        + _sub_keys_prefix
                        + _show_ml_label_prefix
                        + _column_align_prefix
                        + '<div class="report-kpi-table">'
                        + "{% if kpis %}"
                        + "{% for kpi in kpis %}"
                        + heading_html
                        + '<table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%;">'
                        + '<tbody>'
                        + "{% for f in kpi.entries[0].fields if kpi.entries %}"
                        + "{% if f.field_type != 'multi_line_items' %}"
                        "<tr><td" + _td_style_f + ">" + _label_f_cond + "</td>"
                        + "{% for entry in kpi.entries %}"
                        "{% for ef in entry.fields %}{% if ef.field_key == f.field_key %}<td" + _td_style_ef + ">" + _cell_scalar_ef + "</td>{% endif %}{% endfor %}"
                        + "{% endfor %}"
                        "</tr>"
                        + "{% endif %}"
                        + "{% endfor %}"
                        "</tbody></table>"
                        + _multi_section
                        + "{% endfor %}{% else %}<p>No data.</p>{% endif %}</div>"
                    )
                    out.append(_inject_time_dimension_filter(_content, _td_prefix))
                else:
                    fid_list = ", ".join(str(i) for i in kpi_ids)
                    fkeys_list = ", ".join(repr(k) for k in field_keys)
                    _cell_by_key = "{% for f in entry.fields %}{% if f.field_key == key %}<td" + _td_style_key + ">" + _cell_scalar + "</td>{% endif %}{% endfor %}"
                    _multi_section = (
                        "{% for entry in kpi.entries %}"
                        "{% for f in entry.fields %}"
                        "{% if f.field_type == 'multi_line_items' and f.field_key in field_keys_list %}"
                        "{% if show_multi_line_field_label %}<div style=\"margin-top: 0.75rem; font-weight: 600;\">"
                        + _label_f
                        + "</div>{% endif %}"
                        + _multi_table_f
                        + "{% endif %}"
                        + "{% endfor %}"
                        + "{% endfor %}"
                    )
                    _content = (
                        _display_prefix
                        + _sub_display_prefix
                        + _sub_keys_prefix
                        + _show_ml_label_prefix
                        + _column_align_prefix
                        + f"{{% set kpi_ids_set = [{fid_list}] %}}"
                        + f"{{% set field_keys_list = [{fkeys_list}] %}}"
                        '<div class="report-kpi-table">'
                        + "{% if kpis %}"
                        + "{% for kpi in kpis %}"
                        + "{% if kpi.kpi_id in kpi_ids_set %}"
                        + heading_html
                        + '<table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%;">'
                        + '<tbody>'
                        + "{% for key in field_keys_list %}"
                        + "{% set _fl = (kpi.entries[0].fields | default([]) | selectattr('field_key', 'equalto', key) | list) %}"
                        + "{% if (_fl | length == 0) or (((_fl|first)|default({})).field_type != 'multi_line_items') %}"
                        "<tr><td>" + _label_key_cond + "</td>{% for entry in kpi.entries %}" + _cell_by_key + "{% endfor %}</tr>"
                        + "{% endif %}"
                        + "{% endfor %}"
                        "</tbody></table>"
                        + _multi_section
                        + "{% endif %}"
                        + "{% endfor %}{% else %}<p>No data.</p>{% endif %}</div>"
                    )
                    out.append(_inject_time_dimension_filter(_content, _td_prefix))
            else:
                if not kpi_ids and not field_keys:
                    _multi_section = (
                        "{% for entry in kpi.entries %}"
                        "{% for f in entry.fields %}"
                        "{% if f.field_type == 'multi_line_items' %}"
                        "{% if show_multi_line_field_label %}<div style=\"margin-top: 0.75rem; font-weight: 600;\">"
                        + _label_f
                        + "</div>{% endif %}"
                        + _multi_table_f
                        + "{% endif %}"
                        + "{% endfor %}"
                        + "{% endfor %}"
                    )
                    _content = (
                        _display_prefix
                        + _sub_display_prefix
                        + _sub_keys_prefix
                        + _show_ml_label_prefix
                        + _column_align_prefix
                        + '<div class="report-kpi-table">'
                        + "{% if kpis %}"
                        + "{% for kpi in kpis %}"
                        + heading_html
                        + '<table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%;">'
                        '<thead><tr>{% for f in kpi.entries[0].fields if kpi.entries %}{% if f.field_type != "multi_line_items" %}<th' + _td_style_f + '>' + _label_f + '</th>{% endif %}{% endfor %}</tr></thead>'
                        + '<tbody>'
                        + "{% for entry in kpi.entries %}"
                        '<tr>{% for f in entry.fields %}{% if f.field_type != "multi_line_items" %}<td' + _td_style_f + '>' + _cell_scalar + '</td>{% endif %}{% endfor %}</tr>'
                        + "{% endfor %}"
                        "</tbody></table>"
                        + _multi_section
                        + "{% endfor %}{% else %}<p>No data.</p>{% endif %}</div>"
                    )
                    out.append(_inject_time_dimension_filter(_content, _td_prefix))
                else:
                    fid_list = ", ".join(str(i) for i in kpi_ids)
                    fkeys_list = ", ".join(repr(k) for k in field_keys)
                    _cell_by_key = "{% for f in entry.fields %}{% if f.field_key == key %}<td" + _td_style_key + ">" + _cell_scalar + "</td>{% endif %}{% endfor %}"
                    _multi_section = (
                        "{% for entry in kpi.entries %}"
                        "{% for f in entry.fields %}"
                        "{% if f.field_type == 'multi_line_items' and f.field_key in field_keys_list %}"
                        "{% if show_multi_line_field_label %}<div style=\"margin-top: 0.75rem; font-weight: 600;\">"
                        + _label_f
                        + "</div>{% endif %}"
                        + _multi_table_f
                        + "{% endif %}"
                        + "{% endfor %}"
                        + "{% endfor %}"
                    )
                    _content = (
                        _display_prefix
                        + _sub_display_prefix
                        + _sub_keys_prefix
                        + _show_ml_label_prefix
                        + _column_align_prefix
                        + f"{{% set kpi_ids_set = [{fid_list}] %}}"
                        + f"{{% set field_keys_list = [{fkeys_list}] %}}"
                        '<div class="report-kpi-table">'
                        + "{% if kpis %}"
                        + "{% for kpi in kpis %}"
                        + "{% if kpi.kpi_id in kpi_ids_set %}"
                        + heading_html
                        + '<table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%;">'
                        '<thead><tr>{% for key in field_keys_list %}{% set _fl = (kpi.entries[0].fields | default([]) | selectattr("field_key", "equalto", key) | list) %}{% if (_fl | length == 0) or (((_fl|first)|default({})).field_type != "multi_line_items") %}<th' + _th_style_key + '>' + _label_key + '</th>{% endif %}{% endfor %}</tr></thead>'
                        + '<tbody>'
                        + "{% for entry in kpi.entries %}"
                        '<tr>{% for key in field_keys_list %}{% set _fl = (entry.fields | default([]) | selectattr("field_key", "equalto", key) | list) %}{% if (_fl | length == 0) or (((_fl|first)|default({})).field_type != "multi_line_items") %}' + _cell_by_key + "{% endif %}{% endfor %}</tr>"
                        + "{% endfor %}"
                        "</tbody></table>"
                        + _multi_section
                        + "{% endif %}"
                        + "{% endfor %}{% else %}<p>No data.</p>{% endif %}</div>"
                    )
                    out.append(_inject_time_dimension_filter(_content, _td_prefix))
        elif block_type == "kpi_multi_table":
            kpi_id = int(b.get("kpiId") or 0)
            field_key = (b.get("fieldKey") or "").strip()
            if not kpi_id or not field_key:
                continue
            field_key_escaped = field_key.replace("\\", "\\\\").replace("'", "\\'")
            out.append(
                "<div class=\"report-kpi-multi-table\">"
                "{% set _ml = get_multi_line_field(kpis, "
                + str(kpi_id)
                + ", '"
                + field_key_escaped
                + "', 0) %}"
                "{% if _ml %}"
                "<table border=\"1\" cellpadding=\"4\" style=\"border-collapse: collapse; width: 100%;\">"
                "<tr>{% for sf in (_ml.sub_fields | default([])) %}<th>{{ sf.name }}</th>{% endfor %}</tr>"
                "{% for item in _ml.value_items %}<tr>{% for sf in (_ml.sub_fields | default([])) %}<td>{{ item[sf.key] }}</td>{% endfor %}</tr>{% endfor %}"
                "</table>"
                "{% endif %}</div>"
            )
        elif block_type == "simple_table":
            rows = b.get("rows") or []
            row_parts = []
            for row in rows:
                cells = row.get("cells") if isinstance(row, dict) else []
                cell_parts = []
                for cell in cells:
                    if not isinstance(cell, dict):
                        cell_parts.append("<td></td>")
                        continue
                    ctype = cell.get("type") or "text"
                    align = cell.get("align") or "left"
                    if align not in ("left", "center", "right", "justify"):
                        align = "left"
                    td_style = f' style="text-align: {align}"'
                    if ctype == "text":
                        content = (cell.get("content") or "").strip()
                        cell_parts.append(f"<td{td_style}>{html_escape(content)}</td>")
                    elif ctype == "kpi":
                        kpi_id = int(cell.get("kpiId") or 0)
                        field_key = (cell.get("fieldKey") or "").strip().replace("\\", "\\\\").replace("'", "\\'")
                        sub_key = (cell.get("subFieldKey") or "").strip()
                        sub_field_group_fn = (cell.get("subFieldGroupFn") or "SUM_ITEMS").strip() or "SUM_ITEMS"
                        entry_idx = int(cell.get("entryIndex") or 0)
                        if cell.get("asGroup"):
                            cell_parts.append(
                                "<td" + td_style + ">{% set _ml = get_multi_line_field(kpis, " + str(kpi_id) + ", '" + field_key + "', " + str(entry_idx) + ") %}"
                                "{% if _ml %}<table border=\"1\" cellpadding=\"4\" style=\"border-collapse: collapse;\">"
                                "<tr>{% for sf in (_ml.sub_fields | default([])) %}<th>{{ sf.name }}</th>{% endfor %}</tr>"
                                "{% for item in _ml.value_items %}<tr>{% for sf in (_ml.sub_fields | default([])) %}<td>{{ item[sf.key] }}</td>{% endfor %}</tr>{% endfor %}"
                                "</table>{% endif %}</td>"
                            )
                        elif sub_key:
                            raw_field_key = (cell.get("fieldKey") or "").strip()
                            formula_expr = f"{sub_field_group_fn}({raw_field_key}, {sub_key})"
                            formula_escaped = formula_expr.replace("\\", "\\\\").replace("'", "\\'")
                            cell_parts.append(
                                f"<td{td_style}>{{{{ evaluate_report_formula(kpis, '{formula_escaped}', {kpi_id}, {entry_idx}) }}}}</td>"
                            )
                        else:
                            sub_arg = ", none"
                            cell_parts.append(
                                f"<td{td_style}>{{{{ get_kpi_field_value(kpis, {kpi_id}, '{field_key}'{sub_arg}, {entry_idx}) }}}}</td>"
                            )
                    elif ctype == "formula":
                        kpi_id = int(cell.get("kpiId") or 0)
                        entry_idx = int(cell.get("entryIndex") or 0)
                        formula = (cell.get("formula") or "").strip().replace("\\", "\\\\").replace("'", "\\'")
                        cell_parts.append(
                            f"<td{td_style}>{{{{ evaluate_report_formula(kpis, '{formula}', {kpi_id}, {entry_idx}) }}}}</td>"
                        )
                    else:
                        cell_parts.append("<td" + td_style + "></td>")
                row_parts.append("<tr>" + "".join(cell_parts) + "</tr>")
            out.append(
                '<div class="report-simple-table"><table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%;">'
                "<tbody>" + "".join(row_parts) + "</tbody></table></div>"
            )
        elif block_type == "kpi_grid":
            block_uid = str(b.get("id") or f"__idx_{bi}__")
            _buid_g = _jinja_quote_block_uid(block_uid)
            _grid_ml_prefix_f = (
                "{% set _eid = (entry.entry_id|string) if entry.entry_id is not none else '__none__' %}"
                "{% set _by_ent = ((multi_line_block_rows.get('" + _buid_g + "', {}) | default({})).get((kpi.kpi_id|string), {}) | default({})).get(_eid, {}) | default({}) %}"
                "{% set _f_rows = _by_ent.get(f.field_key, none) %}"
                "{% set _vi = _f_rows if _f_rows is not none else f.value_items %}"
            )
            kpi_ids = b.get("kpiIds") or []
            field_keys = b.get("fieldKeys") or []
            _field_display = b.get("fieldDisplayNames") or b.get("field_display_names") or {}
            _display_parts = [
                f"'{str(k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': '{str(v).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}'"
                for k, v in _field_display.items()
                if v is not None and str(v).strip()
            ]
            _display_prefix = "{% set field_display_names = {" + ", ".join(_display_parts) + "} %}"
            _label_f = "{{ (field_display_names.get(f.field_key) or f.field_name) | default(f.field_name) }}"
            _label_key = "{{ (field_display_names.get(key) or kpi.field_names.get(key, key)) | default(key) }}"
            _sub_field_display = b.get("subFieldDisplayNames") or b.get("sub_field_display_names") or {}
            _sub_display_outer = []
            for _fk, _inner in (_sub_field_display or {}).items():
                if not _inner or not isinstance(_inner, dict):
                    continue
                _inner_parts = [
                    f"'{str(_k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': '{str(_v).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}'"
                    for _k, _v in _inner.items()
                    if _v is not None and str(_v).strip()
                ]
                if _inner_parts:
                    _sub_display_outer.append(f"'{str(_fk).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': {{{', '.join(_inner_parts)}}}")
            _sub_display_prefix = "{% set sub_field_display_names = {" + ", ".join(_sub_display_outer) + "} %}" if _sub_display_outer else "{% set sub_field_display_names = {} %}"
            _sub_label_sf_f = "{{ ((sub_field_display_names.get(f.field_key) or {}) | default({})).get(sf.key) or sf.name | default(sf.name) }}"
            _ml_sub = b.get("multiLineSubFieldKeys") or b.get("multi_line_sub_field_keys") or {}
            _sub_keys_parts = [
                f"'{str(k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': [{', '.join(repr(str(s)) for s in (v or []))}]"
                for k, v in _ml_sub.items()
            ]
            _sub_keys_prefix = "{% set field_sub_field_keys = {" + ", ".join(_sub_keys_parts) + "} %}"
            _grid_cell_multi = (
                _grid_ml_prefix_f
                + "{% set show_sub_keys = field_sub_field_keys.get(f.field_key, []) | default([]) %}"
                "{% if f.field_type == 'multi_line_items' and _vi %}"
                "<table border=\"1\" cellpadding=\"4\" style=\"border-collapse: collapse;\">"
                "<tr>{% for sf in (f.sub_fields | default([])) %}{% if not show_sub_keys or sf.key in show_sub_keys %}<th>" + _sub_label_sf_f + "</th>{% endif %}{% endfor %}</tr>"
                "{% for item in _vi %}<tr>{% for sf in (f.sub_fields | default([])) %}{% if not show_sub_keys or sf.key in show_sub_keys %}<td>{{ item[sf.key] }}</td>{% endif %}{% endfor %}</tr>{% endfor %}"
                "</table>{% else %}{{ f.value }}{% endif %}"
            )
            _td_prefix_grid, _ = _block_time_dimension_vars(b)
            if not kpi_ids and not field_keys:
                _content = (
                    _display_prefix
                    + _sub_display_prefix
                    + _sub_keys_prefix
                    + '<div class="report-kpi-grid" style="display: grid; grid-template-columns: repeat(auto-fill, minmax(260px, 1fr)); gap: 1rem;">'
                    "{% if kpis %}{% for kpi in kpis %}"
                    "{% for entry in kpi.entries %}"
                    '<div class="report-card" style="border: 1px solid #ddd; padding: 1rem; border-radius: 8px;">'
                    '<h4 style="margin-top: 0;">{{ kpi.kpi_name }}</h4>'
                    "{% for f in entry.fields %}"
                    '<p style="margin: 0.25rem 0;"><strong>' + _label_f + ':</strong> ' + _grid_cell_multi + '</p>'
                    "{% endfor %}</div>"
                    "{% endfor %}{% endfor %}{% endif %}</div>"
                )
                out.append(_inject_time_dimension_filter(_content, _td_prefix_grid))
            else:
                fid_list = ", ".join(str(i) for i in kpi_ids)
                fkeys_list = ", ".join(repr(k) for k in field_keys)
                _grid_cell_by_key = (
                    "{% for f in entry.fields %}{% if f.field_key == key %}" + _grid_cell_multi + "{% endif %}{% endfor %}"
                )
                _content = (
                    _display_prefix
                    + _sub_display_prefix
                    + _sub_keys_prefix
                    + f"{{% set kpi_ids_set = [{fid_list}] %}}"
                    f"{{% set field_keys_list = [{fkeys_list}] %}}"
                    '<div class="report-kpi-grid" style="display: grid; grid-template-columns: repeat(auto-fill, minmax(260px, 1fr)); gap: 1rem;">'
                    "{% if kpis %}{% for kpi in kpis %}"
                    "{% if kpi.kpi_id in kpi_ids_set %}"
                    "{% for entry in kpi.entries %}"
                    '<div class="report-card" style="border: 1px solid #ddd; padding: 1rem; border-radius: 8px;">'
                    '<h4 style="margin-top: 0;">{{ kpi.kpi_name }}</h4>'
                    "{% for key in field_keys_list %}"
                    '<p style="margin: 0.25rem 0;"><strong>' + _label_key + ':</strong> ' + _grid_cell_by_key + '</p>'
                    "{% endfor %}</div>"
                    "{% endfor %}{% endif %}{% endfor %}{% endif %}</div>"
                )
                out.append(_inject_time_dimension_filter(_content, _td_prefix_grid))
        elif block_type == "kpi_list":
            block_uid = str(b.get("id") or f"__idx_{bi}__")
            _buid_l = _jinja_quote_block_uid(block_uid)
            _list_ml_prefix_f = (
                "{% set _eid = (entry.entry_id|string) if entry.entry_id is not none else '__none__' %}"
                "{% set _by_ent = ((multi_line_block_rows.get('" + _buid_l + "', {}) | default({})).get((kpi.kpi_id|string), {}) | default({})).get(_eid, {}) | default({}) %}"
                "{% set _f_rows = _by_ent.get(f.field_key, none) %}"
                "{% set _vi = _f_rows if _f_rows is not none else f.value_items %}"
            )
            kpi_ids = b.get("kpiIds") or []
            field_keys = b.get("fieldKeys") or []
            _field_display = b.get("fieldDisplayNames") or b.get("field_display_names") or {}
            _display_parts = [
                f"'{str(k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': '{str(v).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}'"
                for k, v in _field_display.items()
                if v is not None and str(v).strip()
            ]
            _display_prefix = "{% set field_display_names = {" + ", ".join(_display_parts) + "} %}"
            _label_f = "{{ (field_display_names.get(f.field_key) or f.field_name) | default(f.field_name) }}"
            _label_key = "{{ (field_display_names.get(key) or kpi.field_names.get(key, key)) | default(key) }}"
            _sub_field_display = b.get("subFieldDisplayNames") or b.get("sub_field_display_names") or {}
            _sub_display_outer = []
            for _fk, _inner in (_sub_field_display or {}).items():
                if not _inner or not isinstance(_inner, dict):
                    continue
                _inner_parts = [
                    f"'{str(_k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': '{str(_v).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}'"
                    for _k, _v in _inner.items()
                    if _v is not None and str(_v).strip()
                ]
                if _inner_parts:
                    _sub_display_outer.append(f"'{str(_fk).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': {{{', '.join(_inner_parts)}}}")
            _sub_display_prefix = "{% set sub_field_display_names = {" + ", ".join(_sub_display_outer) + "} %}" if _sub_display_outer else "{% set sub_field_display_names = {} %}"
            _sub_label_sf_f = "{{ ((sub_field_display_names.get(f.field_key) or {}) | default({})).get(sf.key) or sf.name | default(sf.name) }}"
            _ml_sub = b.get("multiLineSubFieldKeys") or b.get("multi_line_sub_field_keys") or {}
            _sub_keys_parts = [
                f"'{str(k).replace(chr(92), chr(92)*2).replace(chr(39), chr(92)+chr(39))}': [{', '.join(repr(str(s)) for s in (v or []))}]"
                for k, v in _ml_sub.items()
            ]
            _sub_keys_prefix = "{% set field_sub_field_keys = {" + ", ".join(_sub_keys_parts) + "} %}"
            _list_cell_multi = (
                _list_ml_prefix_f
                + "{% set show_sub_keys = field_sub_field_keys.get(f.field_key, []) | default([]) %}"
                "{% if f.field_type == 'multi_line_items' and _vi %}"
                "<ul style=\"margin: 0.25rem 0;\">{% for item in _vi %}<li>{% for sf in (f.sub_fields | default([])) %}{% if not show_sub_keys or sf.key in show_sub_keys %}{{ item[sf.key] }}{% if not loop.last %} – {% endif %}{% endif %}{% endfor %}</li>{% endfor %}</ul>"
                "{% else %}{{ f.value }}{% endif %}"
            )
            _td_prefix_list, _ = _block_time_dimension_vars(b)
            if not kpi_ids and not field_keys:
                _content = (
                    _display_prefix
                    + _sub_display_prefix
                    + _sub_keys_prefix
                    + '<div class="report-kpi-list">'
                    "{% if kpis %}{% for kpi in kpis %}"
                    '<h4>{{ kpi.kpi_name }}</h4><dl style="margin: 0.5rem 0;">'
                    "{% for entry in kpi.entries %}"
                    "{% for f in entry.fields %}"
                    '<dt style="font-weight: 600;">' + _label_f + '</dt><dd style="margin-left: 1rem;">' + _list_cell_multi + '</dd>'
                    "{% endfor %}{% endfor %}</dl>"
                    "{% endfor %}{% else %}<p>No data.</p>{% endif %}</div>"
                )
                out.append(_inject_time_dimension_filter(_content, _td_prefix_list))
            else:
                fid_list = ", ".join(str(i) for i in kpi_ids)
                fkeys_list = ", ".join(repr(k) for k in field_keys)
                _list_cell_by_key = (
                    "{% for f in entry.fields %}{% if f.field_key == key %}" + _list_cell_multi + "{% endif %}{% endfor %}"
                )
                _content = (
                    _display_prefix
                    + _sub_display_prefix
                    + _sub_keys_prefix
                    + f"{{% set kpi_ids_set = [{fid_list}] %}}"
                    f"{{% set field_keys_list = [{fkeys_list}] %}}"
                    '<div class="report-kpi-list">'
                    "{% if kpis %}{% for kpi in kpis %}"
                    "{% if kpi.kpi_id in kpi_ids_set %}"
                    '<h4>{{ kpi.kpi_name }}</h4><dl style="margin: 0.5rem 0;">'
                    "{% for entry in kpi.entries %}"
                    "{% for key in field_keys_list %}"
                    '<dt style="font-weight: 600;">' + _label_key + '</dt><dd style="margin-left: 1rem;">' + _list_cell_by_key + '</dd>'
                    "{% endfor %}{% endfor %}</dl>"
                    "{% endif %}{% endfor %}{% endif %}</div>"
                )
                out.append(_inject_time_dimension_filter(_content, _td_prefix_list))
    if not out:
        return "<p>No content. Add blocks in the visual designer.</p>"
    return "\n".join(out)


def _jinja_quote_block_uid(s: str) -> str:
    """Escape a block id for embedding in single-quoted Jinja string literals."""
    return s.replace("\\", "\\\\").replace("'", "\\'")


def _multi_line_filter_payload_nonempty(raw: object) -> bool:
    if raw is None:
        return False
    if not isinstance(raw, dict):
        return False
    if raw.get("_version") == 2:
        conds = raw.get("conditions")
        return isinstance(conds, list) and len(conds) > 0
    return len(raw) > 0


async def _build_multi_line_block_rows(
    db: AsyncSession,
    org_id: int,
    prefer_year: int,
    body_blocks: list,
    kpis_payload: list,
    template_kpis: list,
) -> dict:
    """
    Per visual block, per KPI, per entry, per multi-line field key: filtered row list for report rendering.
    Keys: block_uid -> kpi_id str -> entry_id str ('__none__' for placeholder) -> field_key -> list[dict].
    """
    out: dict[str, dict[str, dict[str, dict[str, list]]]] = {}
    kpi_field_by_pair: dict[tuple[int, str], KPIField] = {}
    for kpi in template_kpis:
        for f in kpi.fields or []:
            kpi_field_by_pair[(kpi.id, str(f.key))] = f
    payload_by_kpi_id = {p["kpi_id"]: p for p in kpis_payload if isinstance(p, dict) and p.get("kpi_id") is not None}

    for bi, b in enumerate(body_blocks or []):
        if not isinstance(b, dict):
            continue
        bt = (b.get("type") or "").strip()
        if bt not in ("kpi_table", "kpi_grid", "kpi_list"):
            continue
        block_key = str(b.get("id") or f"__idx_{bi}__")
        mfilters = b.get("multiLineFilters") or b.get("multi_line_filters") or {}
        if not isinstance(mfilters, dict):
            continue
        kpi_ids_block = b.get("kpiIds") or []
        if not isinstance(kpi_ids_block, list):
            kpi_ids_block = []

        for multi_field_key, raw_filter in mfilters.items():
            if not isinstance(multi_field_key, str) or not multi_field_key.strip():
                continue
            if not _multi_line_filter_payload_nonempty(raw_filter):
                continue
            raw_filters = raw_filter if isinstance(raw_filter, dict) else {}

            for kpi in template_kpis:
                if kpi_ids_block and kpi.id not in kpi_ids_block:
                    continue
                field_orm = kpi_field_by_pair.get((kpi.id, multi_field_key))
                if field_orm is None or field_orm.field_type != FieldType.multi_line_items:
                    continue
                kpi_payload = payload_by_kpi_id.get(kpi.id)
                if not kpi_payload:
                    continue
                reference_field_types: dict[str, str] = {}
                for sf in field_orm.sub_fields or []:
                    fk = getattr(sf, "key", "") or ""
                    if not fk:
                        continue
                    ft = getattr(sf.field_type, "value", sf.field_type)
                    reference_field_types[str(fk)] = str(ft)

                entries_list = [e for e in (kpi_payload.get("entries") or []) if isinstance(e, dict)]

                resolution_maps = None
                if raw_filters.get("_version") == 2:
                    conds = raw_filters.get("conditions")
                    if isinstance(conds, list) and conds:
                        # One resolver pass for all rows in this KPI/block (was: once per entry → huge duplicate DB work).
                        all_rows_for_resolve: list[dict] = []
                        for ent in entries_list:
                            field_pl = None
                            for fp in ent.get("fields") or []:
                                if isinstance(fp, dict) and fp.get("field_key") == multi_field_key:
                                    field_pl = fp
                                    break
                            if not field_pl:
                                continue
                            raw_items = field_pl.get("value_items")
                            if not isinstance(raw_items, list):
                                continue
                            for r in raw_items:
                                if isinstance(r, dict):
                                    all_rows_for_resolve.append(r)
                        resolution_maps = await build_reference_resolution_map(
                            db,
                            org_id,
                            prefer_year,
                            field_orm,
                            conds,
                            all_rows_for_resolve,
                        )

                for entry in entries_list:
                    entry_id = entry.get("entry_id")
                    eid_key = str(entry_id) if entry_id is not None else "__none__"
                    field_pl = None
                    for fp in entry.get("fields") or []:
                        if isinstance(fp, dict) and fp.get("field_key") == multi_field_key:
                            field_pl = fp
                            break
                    if not field_pl:
                        continue
                    raw_items = field_pl.get("value_items")
                    if not isinstance(raw_items, list):
                        continue
                    rows_copy = [dict(r) for r in raw_items if isinstance(r, dict)]
                    if raw_filters.get("_version") == 2:
                        filtered = [
                            r
                            for r in rows_copy
                            if row_passes_filters(
                                r,
                                raw_filters,
                                resolution_maps=resolution_maps,
                                reference_field_types=reference_field_types,
                            )
                        ]
                    else:
                        filtered = [r for r in rows_copy if row_passes_filters(r, raw_filters)]

                    out.setdefault(block_key, {}).setdefault(str(kpi.id), {}).setdefault(eid_key, {})[
                        multi_field_key
                    ] = filtered

    return out


def _formulas_need_other_kpi_values(fields: list[KPIField]) -> bool:
    """True if any formula uses KPI_FIELD(...) — only then we need _load_other_kpi_values."""
    if not fields:
        return False
    pat = re.compile(r"\bKPI_FIELD\s*\(", re.IGNORECASE)
    for f in fields:
        if getattr(f, "field_type", None) != FieldType.formula:
            continue
        expr = getattr(f, "formula_expression", None) or ""
        if pat.search(str(expr)):
            return True
    return False


def _extract_kpi_ids_from_blocks(body_blocks: list) -> set[int]:
    """
    Extract KPI ids referenced by designer blocks.

    The frontend stores targeted KPIs on blocks as:
      - kpiIds: number[]
      - kpiId: number (rare, single-target blocks)

    Empty means "all" in the designer UI, so this only returns ids when explicitly provided.
    """
    out: set[int] = set()
    if not isinstance(body_blocks, list) or not body_blocks:
        return out

    stack: list[object] = list(body_blocks)
    while stack:
        cur = stack.pop()
        if isinstance(cur, dict):
            kids = cur.get("kpiIds") or cur.get("kpi_ids") or []
            if isinstance(kids, list):
                for x in kids:
                    try:
                        out.add(int(x))
                    except (TypeError, ValueError):
                        continue
            kid = cur.get("kpiId") or cur.get("kpi_id")
            if kid is not None and kid != "":
                try:
                    out.add(int(kid))
                except (TypeError, ValueError):
                    pass
            # Traverse nested structures defensively (some blocks nest config objects/lists).
            for v in cur.values():
                if isinstance(v, (dict, list)):
                    stack.append(v)
        elif isinstance(cur, list):
            stack.extend(cur)
    return out


def _block_targets_for_time_scope(b: dict, template_kpi_ids: set[int]) -> set[int]:
    """KPI ids a block applies to; empty kpiIds means all template KPIs."""
    raw = b.get("kpiIds") or b.get("kpi_ids") or []
    if not isinstance(raw, list) or not raw:
        return set(template_kpi_ids)
    out: set[int] = set()
    for x in raw:
        try:
            out.add(int(x))
        except (TypeError, ValueError):
            continue
    return out & template_kpi_ids


def _resolve_kpi_time_scope_from_blocks(
    body_blocks: list,
    template_kpi_ids: set[int],
) -> tuple[dict[int, str], dict[int, str]]:
    """
    Match generate_report_data entry loading to designer time settings (kpi_table / grid / list).

    Previously Jinja filtered to latest period at render time only; the backend still loaded every
    period (heavy multi-line + reference resolution). Scope:
      - 'all' — keep all entries for the year
      - 'single' — one period_key
      - 'latest' — last entry after period sort (default in designer)
    """
    scope: dict[int, str] = {k: "latest" for k in template_kpi_ids}
    single_period: dict[int, str] = {}
    if not body_blocks or not template_kpi_ids:
        return scope, single_period

    blocks = [b for b in body_blocks if isinstance(b, dict)]

    for b in blocks:
        bt = (b.get("type") or "").strip()
        if bt not in ("kpi_table", "kpi_grid", "kpi_list"):
            continue
        targets = _block_targets_for_time_scope(b, template_kpi_ids)
        if not targets:
            continue
        mode = (b.get("timeDimensionMode") or b.get("time_dimension_mode") or "latest").strip().lower()
        if mode in ("all_periods", "all"):
            for kid in targets:
                scope[kid] = "all"

    for b in blocks:
        bt = (b.get("type") or "").strip()
        if bt not in ("kpi_table", "kpi_grid", "kpi_list"):
            continue
        targets = _block_targets_for_time_scope(b, template_kpi_ids)
        if not targets:
            continue
        mode = (b.get("timeDimensionMode") or b.get("time_dimension_mode") or "latest").strip().lower()
        if mode != "single_period":
            continue
        pk = b.get("periodKey") or b.get("period_key") or ""
        pk = str(pk).strip() if pk is not None else ""
        if not pk:
            continue
        for kid in targets:
            if scope.get(kid) == "all":
                continue
            scope[kid] = "single"
            single_period[kid] = pk

    return scope, single_period


def _report_period_display(year: int, period_key: str, dimension: TimeDimension) -> str:
    """Human-readable period for report (e.g. '2026 Q1')."""
    if not period_key or not period_key.strip():
        return str(year)
    pk = period_key.strip().upper()
    if pk in ("H1", "H2"):
        return f"{year} H{pk[1]}"
    if pk in ("Q1", "Q2", "Q3", "Q4"):
        return f"{year} {pk}"
    if period_key.isdigit() and 1 <= int(period_key) <= 12:
        months = "Jan Feb Mar Apr May Jun Jul Aug Sep Oct Nov Dec".split()
        return f"{year} {months[int(period_key) - 1]}"
    return f"{year} {period_key}"


async def generate_report_data(
    db: AsyncSession,
    template_id: int,
    org_id: int,
    year: int | None = None,
    include_drafts: bool = False,
) -> dict | None:
    """
    Compile report data from KPI entries for the template.
    Uses KPIs linked on the report template (ReportTemplateKPI); if none are linked, falls back to all org KPIs.
    Returns structured dict: { template_name, year, kpis: [ { kpi_name, entries: [ { fields } ] } ] }
    Formula fields are evaluated.
    """
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None
    yr = year if year is not None else datetime.date.today().year
    t0 = time.perf_counter()
    cache_key = (template_id, org_id, int(yr), bool(include_drafts), "v3")
    cached = _cache_get(cache_key)
    if cached is not None:
        _prof(f"CACHE HIT key={cache_key}")
        return cached

    # Collect requested reference-derived columns for multi-line items from designer blocks.
    # Frontend encodes a reference-derived column key as "__ref__{subKey}__{encodedChain}" where encodedChain is
    # a ">>"-joined list of paths ("fieldKey" or "fieldKey|subKey") URL-encoded.
    ref_cols_by_multi_field: dict[str, list[tuple[str, str, list[str]]]] = {}
    try:
        blocks = getattr(rt, "body_blocks", None) or []
        for b in blocks if isinstance(blocks, list) else []:
            if not isinstance(b, dict):
                continue
            if (b.get("type") or "").strip() != "kpi_table":
                continue
            ml = b.get("multiLineSubFieldKeys") or b.get("multi_line_sub_field_keys") or {}
            if not isinstance(ml, dict):
                continue
            for multi_field_key, keys in ml.items():
                if not isinstance(keys, list):
                    continue
                for k in keys:
                    if not isinstance(k, str) or not k.startswith("__ref__"):
                        continue
                    # "__ref__{subKey}__{encodedChain}"
                    parts = k.split("__", 3)
                    if len(parts) < 4:
                        continue
                    sub_key = parts[2]
                    enc = parts[3]
                    try:
                        from urllib.parse import unquote
                        chain_s = unquote(enc)
                    except Exception:
                        chain_s = enc
                    chain = [p for p in (chain_s.split(">>") if chain_s else []) if p]
                    if not chain:
                        continue
                    ref_cols_by_multi_field.setdefault(str(multi_field_key), []).append((k, sub_key, chain))
    except Exception:
        ref_cols_by_multi_field = {}
    _prof(f"template={template_id} org={org_id} year={yr} ref_col_fields={len(ref_cols_by_multi_field)}")

    # KPIs attached to this template (not the whole org — that was O(org_kpis × entries) and unusably slow).
    t_kpis0 = time.perf_counter()
    rtk_result = await db.execute(
        select(ReportTemplateKPI)
        .where(ReportTemplateKPI.report_template_id == template_id)
        .order_by(ReportTemplateKPI.sort_order, ReportTemplateKPI.id)
        .options(
            selectinload(ReportTemplateKPI.kpi).selectinload(KPI.fields).selectinload(KPIField.sub_fields),
            selectinload(ReportTemplateKPI.fields),
        )
    )
    rtk_list = list(rtk_result.scalars().unique().all())

    kpi_worklist: list[tuple[KPI, list[KPIField]]] = []
    if rtk_list:
        seen_kpi_ids: set[int] = set()
        for rtk in rtk_list:
            kpi = rtk.kpi
            if not kpi or kpi.id in seen_kpi_ids:
                continue
            seen_kpi_ids.add(kpi.id)
            if rtk.include_all_fields:
                fts = sorted(list(kpi.fields or []), key=lambda f: (f.sort_order, f.id))
            else:
                tf_by_field_id = {tf.kpi_field_id: tf for tf in (rtk.fields or [])}
                fts = [f for f in (kpi.fields or []) if f.id in tf_by_field_id]
                fts.sort(
                    key=lambda f: (
                        tf_by_field_id[f.id].sort_order,
                        f.sort_order,
                        f.id,
                    )
                )
            kpi_worklist.append((kpi, fts))
    else:
        # Template has no attached KPI list yet.
        # For designer templates, use block-scoped KPI ids to avoid loading the entire org
        # (which can be very slow with multi_line_items).
        raw_blocks = getattr(rt, "body_blocks", None) or []
        block_kpi_ids = _extract_kpi_ids_from_blocks(raw_blocks)
        if block_kpi_ids:
            result = await db.execute(
                select(KPI)
                .where(KPI.organization_id == org_id, KPI.id.in_(block_kpi_ids))
                .order_by(KPI.sort_order, KPI.name)
                .options(selectinload(KPI.fields).selectinload(KPIField.sub_fields))
            )
        else:
            # Legacy/code templates: fall back to all org KPIs (previous behavior).
            result = await db.execute(
                select(KPI)
                .where(KPI.organization_id == org_id)
                .order_by(KPI.sort_order, KPI.name)
                .options(selectinload(KPI.fields).selectinload(KPIField.sub_fields))
            )
        for kpi in result.unique().scalars().all():
            fts = sorted(list(kpi.fields or []), key=lambda f: (f.sort_order, f.id))
            kpi_worklist.append((kpi, fts))

    template_kpis = [kp for kp, _ in kpi_worklist]
    _prof(
        f"loaded_kpis={len(template_kpis)} rtk={len(rtk_list)} ms={(time.perf_counter()-t_kpis0)*1000:.1f}"
    )

    raw_blocks = getattr(rt, "body_blocks", None) or []
    template_kpi_ids_set = {kp.id for kp, _ in kpi_worklist}
    if not raw_blocks:
        # Code-only / legacy templates: keep full year (same as before visual blocks).
        kpi_td_scope = {kid: "all" for kid in template_kpi_ids_set}
        kpi_td_single_period: dict[int, str] = {}
    else:
        kpi_td_scope, kpi_td_single_period = _resolve_kpi_time_scope_from_blocks(
            raw_blocks, template_kpi_ids_set
        )

    # Only build domains/categories when the template actually uses domain-driven blocks.
    _has_domain_blocks = False
    if isinstance(raw_blocks, list) and raw_blocks:
        for b in raw_blocks:
            if not isinstance(b, dict):
                continue
            bt = (b.get("type") or "").strip()
            if bt in ("domain_list", "domain_categories", "domain_kpis"):
                _has_domain_blocks = True
                break

    # Load text blocks
    t_tb0 = time.perf_counter()
    text_blocks_result = await db.execute(
        select(ReportTemplateTextBlock)
        .where(ReportTemplateTextBlock.report_template_id == template_id)
        .order_by(ReportTemplateTextBlock.sort_order, ReportTemplateTextBlock.id)
    )
    text_blocks = [
        {"id": tb.id, "title": tb.title, "content": tb.content, "sort_order": tb.sort_order}
        for tb in text_blocks_result.scalars().all()
    ]
    _prof(f"text_blocks={len(text_blocks)} ms={(time.perf_counter()-t_tb0)*1000:.1f}")

    out = {
        "template_name": rt.name,
        "template_id": rt.id,
        "year": yr,
        "text_blocks": text_blocks,
        "kpis": [],
    }
    org = await db.get(Organization, org_id)
    org_td = TimeDimension(getattr(org, "time_dimension", None) or "yearly") if org else TimeDimension.YEARLY

    total_entries_loaded = 0
    total_ml_rows = 0
    total_entries_query_ms = 0.0
    total_ml_load_ms = 0.0
    total_ref_col_ms = 0.0

    for kpi, fields_to_include in kpi_worklist:
        t_kpi0 = time.perf_counter()
        kpi_td_raw = getattr(kpi, "time_dimension", None)
        kpi_td = TimeDimension(kpi_td_raw) if kpi_td_raw else None
        effective_td = effective_kpi_time_dimension(kpi_td, org_td)

        entry_filters = [
            KPIEntry.organization_id == org_id,
            KPIEntry.kpi_id == kpi.id,
            KPIEntry.year == yr,
        ]
        if not include_drafts:
            entry_filters.append(KPIEntry.is_draft == False)
        t_eq0 = time.perf_counter()
        entries_result = await db.execute(
            select(KPIEntry)
            .where(*entry_filters)
            .options(selectinload(KPIEntry.field_values))
        )
        total_entries_query_ms += (time.perf_counter() - t_eq0) * 1000.0
        all_entries = list(entries_result.scalars().all())
        # Sort by period (e.g. Q1, Q2, Q3, Q4) so "latest" filter returns last; report can show all or one period
        entries_sorted = sorted(
            all_entries,
            key=lambda e: period_key_sort_order(getattr(e, "period_key", "") or "", effective_td),
        )
        # Match designer time filter (latest / single period / all) — was Jinja-only; loading all periods was the main cost.
        td_scope = kpi_td_scope.get(kpi.id, "latest")
        td_pk = kpi_td_single_period.get(kpi.id)
        if td_scope == "latest" and len(entries_sorted) > 1:
            entries_sorted = [entries_sorted[-1]]
        elif td_scope == "single" and td_pk is not None:
            pk_s = str(td_pk).strip()
            entries_sorted = [
                e
                for e in entries_sorted
                if (getattr(e, "period_key", "") or "").strip() == pk_s
            ]

        need_cross_kpi = _formulas_need_other_kpi_values(fields_to_include)
        other_kpi_values = (
            await _load_other_kpi_values(db, yr, org_id, kpi.id)
            if entries_sorted and need_cross_kpi
            else {}
        )
        entry_ids_sorted = [e.id for e in entries_sorted]
        total_entries_loaded += len(entry_ids_sorted)
        ml_fields = [f for f in fields_to_include if f.field_type == FieldType.multi_line_items]
        ml_rows_by_field_id: dict[int, dict[int, list[dict]]] = {}
        for mf in ml_fields:
            t_ml0 = time.perf_counter()
            ml_rows_by_field_id[mf.id] = await _load_multi_line_items_rows_batch(
                db, entry_ids=entry_ids_sorted, field=mf
            )
            total_ml_load_ms += (time.perf_counter() - t_ml0) * 1000.0
            for _eid, _rows in ml_rows_by_field_id[mf.id].items():
                total_ml_rows += len(_rows or [])
        # Build value by entry and field (one row per entry; each row has period_key and period_display)
        rows = []
        if not entries_sorted:
            # No submitted data for this KPI: provide one placeholder entry so the report shows "No data entered"
            _NO_DATA_PLACEHOLDER = "No data entered"
            field_values_out = []
            for f in fields_to_include:
                card_ids = kpi.card_display_field_ids or []
                show_on_card = f.id in card_ids if isinstance(card_ids, list) else False
                field_payload = {
                    "field_key": f.key,
                    "field_name": f.name,
                    "value": _NO_DATA_PLACEHOLDER,
                    "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                    "show_on_card": show_on_card,
                }
                if f.field_type == FieldType.multi_line_items:
                    sub_fields_orm = getattr(f, "sub_fields") or []
                    field_payload["value_items"] = []
                    field_payload["sub_field_keys"] = [sf.key for sf in sub_fields_orm]
                    field_payload["sub_fields"] = [{"key": sf.key, "name": getattr(sf, "name", sf.key)} for sf in sub_fields_orm]
                field_values_out.append(field_payload)
            rows.append({"entry_id": None, "fields": field_values_out, "period_key": "", "period_display": str(yr)})
        else:
            # One reference-resolution pass per multi-line field (was: once per time period / entry).
            ref_col_resolve_by_field_id: dict[int, tuple[list[tuple[str, str, dict]], dict]] = {}
            for f in fields_to_include:
                if f.field_type != FieldType.multi_line_items:
                    continue
                requested_pre = ref_cols_by_multi_field.get(str(f.key), [])
                if not requested_pre:
                    continue
                conditions_pre: list[dict] = []
                syn_meta_pre: list[tuple[str, str, dict]] = []
                for syn_key, sub_key, chain in requested_pre:
                    steps = []
                    for p in chain:
                        if "|" in p:
                            fk, sk = p.split("|", 1)
                            steps.append({"compare_field_key": fk, "compare_sub_field_key": sk})
                        else:
                            steps.append({"compare_field_key": p})
                    rr = {"chain": steps}
                    conditions_pre.append({"field": sub_key, "reference_resolution": rr})
                    syn_meta_pre.append((syn_key, sub_key, rr))
                all_rows_merge: list[dict] = []
                for ent in entries_sorted:
                    for r in ml_rows_by_field_id.get(f.id, {}).get(ent.id, []):
                        if isinstance(r, dict):
                            all_rows_merge.append(r)
                t_ref0 = time.perf_counter()
                res_map_pre = await build_reference_resolution_map(
                    db,
                    org_id,
                    prefer_year=yr,
                    field=f,
                    conditions=conditions_pre,
                    row_dicts=all_rows_merge,
                )
                total_ref_col_ms += (time.perf_counter() - t_ref0) * 1000.0
                ref_col_resolve_by_field_id[f.id] = (syn_meta_pre, res_map_pre)

            for entry in entries_sorted:
                _pk = getattr(entry, "period_key", "") or ""
                _pd = _report_period_display(yr, _pk, effective_td)
                fv_by_field = {fv.field_id: fv for fv in entry.field_values}
                value_by_key = {}
                field_values_out = []
                multi_line_items_data = {}
                for f in fields_to_include:
                    # Skip formula fields here; they are added once with computed value in the loop below
                    if f.field_type == FieldType.formula:
                        continue
                    fv = fv_by_field.get(f.id)
                    val = None
                    if fv:
                        # IMPORTANT: preserve falsy values like 0, False, and empty list.
                        if fv.value_date is not None:
                            val = fv.value_date.isoformat() if hasattr(fv.value_date, "isoformat") else str(fv.value_date)
                        elif fv.value_text is not None:
                            val = fv.value_text
                        elif fv.value_number is not None:
                            val = fv.value_number
                        elif fv.value_json is not None:
                            val = fv.value_json
                        elif fv.value_boolean is not None:
                            val = fv.value_boolean
                        if f.field_type == FieldType.number and fv.value_number is not None:
                            value_by_key[f.key] = fv.value_number
                        if f.field_type == FieldType.multi_line_items:
                            # Multi-line items are stored relationally (loaded in batch per KPI above).
                            rows_items = ml_rows_by_field_id.get(f.id, {}).get(entry.id, [])
                            multi_line_items_data[f.key] = rows_items
                            val = rows_items
                    card_ids = kpi.card_display_field_ids or []
                    show_on_card = f.id in card_ids if isinstance(card_ids, list) else False
                    field_payload = {
                        "field_key": f.key,
                        "field_name": f.name,
                        "value": val,
                        "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                        "show_on_card": show_on_card,
                    }
                    if f.field_type == FieldType.multi_line_items:
                        sub_fields_orm = getattr(f, "sub_fields") or []
                        field_payload["sub_field_keys"] = [sf.key for sf in sub_fields_orm]
                        field_payload["sub_fields"] = [{"key": sf.key, "name": getattr(sf, "name", sf.key)} for sf in sub_fields_orm]
                        if isinstance(val, list):
                            field_payload["value_items"] = val
                        else:
                            field_payload["value_items"] = []
                        # Inject reference-derived columns (synthetic keys) into value_items and sub_fields so Jinja can render them.
                        requested = ref_cols_by_multi_field.get(str(f.key), [])
                        if requested and isinstance(field_payload.get("value_items"), list):
                            rows_list = field_payload["value_items"]
                            packed_ref = ref_col_resolve_by_field_id.get(f.id)
                            if not packed_ref:
                                syn_meta = []
                                res_map = {}
                            else:
                                syn_meta, res_map = packed_ref
                            # Apply to each row; keep raw ref cell for lookup.
                            for row in rows_list:
                                if not isinstance(row, dict):
                                    continue
                                for cond_idx, (syn_key, sub_key, _rr) in enumerate(syn_meta):
                                    cell = row.get(sub_key)
                                    # multi_reference: resolve each piece; else resolve single.
                                    ft = None
                                    sub_obj = next((s for s in (sub_fields_orm or []) if getattr(s, "key", None) == sub_key), None)
                                    if sub_obj is not None:
                                        ft = getattr(sub_obj, "field_type", None)
                                        ft = ft.value if hasattr(ft, "value") else ft
                                    if ft == "multi_reference":
                                        vals = []
                                        for piece in _multi_raw_pieces(cell):
                                            lab = _normalize_reference_value(_extract_ref_label(piece))
                                            if not lab:
                                                continue
                                            v = res_map.get((cond_idx, lab))
                                            if v is None:
                                                continue
                                            vals.append(v)
                                        row[syn_key] = ", ".join(str(x) for x in vals) if vals else None
                                    else:
                                        lab = _normalize_reference_value(_extract_ref_label(cell))
                                        row[syn_key] = res_map.get((cond_idx, lab)) if lab else None
                            # Expose synthetic keys as sub_fields (so the table header can include them)
                            sub_fields_out = field_payload.get("sub_fields") or []
                            if isinstance(sub_fields_out, list):
                                for syn_key, sub_key, chain in requested:
                                    # Default name: "subKey → <terminal>"
                                    terminal = chain[-1] if chain else ""
                                    sub_fields_out.append({"key": syn_key, "name": f"{sub_key} → {terminal}", "field_type": "resolved_reference"})
                                field_payload["sub_fields"] = sub_fields_out
                                field_payload["sub_field_keys"] = [sf.get("key") for sf in sub_fields_out if isinstance(sf, dict) and sf.get("key")]
                    field_values_out.append(field_payload)
                    if val is not None and f.field_type == FieldType.number:
                        value_by_key[f.key] = val

                # Seed existing stored formula values as baseline for dependencies.
                # This mirrors entries.service.recompute_formula_fields_for_kpi and prevents
                # report formulas that reference other formula fields from evaluating to blank.
                for f in fields_to_include:
                    if f.field_type != FieldType.formula:
                        continue
                    fv_formula = fv_by_field.get(f.id)
                    if not fv_formula or fv_formula.value_number is None:
                        continue
                    try:
                        value_by_key[f.key] = float(fv_formula.value_number)
                    except (TypeError, ValueError):
                        continue

                # Formula fields (with multi_line_items support for SUM_ITEMS etc.)
                for f in fields_to_include:
                    if f.field_type == FieldType.formula and f.formula_expression:
                        computed = evaluate_formula(
                            f.formula_expression,
                            value_by_key,
                            multi_line_items_data,
                            other_kpi_values,
                        )
                        # If evaluation fails (returns None), fall back to the stored formula value
                        # so reports can still display existing computed values.
                        if computed is None:
                            fv_formula = fv_by_field.get(f.id)
                            if fv_formula and fv_formula.value_number is not None:
                                computed = fv_formula.value_number
                        card_ids_f = kpi.card_display_field_ids or []
                        show_on_card_f = f.id in card_ids_f if isinstance(card_ids_f, list) else False
                        field_values_out.append({
                            "field_key": f.key,
                            "field_name": f.name,
                            "value": computed,
                            "field_type": f.field_type.value if hasattr(f.field_type, "value") else str(f.field_type),
                            "show_on_card": show_on_card_f,
                        })
                        if computed is not None:
                            value_by_key[f.key] = computed
                rows.append({
                    "entry_id": entry.id,
                    "fields": field_values_out,
                    "period_key": _pk,
                    "period_display": _pd,
                })
        # Map field_key -> field_name for template headers when only key is in scope
        field_names_map = {}
        if rows and rows[0].get("fields"):
            for f in rows[0]["fields"]:
                fkey = f.get("field_key")
                if fkey is not None:
                    field_names_map[fkey] = f.get("field_name", fkey)
        _latest_period = ""
        _latest_display = str(yr)
        if rows and rows[-1].get("period_display"):
            _latest_display = rows[-1]["period_display"]
            _latest_period = rows[-1].get("period_key") or ""
        out["kpis"].append({
            "kpi_id": kpi.id,
            "kpi_name": kpi.name,
            "entries": rows,
            "field_names": field_names_map,
            "period_display": _latest_display,
            "time_dimension_used": effective_td.value,
        })
        _prof(
            f"kpi={kpi.id} scope={td_scope} entries={len(entries_sorted)} ml_fields={len(ml_fields)} ms={(time.perf_counter()-t_kpi0)*1000:.1f}"
        )

    t_bl0 = time.perf_counter()
    # Only build multi_line_block_rows when there are active multi-line filters in blocks.
    # Otherwise the Jinja template falls back to f.value_items directly and this work is wasted.
    body_blocks = getattr(rt, "body_blocks", None) or []
    has_ml_filters = False
    if isinstance(body_blocks, list) and body_blocks:
        for b in body_blocks:
            if not isinstance(b, dict):
                continue
            mfilters = b.get("multiLineFilters") or b.get("multi_line_filters") or {}
            if isinstance(mfilters, dict) and any(_multi_line_filter_payload_nonempty(v) for v in mfilters.values()):
                has_ml_filters = True
                break
    if has_ml_filters:
        out["multi_line_block_rows"] = await _build_multi_line_block_rows(
            db,
            org_id,
            yr,
            body_blocks,
            out["kpis"],
            template_kpis,
        )
        _prof(f"multi_line_block_rows ms={(time.perf_counter()-t_bl0)*1000:.1f}")
    else:
        out["multi_line_block_rows"] = {}
        _prof("multi_line_block_rows skipped (no active multi-line filters)")

    # Build domains → categories → KPIs for template access (all org domains; KPIs in template)
    out["domains"] = []
    if _has_domain_blocks and out["kpis"]:
        t_dom0 = time.perf_counter()
        kpi_payload_by_id = {p["kpi_id"]: p for p in out["kpis"]}
        template_kpi_ids = set(kpi_payload_by_id.keys())
        domains_result = await db.execute(
            select(Domain)
            .where(Domain.organization_id == org_id)
            .order_by(Domain.sort_order, Domain.name)
            .options(selectinload(Domain.categories))
        )
        domains_orm = list(domains_result.unique().scalars().all())
        category_ids = [c.id for d in domains_orm for c in (d.categories or [])]
        category_to_kpi_ids = defaultdict(list)
        if category_ids:
            kc_result = await db.execute(
                select(KPICategory.kpi_id, KPICategory.category_id).where(
                    KPICategory.category_id.in_(category_ids),
                    KPICategory.kpi_id.in_(template_kpi_ids),
                )
            )
            for kpi_id, cat_id in kc_result.all():
                category_to_kpi_ids[cat_id].append(kpi_id)
        for d in domains_orm:
            categories_out = []
            for cat in sorted(d.categories or [], key=lambda c: (c.sort_order, c.name)):
                kpi_ids_in_cat = category_to_kpi_ids.get(cat.id, [])
                kpis_in_cat = [kpi_payload_by_id[kid] for kid in kpi_ids_in_cat if kid in kpi_payload_by_id]
                categories_out.append({
                    "id": cat.id,
                    "name": cat.name,
                    "kpis": kpis_in_cat,
                })
            out["domains"].append({
                "id": d.id,
                "name": d.name,
                "categories": categories_out,
            })
        _prof(f"domains ms={(time.perf_counter()-t_dom0)*1000:.1f}")

    _prof(
        f"TOTAL entries={total_entries_loaded} ml_rows={total_ml_rows} "
        f"entries_q_ms={total_entries_query_ms:.1f} ml_ms={total_ml_load_ms:.1f} refcol_ms={total_ref_col_ms:.1f} "
        f"total_ms={(time.perf_counter()-t0)*1000:.1f}"
    )
    _cache_set(cache_key, out)
    return out


async def render_report_html(
    db: AsyncSession,
    template_id: int,
    org_id: int,
    year: int | None = None,
    include_drafts: bool = False,
    report_data: dict | None = None,
) -> str | None:
    """
    Render report using the template's body_template or body_blocks and
    the structured KPI data produced by generate_report_data.
    When body_blocks is set, body_template is generated from it first.
    Pass report_data to reuse data from a prior generate_report_data call.
    """
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None
    body_template = rt.body_template
    mode = (getattr(rt, "template_mode", "designer") or "designer").strip().lower()
    if mode != "code" and getattr(rt, "body_blocks", None):
        body_template = _blocks_to_jinja(rt.body_blocks)
    if not body_template:
        return None
    return await render_report_html_with_template(
        db,
        template_id,
        org_id,
        year=year,
        body_template_override=body_template,
        include_drafts=include_drafts,
        report_data=report_data,
    )


async def render_report_html_with_template(
    db: AsyncSession,
    template_id: int,
    org_id: int,
    year: int | None = None,
    body_template_override: str | None = None,
    include_drafts: bool = False,
    report_data: dict | None = None,
) -> str | None:
    """
    Render report with given template string (for live preview) or from DB.
    Uses generate_report_data and Jinja2 render.
    If report_data is provided (e.g. caller already ran generate_report_data), skips regenerating.
    """
    if report_data is not None:
        data = report_data
    else:
        data = await generate_report_data(
            db,
            template_id,
            org_id,
            year=year,
            include_drafts=include_drafts,
        )
    if not data:
        return None
    if body_template_override:
        body_template = body_template_override
    else:
        rt = await get_report_template(db, template_id, org_id)
        if not rt:
            return None
        mode = (getattr(rt, "template_mode", "designer") or "designer").strip().lower()
        body_template = rt.body_template or ""
        if mode != "code" and getattr(rt, "body_blocks", None):
            body_template = _blocks_to_jinja(rt.body_blocks)
    if not body_template:
        return None
    template = _jinja_env.from_string(body_template)
    return template.render(**data)


async def evaluate_report_snippet(
    db: AsyncSession,
    template_id: int,
    org_id: int,
    snippet_type: str,
    year: int | None = None,
    kpi_id: int | None = None,
    field_key: str | None = None,
    sub_field_key: str | None = None,
    sub_field_group_fn: str | None = None,
    entry_index: int = 0,
    expression: str | None = None,
    include_drafts: bool = False,
) -> str | int | float | None:
    """
    Evaluate a single KPI value or formula in report context for preview.
    Returns the computed value or None if not found / error.
    """
    rt = await get_report_template(db, template_id, org_id)
    if not rt:
        return None
    yr = year if year is not None else datetime.date.today().year
    data = await generate_report_data(
        db,
        template_id,
        org_id,
        year=yr,
        include_drafts=include_drafts,
    )
    if not data or "kpis" not in data:
        return None
    kpis = data["kpis"]

    if snippet_type == "formula":
        if expression is None or kpi_id is None:
            return None
        result = _evaluate_report_formula(kpis, expression.strip(), kpi_id, entry_index)
        return result if result != "" else None

    if snippet_type == "kpi_value":
        if kpi_id is None or not field_key:
            return None
        if sub_field_key and sub_field_group_fn:
            formula_expr = f"{sub_field_group_fn.strip()}({field_key}, {sub_field_key})"
            result = _evaluate_report_formula(kpis, formula_expr, kpi_id, entry_index)
            return result if result != "" else None
        val = _get_kpi_field_value(
            kpis, kpi_id, field_key, sub_field_key or None, entry_index
        )
        return val if val != "" else None

    return None


def group_dependent_fields(original_fields: list) -> list:
    """Topologically sort fields so that dependents appear directly below their trigger field."""
    result = []
    visited = set()
    dependents_map = {}
    for f in original_fields:
        trigger_id = f.config.get("condition_trigger_field_id") if f.config else None
        if trigger_id is not None:
            try:
                trigger_id = int(trigger_id)
                dependents_map.setdefault(trigger_id, []).append(f)
            except (ValueError, TypeError):
                continue

    def insert_field(f):
        if f.id in visited:
            return
        visited.add(f.id)
        result.append(f)
        dependents = dependents_map.get(f.id, [])
        dependents.sort(key=lambda x: (x.sort_order, x.id))
        for dep in dependents:
            insert_field(dep)

    sorted_orig = sorted(original_fields, key=lambda x: (x.sort_order, x.id))
    for f in sorted_orig:
        trigger_id = f.config.get("condition_trigger_field_id") if f.config else None
        trigger_exists = False
        if trigger_id is not None:
            try:
                trigger_id = int(trigger_id)
                trigger_exists = any(x.id == trigger_id for x in sorted_orig)
            except (ValueError, TypeError):
                pass
        if not trigger_exists:
            insert_field(f)

    for f in sorted_orig:
        if f.id not in visited:
            insert_field(f)
    return result


def is_field_visible(f, fields_by_id: dict, values_by_field_id: dict) -> bool:
    """Recursively check conditional visibility of a field."""
    from app.fields.conditional import is_field_visible as _is_field_visible
    return _is_field_visible(f, fields_by_id, values_by_field_id)


from reportlab.pdfgen import canvas

class NumberedCanvas(canvas.Canvas):
    """Two-pass ReportLab canvas to draw exact page numbers, headers, and footers on each page."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []
        self.organization_name = ""
        self.confidentiality_text = "Confidential Document"
        self.include_date = True

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_elements(num_pages)
            super().showPage()
        super().save()

    def draw_page_elements(self, page_count):
        from reportlab.lib import colors
        import datetime
        self.saveState()
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(colors.HexColor("#4b5563"))
        
        # Running Header
        pass

        # Running Footer
        self.setStrokeColor(colors.HexColor("#e5e7eb"))
        self.setLineWidth(0.5)
        self.line(54, 55, 612 - 54, 55)
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#6b7280"))
        
        footer_left = f"{self.confidentiality_text}"
        if self.organization_name:
            footer_left += f" | {self.organization_name}"
        self.drawString(54, 40, footer_left)
        
        date_str = ""
        if self.include_date:
            date_str = f"Generated on {datetime.date.today().strftime('%B %d, %Y')} | "
        footer_right = f"{date_str}Page {self._pageNumber} of {page_count}"
        self.drawRightString(612 - 54, 40, footer_right)
        self.restoreState()


async def generate_kpi_pdf_report(
    db: AsyncSession,
    organization_id: int,
    kpi_id: int,
    year: int,
    period_key: str,
    configuration: dict,
    requesting_user_id: int | None = None
) -> bytes:
    import html
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    
    from app.core.models import KPI, KPIField, KPIEntry, KPIFieldValue, Organization, FieldType
    from app.entries.routes import _display_string_for_pdf_export
    from sqlalchemy.orm import selectinload
    
    # 1. Fetch KPI
    kpi_stmt = (
        select(KPI)
        .where(KPI.id == kpi_id, KPI.organization_id == organization_id)
        .options(selectinload(KPI.fields).selectinload(KPIField.sub_fields))
    )
    kpi_res = await db.execute(kpi_stmt)
    kpi = kpi_res.scalar_one_or_none()
    if not kpi:
        raise ValueError("KPI not found")
 
    # 2. Fetch Entry
    entry = None
    if requesting_user_id is not None:
        entry_stmt = select(KPIEntry).where(
            KPIEntry.organization_id == organization_id,
            KPIEntry.kpi_id == kpi_id,
            KPIEntry.year == year,
            KPIEntry.period_key == period_key,
            KPIEntry.is_draft == True,
            KPIEntry.user_id == requesting_user_id
        )
        entry_res = await db.execute(entry_stmt)
        entry = entry_res.scalar_one_or_none()
        
    if not entry:
        entry_stmt = select(KPIEntry).where(
            KPIEntry.organization_id == organization_id,
            KPIEntry.kpi_id == kpi_id,
            KPIEntry.year == year,
            KPIEntry.period_key == period_key,
            KPIEntry.is_draft == False
        )
        entry_res = await db.execute(entry_stmt)
        entry = entry_res.scalar_one_or_none()
        
    if not entry:
        raise ValueError(f"KPI Entry for year {year} and period '{period_key}' not found")

    # 3. Fetch Scalar Values
    scalar_values_query = select(KPIFieldValue).where(KPIFieldValue.entry_id == entry.id)
    scalar_values_res = await db.execute(scalar_values_query)
    scalar_values = scalar_values_res.scalars().all()
    
    values_by_field_id = {}
    for val in scalar_values:
        if val.value_boolean is not None:
            values_by_field_id[val.field_id] = val.value_boolean
        elif val.value_number is not None:
            values_by_field_id[val.field_id] = val.value_number
        elif val.value_date is not None:
            values_by_field_id[val.field_id] = val.value_date
        elif val.value_json is not None:
            values_by_field_id[val.field_id] = val.value_json
        else:
            values_by_field_id[val.field_id] = val.value_text

    # 4. Extract Configuration
    title = (configuration.get("report_header") or configuration.get("title") or "").strip() or f"{kpi.name} Report"
    kpi_name_override = (configuration.get("kpi_name_override") or "").strip() or kpi.name
    custom_header = (configuration.get("custom_header") or "").strip()
    custom_subheader = (configuration.get("custom_subheader") or "").strip()
    organization_info = (configuration.get("organization_info") or "").strip()
    include_generation_date = configuration.get("include_generation_date", True)

    # Resolve time dimensions
    org = await db.get(Organization, organization_id)
    org_td = TimeDimension(getattr(org, "time_dimension", None) or "yearly") if org else TimeDimension.YEARLY
    kpi_td_raw = getattr(kpi, "time_dimension", None)
    kpi_td = TimeDimension(kpi_td_raw) if kpi_td_raw else None
    effective_td = effective_kpi_time_dimension(kpi_td, org_td)

    # 5. Setup Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "KpiReportTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#1e3a8a"),
        alignment=TA_CENTER,
        spaceAfter=12
    )
    kpi_name_style = ParagraphStyle(
        "KpiReportName",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#0f766e"),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )
    section_heading_style = ParagraphStyle(
        "KpiSectionHeading",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#111827"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    body_label_style = ParagraphStyle(
        "KpiBodyLabel",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#374151")
    )
    body_value_style = ParagraphStyle(
        "KpiBodyValue",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#1f2937")
    )
    table_title_style = ParagraphStyle(
        "KpiTableTitle",
        parent=styles["Heading4"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1e3a8a"),
        spaceBefore=14,
        spaceAfter=4,
        keepWithNext=False
    )
    table_heading_style = ParagraphStyle(
        "KpiTableHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#4b5563"),
        spaceAfter=2,
        keepWithNext=False
    )
    table_subheader_style = ParagraphStyle(
        "KpiTableSubheader",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#6b7280"),
        spaceAfter=6,
        keepWithNext=False
    )
    th_style = ParagraphStyle(
        "KpiTableHeaderCell",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8,
        leading=10,
        textColor=colors.white
    )
    th_sr_style = ParagraphStyle("KpiTableHeaderSrCell", parent=th_style, alignment=TA_CENTER)
    td_style = ParagraphStyle(
        "KpiTableBodyCell",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#1f2937")
    )
    td_sr_style = ParagraphStyle("KpiTableBodySrCell", parent=td_style, alignment=TA_CENTER)

    elements = []

    # Main Title (Report Header)
    elements.append(Paragraph(html.escape(title), title_style))
    elements.append(Spacer(1, 10))

    # Optional Description
    description = (configuration.get("description") or "").strip()
    if description:
        desc_heading_style = ParagraphStyle(
            "DescHeading",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#1e3a8a"),
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True
        )
        elements.append(Paragraph("Description", desc_heading_style))
        
        desc_text_style = ParagraphStyle(
            "DescText",
            parent=styles["Normal"],
            fontSize=9,
            leading=13,
            textColor=colors.HexColor("#374151"),
            spaceAfter=10
        )
        elements.append(Paragraph(html.escape(description), desc_text_style))

    # 6. Scalar Fields Table (Section Grouped & Borderless)
    scalar_fields = [f for f in kpi.fields if f.field_type != FieldType.multi_line_items]
    fields_by_id = {f.id: f for f in kpi.fields}
    sorted_scalars = group_dependent_fields(scalar_fields)
    
    excluded_fields = configuration.get("excluded_scalar_fields") or []
    excluded_set = {str(item) for item in excluded_fields}
    
    visible_scalars = [
        f for f in sorted_scalars 
        if is_field_visible(f, fields_by_id, values_by_field_id)
        and f.key not in excluded_set
        and str(f.id) not in excluded_set
    ]

    # Apply custom scalar fields order if specified
    ordered_scalar_ids = configuration.get("ordered_scalar_fields")
    if ordered_scalar_ids:
        order_map = {str(val): idx for idx, val in enumerate(ordered_scalar_ids)}
        def get_sort_key(f):
            if str(f.id) in order_map:
                return order_map[str(f.id)]
            if f.key in order_map:
                return order_map[f.key]
            return 999999
        visible_scalars.sort(key=get_sort_key)

    # Get excluded multi-line fields and compute visible ones
    multi_line_fields = [f for f in kpi.fields if f.field_type == FieldType.multi_line_items]
    excluded_ml_keys = configuration.get("excluded_multi_line_fields") or []
    excluded_ml_set = {str(item) for item in excluded_ml_keys}
    visible_multi_line = [
        f for f in multi_line_fields
        if is_field_visible(f, fields_by_id, values_by_field_id)
        and f.key not in excluded_ml_set and str(f.id) not in excluded_ml_set
    ]

    # Build parent blocks for scalar fields to keep parents and their children grouped
    scalar_blocks = []
    current_block = []
    
    for f in visible_scalars:
        parent_id = f.config.get("condition_trigger_field_id") if f.config else None
        if parent_id is not None:
            try:
                parent_id = int(parent_id)
            except (ValueError, TypeError):
                parent_id = None
                
        is_child = False
        if parent_id is not None:
            parent_field = fields_by_id.get(parent_id)
            if parent_field and parent_field in visible_scalars:
                is_child = True
                
        if is_child:
            if current_block:
                current_block.append(f)
            else:
                current_block = [f]
        else:
            if current_block:
                scalar_blocks.append(current_block)
            current_block = [f]
            
    if current_block:
        scalar_blocks.append(current_block)

    # Construct main rendering blocks (scalars + multi-line)
    main_blocks = [("scalar_block", block) for block in scalar_blocks]
    
    # Separate multi-line fields into ordered and unordered
    ordered_ml = []
    unordered_ml = []
    
    for f in visible_multi_line:
        field_config = configuration.get("multi_line_fields", {}).get(f.key, {})
        sort_order = field_config.get("sort_order")
        try:
            if sort_order is not None and str(sort_order).strip():
                sort_order = int(sort_order)
                ordered_ml.append((sort_order, f))
            else:
                unordered_ml.append(f)
        except (ValueError, TypeError):
            unordered_ml.append(f)
            
    # Sort ordered_ml by their sort_order (stable sort)
    ordered_ml.sort(key=lambda x: x[0])
    
    # Interleave ordered multi-line fields into main_blocks (1-indexed position)
    for sort_order, f in ordered_ml:
        idx = max(0, sort_order - 1)
        if idx >= len(main_blocks):
            main_blocks.append(("multi_line", f))
        else:
            main_blocks.insert(idx, ("multi_line", f))
            
    # Append unordered multi-line fields to the end of main_blocks
    for f in unordered_ml:
        main_blocks.append(("multi_line", f))

    # Calculate unified serial numbers for all visible elements in main_blocks
    field_serial_numbers = {}
    top_counter = 0
    
    for item_type, val in main_blocks:
        if item_type == "scalar_block":
            parent_field = val[0]
            top_counter += 1
            field_serial_numbers[parent_field.id] = str(top_counter)
            
            parent_child_counters = {}
            for child in val[1:]:
                parent_id = child.config.get("condition_trigger_field_id") if child.config else None
                if parent_id is not None:
                    try:
                        parent_id = int(parent_id)
                    except (ValueError, TypeError):
                        parent_id = None
                
                if parent_id is not None:
                    parent_num = field_serial_numbers.get(parent_id, "")
                    child_idx = parent_child_counters.get(parent_id, 0) + 1
                    parent_child_counters[parent_id] = child_idx
                    field_serial_numbers[child.id] = f"{parent_num}.{child_idx}" if parent_num else f"{child_idx}"
                else:
                    field_serial_numbers[child.id] = f"{top_counter}.x"
        elif item_type == "multi_line":
            top_counter += 1
            field_serial_numbers[val.id] = str(top_counter)

    # 6. Render main blocks (interleaved scalars & tables)
    custom_labels = configuration.get("scalar_fields", {})
    
    for item_type, val in main_blocks:
        if item_type == "scalar_block":
            scalar_data = []
            for f in val:
                val_data = values_by_field_id.get(f.id)
                # Check for boolean or has options and render checkboxes
                if f.field_type == FieldType.boolean:
                    val_bool = None
                    if val_data is not None:
                        if isinstance(val_data, bool):
                            val_bool = val_data
                        elif isinstance(val_data, str):
                            val_bool = val_data.strip().lower() in ("1", "true", "yes", "y")
                        elif isinstance(val_data, (int, float)):
                            val_bool = bool(val_data)
                    if val_bool is True:
                        val_str = "Yes"
                    elif val_bool is False:
                        val_str = "No"
                    else:
                        val_str = "—"
                elif f.options:
                    opts_str = []
                    selected_val = str(val_data).strip() if val_data is not None else ""
                    for opt in f.options:
                        box = "☑" if opt.value == selected_val else "☐"
                        opts_str.append(f"{box} {html.escape(opt.label or opt.value)}")
                    val_str = " &nbsp;&nbsp;&nbsp;&nbsp; ".join(opts_str)
                else:
                    val_str = _display_string_for_pdf_export(val_data, f.field_type) if val_data is not None else "—"

                # calculate depth
                depth = 0
                curr = f
                while curr.config and curr.config.get("condition_trigger_field_id") is not None:
                    try:
                        parent_id = int(curr.config.get("condition_trigger_field_id"))
                        parent = fields_by_id.get(parent_id)
                        if not parent:
                            break
                        depth += 1
                        curr = parent
                    except (ValueError, TypeError):
                        break
                    if depth > 10:
                        break

                display_label = custom_labels.get(f.key) or custom_labels.get(str(f.id)) or f.name
                serial_prefix = field_serial_numbers.get(f.id, "")
                if serial_prefix:
                    display_label = f"{serial_prefix}. {display_label}"
                    
                indented_label_style = ParagraphStyle(
                    f"KpiBodyLabel_{f.id}",
                    parent=body_label_style,
                    leftIndent=depth * 15
                )
                
                scalar_data.append([
                    Paragraph(html.escape(display_label), indented_label_style),
                    Paragraph(val_str, body_value_style)
                ])

            scalar_table = Table(scalar_data, colWidths=[384, 120])
            scalar_table.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.HexColor("#f3f4f6")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]))
            elements.append(scalar_table)
            elements.append(Spacer(1, 10))
            
        elif item_type == "multi_line":
            f = val
            field_config = configuration.get("multi_line_fields", {}).get(f.key, {})
            selected_column_keys = field_config.get("selected_columns")
            if not selected_column_keys:
                selected_column_keys = [sf.key for sf in f.sub_fields]

            raw_filters = field_config.get("filters", {})

            from app.entries.multi_line_load import load_multi_line_row_dicts
            row_pairs = await load_multi_line_row_dicts(db, entry_id=entry.id, field=f)
            rows = [r for _, r in row_pairs]

            filtered_rows = []
            if rows:
                from app.entries.multi_item_filters import row_passes_filters
                from app.entries.reference_filter_resolve import build_reference_resolution_map

                if raw_filters and raw_filters.get("conditions"):
                    conds = raw_filters.get("conditions")
                    resolution_maps = await build_reference_resolution_map(
                        db, organization_id, entry.year, f, conds, rows
                    )
                    reference_field_types = {sf.key: sf.field_type.value if hasattr(sf.field_type, "value") else sf.field_type for sf in f.sub_fields}
                    for r in rows:
                        if row_passes_filters(r, raw_filters, resolution_maps=resolution_maps, reference_field_types=reference_field_types):
                            filtered_rows.append(r)
                else:
                    filtered_rows = rows

            table_name = (field_config.get("table_name") or "").strip() or f.name
            table_heading = (field_config.get("table_heading") or "").strip()
            table_subheader = (field_config.get("table_subheader") or "").strip()

            is_duplicate_heading = (table_heading == table_name)

            serial_prefix = field_serial_numbers.get(f.id, "")
            if serial_prefix:
                table_name = f"{serial_prefix}. {table_name}"

            elements.append(Spacer(1, 10))
            elements.append(Paragraph(html.escape(table_name), table_title_style))
            if table_heading and not is_duplicate_heading:
                elements.append(Paragraph(html.escape(table_heading), table_heading_style))
            if table_subheader:
                elements.append(Paragraph(html.escape(table_subheader), table_subheader_style))

            key_to_sf = {sf.key: sf for sf in f.sub_fields}
            headers = [str(key_to_sf[col].name if col in key_to_sf else col) for col in selected_column_keys]

            num_cols = len(selected_column_keys)
            max_char_limit = None
            if num_cols > 30:
                table_font_size = 4.5
                table_leading = 5.5
                min_col_width = 20
                max_char_limit = 12
            elif num_cols > 15:
                table_font_size = 5.5
                table_leading = 6.5
                min_col_width = 25
                max_char_limit = 25
            elif num_cols > 8:
                table_font_size = 7
                table_leading = 9
                min_col_width = 35
                max_char_limit = 40
            else:
                table_font_size = 8
                table_leading = 10
                min_col_width = 45

            t_th_style = ParagraphStyle(
                f"KpiTH_{f.key}",
                parent=styles["Normal"],
                fontName="Helvetica-Bold",
                fontSize=table_font_size,
                leading=table_leading,
                textColor=colors.white
            )
            t_th_sr_style = ParagraphStyle(f"KpiTH_Sr_{f.key}", parent=t_th_style, alignment=TA_CENTER)
            
            t_td_style = ParagraphStyle(
                f"KpiTD_{f.key}",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=table_font_size,
                leading=table_leading,
                textColor=colors.HexColor("#1f2937")
            )
            t_td_sr_style = ParagraphStyle(f"KpiTD_Sr_{f.key}", parent=t_td_style, alignment=TA_CENTER)

            available_width = 504
            sr_no_width = 24

            sample_rows = filtered_rows[:200]
            col_chars = [len(h) for h in headers]
            for r in sample_rows:
                for idx, col in enumerate(selected_column_keys):
                    sf = key_to_sf.get(col)
                    if sf and not _is_subfield_satisfied_for_row(sf, r, key_to_sf):
                        raw_cell = ""
                    else:
                        ft = getattr(sf, "field_type", None) if sf else None
                        raw_cell = _display_string_for_pdf_export(r.get(col), ft) if r.get(col) is not None else ""
                    col_chars[idx] = max(col_chars[idx], min(len(raw_cell), 60))

            total_chars = sum(col_chars) or len(selected_column_keys)
            raw_widths = [max(min_col_width, (available_width - sr_no_width) * (c / total_chars)) for c in col_chars]
            scale = (available_width - sr_no_width) / sum(raw_widths) if sum(raw_widths) > (available_width - sr_no_width) else 1.0
            
            col_widths = [sr_no_width] + [max(min_col_width, w * scale) for w in raw_widths]

            def get_header_text(h: str) -> str:
                if max_char_limit and len(h) > max_char_limit:
                    return h[:max_char_limit] + "..."
                return h

            table_data = [
                [Paragraph("<b>Sr. No.</b>", t_th_sr_style)] + 
                [Paragraph(f"<b>{html.escape(get_header_text(h))}</b>", t_th_style) for h in headers]
            ]

            def get_cell_text(row: dict, col: str) -> str:
                sf = key_to_sf.get(col)
                if sf and not _is_subfield_satisfied_for_row(sf, row, key_to_sf):
                    return ""
                ft = getattr(sf, "field_type", None) if sf else None
                raw_cell = _display_string_for_pdf_export(row.get(col) if isinstance(row, dict) else None, ft)
                
                limit = max_char_limit or 1000
                if len(raw_cell) > limit:
                    raw_cell = raw_cell[:limit] + "..."
                    
                lines = raw_cell.split("\n")
                if len(lines) > 15:
                    raw_cell = "\n".join(lines[:15]) + "\n[... truncated due to size limit]"
                    
                return html.escape(raw_cell).replace("\n", "<br/>")

            for sr, r in enumerate(filtered_rows, start=1):
                table_data.append([
                    Paragraph(str(sr), t_td_sr_style)
                ] + [
                    Paragraph(get_cell_text(r, col), t_td_style) for col in selected_column_keys
                ])

            grid_style = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a8a")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]

            if not filtered_rows:
                table_data.append([
                    Paragraph("No records found matching filters.", t_td_style)
                ] + [Paragraph("", t_td_style) for _ in selected_column_keys])
                grid_style.append(("SPAN", (0, 1), (-1, 1)))
            else:
                grid_style.append(("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f3f4f6")]))

            ml_table = Table(table_data, colWidths=col_widths, repeatRows=1)
            ml_table.setStyle(TableStyle(grid_style))
            elements.append(ml_table)
            elements.append(Spacer(1, 10))

    # 8. Build Document
    buf = BytesIO()
    margin = 54
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
        title=title
    )

    def canvas_maker(*args, **kwargs):
        canvas_obj = NumberedCanvas(*args, **kwargs)
        canvas_obj.organization_name = organization_info or (getattr(org, "name", "") if org else "")
        canvas_obj.include_date = include_generation_date
        return canvas_obj

    doc.build(elements, canvasmaker=canvas_maker)
    return buf.getvalue()


async def generate_kpi_docx_report(
    db: AsyncSession,
    organization_id: int,
    kpi_id: int,
    year: int,
    period_key: str,
    configuration: dict,
    requesting_user_id: int | None = None
) -> bytes:
    import html
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    from app.core.models import KPI, KPIField, KPIEntry, KPIFieldValue, Organization, FieldType
    from app.entries.routes import _display_string_for_pdf_export
    from sqlalchemy.orm import selectinload
    
    # 1. Fetch KPI
    kpi_stmt = (
        select(KPI)
        .where(KPI.id == kpi_id, KPI.organization_id == organization_id)
        .options(selectinload(KPI.fields).selectinload(KPIField.sub_fields))
    )
    kpi_res = await db.execute(kpi_stmt)
    kpi = kpi_res.scalar_one_or_none()
    if not kpi:
        raise ValueError("KPI not found")

    # 2. Fetch Entry
    entry = None
    if requesting_user_id is not None:
        entry_stmt = select(KPIEntry).where(
            KPIEntry.organization_id == organization_id,
            KPIEntry.kpi_id == kpi_id,
            KPIEntry.year == year,
            KPIEntry.period_key == period_key,
            KPIEntry.is_draft == True,
            KPIEntry.user_id == requesting_user_id
        )
        entry_res = await db.execute(entry_stmt)
        entry = entry_res.scalar_one_or_none()
        
    if not entry:
        entry_stmt = select(KPIEntry).where(
            KPIEntry.organization_id == organization_id,
            KPIEntry.kpi_id == kpi_id,
            KPIEntry.year == year,
            KPIEntry.period_key == period_key,
            KPIEntry.is_draft == False
        )
        entry_res = await db.execute(entry_stmt)
        entry = entry_res.scalar_one_or_none()
        
    if not entry:
        raise ValueError(f"KPI Entry for year {year} and period '{period_key}' not found")

    # 3. Fetch Scalar Values
    scalar_values_query = select(KPIFieldValue).where(KPIFieldValue.entry_id == entry.id)
    scalar_values_res = await db.execute(scalar_values_query)
    scalar_values = scalar_values_res.scalars().all()
    
    values_by_field_id = {}
    for val in scalar_values:
        if val.value_boolean is not None:
            values_by_field_id[val.field_id] = val.value_boolean
        elif val.value_number is not None:
            values_by_field_id[val.field_id] = val.value_number
        elif val.value_date is not None:
            values_by_field_id[val.field_id] = val.value_date
        elif val.value_json is not None:
            values_by_field_id[val.field_id] = val.value_json
        else:
            values_by_field_id[val.field_id] = val.value_text

    # 4. Extract Configuration
    title = (configuration.get("report_header") or configuration.get("title") or "").strip() or f"{kpi.name} Report"
    description = (configuration.get("description") or "").strip()

    org = await db.get(Organization, organization_id)

    # 5. Initialize Document
    doc = Document()
    
    # Add Report Header/Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    
    # Add Description
    if description:
        p_desc_head = doc.add_paragraph()
        p_desc_head.paragraph_format.space_before = Pt(12)
        p_desc_head.paragraph_format.space_after = Pt(4)
        run_dh = p_desc_head.add_run("Description")
        run_dh.bold = True
        run_dh.font.size = Pt(12)
        run_dh.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
        
        p_desc_text = doc.add_paragraph()
        p_desc_text.paragraph_format.space_after = Pt(12)
        run_dt = p_desc_text.add_run(description)
        run_dt.font.size = Pt(10)
        run_dt.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # 6. Extract fields logic (identical to PDF)
    scalar_fields = [f for f in kpi.fields if f.field_type != FieldType.multi_line_items]
    fields_by_id = {f.id: f for f in kpi.fields}
    sorted_scalars = group_dependent_fields(scalar_fields)
    
    excluded_fields = configuration.get("excluded_scalar_fields") or []
    excluded_set = {str(item) for item in excluded_fields}
    
    visible_scalars = [
        f for f in sorted_scalars 
        if is_field_visible(f, fields_by_id, values_by_field_id)
        and f.key not in excluded_set
        and str(f.id) not in excluded_set
    ]

    # Apply custom scalar fields order if specified
    ordered_scalar_ids = configuration.get("ordered_scalar_fields")
    if ordered_scalar_ids:
        order_map = {str(val): idx for idx, val in enumerate(ordered_scalar_ids)}
        def get_sort_key(f):
            if str(f.id) in order_map:
                return order_map[str(f.id)]
            if f.key in order_map:
                return order_map[f.key]
            return 999999
        visible_scalars.sort(key=get_sort_key)

    # Get excluded multi-line fields and compute visible ones
    multi_line_fields = [f for f in kpi.fields if f.field_type == FieldType.multi_line_items]
    excluded_ml_keys = configuration.get("excluded_multi_line_fields") or []
    excluded_ml_set = {str(item) for item in excluded_ml_keys}
    visible_multi_line = [
        f for f in multi_line_fields
        if is_field_visible(f, fields_by_id, values_by_field_id)
        and f.key not in excluded_ml_set and str(f.id) not in excluded_ml_set
    ]

    # Build parent blocks for scalar fields to keep parents and their children grouped
    scalar_blocks = []
    current_block = []
    
    for f in visible_scalars:
        parent_id = f.config.get("condition_trigger_field_id") if f.config else None
        if parent_id is not None:
            try:
                parent_id = int(parent_id)
            except (ValueError, TypeError):
                parent_id = None
                
        is_child = False
        if parent_id is not None:
            parent_field = fields_by_id.get(parent_id)
            if parent_field and parent_field in visible_scalars:
                is_child = True
                
        if is_child:
            if current_block:
                current_block.append(f)
            else:
                current_block = [f]
        else:
            if current_block:
                scalar_blocks.append(current_block)
            current_block = [f]
            
    if current_block:
        scalar_blocks.append(current_block)

    # Construct main rendering blocks (scalars + multi-line)
    main_blocks = [("scalar_block", block) for block in scalar_blocks]
    
    # Separate multi-line fields into ordered and unordered
    ordered_ml = []
    unordered_ml = []
    
    for f in visible_multi_line:
        field_config = configuration.get("multi_line_fields", {}).get(f.key, {})
        sort_order = field_config.get("sort_order")
        try:
            if sort_order is not None and str(sort_order).strip():
                sort_order = int(sort_order)
                ordered_ml.append((sort_order, f))
            else:
                unordered_ml.append(f)
        except (ValueError, TypeError):
            unordered_ml.append(f)
            
    # Sort ordered_ml by their sort_order (stable sort)
    ordered_ml.sort(key=lambda x: x[0])
    
    # Interleave ordered multi-line fields into main_blocks (1-indexed position)
    for sort_order, f in ordered_ml:
        idx = max(0, sort_order - 1)
        if idx >= len(main_blocks):
            main_blocks.append(("multi_line", f))
        else:
            main_blocks.insert(idx, ("multi_line", f))
            
    # Append unordered multi-line fields to the end of main_blocks
    for f in unordered_ml:
        main_blocks.append(("multi_line", f))

    # Calculate unified serial numbers for all visible elements in main_blocks
    field_serial_numbers = {}
    top_counter = 0
    
    for item_type, val in main_blocks:
        if item_type == "scalar_block":
            parent_field = val[0]
            top_counter += 1
            field_serial_numbers[parent_field.id] = str(top_counter)
            
            parent_child_counters = {}
            for child in val[1:]:
                parent_id = child.config.get("condition_trigger_field_id") if child.config else None
                if parent_id is not None:
                    try:
                        parent_id = int(parent_id)
                    except (ValueError, TypeError):
                        parent_id = None
                
                if parent_id is not None:
                    parent_num = field_serial_numbers.get(parent_id, "")
                    child_idx = parent_child_counters.get(parent_id, 0) + 1
                    parent_child_counters[parent_id] = child_idx
                    field_serial_numbers[child.id] = f"{parent_num}.{child_idx}" if parent_num else f"{child_idx}"
                else:
                    field_serial_numbers[child.id] = f"{top_counter}.x"
        elif item_type == "multi_line":
            top_counter += 1
            field_serial_numbers[val.id] = str(top_counter)

    # 7. Render main blocks (interleaved scalars & tables)
    custom_labels = configuration.get("scalar_fields", {})
    
    for item_type, val in main_blocks:
        if item_type == "scalar_block":
            for f in val:
                val_data = values_by_field_id.get(f.id)
                # Check for boolean or has options and render plain values
                if f.field_type == FieldType.boolean:
                    val_bool = None
                    if val_data is not None:
                        if isinstance(val_data, bool):
                            val_bool = val_data
                        elif isinstance(val_data, str):
                            val_bool = val_data.strip().lower() in ("1", "true", "yes", "y")
                        elif isinstance(val_data, (int, float)):
                            val_bool = bool(val_data)
                    box_yes = "☒" if val_bool is True else "☐"
                    box_no = "☒" if val_bool is False else "☐"
                    val_str = f"{box_yes} Yes    {box_no} No"
                elif f.options:
                    opts_str = []
                    selected_val = str(val_data).strip() if val_data is not None else ""
                    for opt in f.options:
                        box = "☒" if opt.value == selected_val else "☐"
                        opts_str.append(f"{box} {opt.label or opt.value}")
                    val_str = "    ".join(opts_str)
                else:
                    val_str = _display_string_for_pdf_export(val_data, f.field_type) if val_data is not None else "—"

                # calculate depth
                depth = 0
                curr = f
                while curr.config and curr.config.get("condition_trigger_field_id") is not None:
                    try:
                        parent_id = int(curr.config.get("condition_trigger_field_id"))
                        parent = fields_by_id.get(parent_id)
                        if not parent:
                            break
                        depth += 1
                        curr = parent
                    except (ValueError, TypeError):
                        break
                    if depth > 10:
                        break

                display_label = custom_labels.get(f.key) or custom_labels.get(str(f.id)) or f.name
                serial_prefix = field_serial_numbers.get(f.id, "")
                if serial_prefix:
                    display_label = f"{serial_prefix}. {display_label}"
                
                # Check ending punctuation for suffix
                suffix = ": "
                if display_label.endswith("?") or display_label.endswith(":") or display_label.endswith("."):
                    suffix = " "

                # Add as styled paragraph in parallel layout flow
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 * depth)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                
                run_lbl = p.add_run(f"{display_label}{suffix}")
                run_lbl.bold = True
                run_lbl.font.size = Pt(10.5)
                run_lbl.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                
                run_val = p.add_run(val_str)
                run_val.font.size = Pt(10.5)
                run_val.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
            
        elif item_type == "multi_line":
            f = val
            field_config = configuration.get("multi_line_fields", {}).get(f.key, {})
            selected_column_keys = field_config.get("selected_columns")
            if not selected_column_keys:
                selected_column_keys = [sf.key for sf in f.sub_fields]

            raw_filters = field_config.get("filters", {})

            from app.entries.multi_line_load import load_multi_line_row_dicts
            row_pairs = await load_multi_line_row_dicts(db, entry_id=entry.id, field=f)
            rows = [r for _, r in row_pairs]

            filtered_rows = []
            if rows:
                from app.entries.multi_item_filters import row_passes_filters
                from app.entries.reference_filter_resolve import build_reference_resolution_map

                if raw_filters and raw_filters.get("conditions"):
                    conds = raw_filters.get("conditions")
                    resolution_maps = await build_reference_resolution_map(
                        db, organization_id, entry.year, f, conds, rows
                    )
                    reference_field_types = {sf.key: sf.field_type.value if hasattr(sf.field_type, "value") else sf.field_type for sf in f.sub_fields}
                    for r in rows:
                        if row_passes_filters(r, raw_filters, resolution_maps=resolution_maps, reference_field_types=reference_field_types):
                            filtered_rows.append(r)
                else:
                    filtered_rows = rows

            table_name = (field_config.get("table_name") or "").strip() or f.name
            serial_prefix = field_serial_numbers.get(f.id, "")
            if serial_prefix:
                table_name = f"{serial_prefix}. {table_name}"
                
            table_heading = (field_config.get("table_heading") or "").strip()
            table_subheader = (field_config.get("table_subheader") or "").strip()

            # Render Table title/headings in Word
            p_tname = doc.add_paragraph()
            p_tname.paragraph_format.space_before = Pt(12)
            p_tname.paragraph_format.space_after = Pt(2)
            run_tn = p_tname.add_run(table_name)
            run_tn.bold = True
            run_tn.font.size = Pt(12)
            run_tn.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
            
            # Deduplicate check
            is_duplicate_heading = (table_heading == (field_config.get("table_name") or "").strip() or f.name)
            if table_heading and not is_duplicate_heading:
                p_thead = doc.add_paragraph()
                p_thead.paragraph_format.space_after = Pt(2)
                run_th = p_thead.add_run(table_heading)
                run_th.bold = True
                run_th.font.size = Pt(10)
                run_th.font.color.rgb = RGBColor(0x4b, 0x55, 0x63)
                
            if table_subheader:
                p_tsub = doc.add_paragraph()
                p_tsub.paragraph_format.space_after = Pt(6)
                run_ts = p_tsub.add_run(table_subheader)
                run_ts.italic = True
                run_ts.font.size = Pt(9)
                run_ts.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

            key_to_sf = {sf.key: sf for sf in f.sub_fields}
            headers = [str(key_to_sf[col].name if col in key_to_sf else col) for col in selected_column_keys]

            # Add Table Grid in Word
            table = doc.add_table(rows=1, cols=len(selected_column_keys) + 1)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].paragraphs[0].text = "Sr. No."
            hdr_cells[0].paragraphs[0].runs[0].bold = True
            hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
            
            for i, h_text in enumerate(headers):
                hdr_cells[i+1].paragraphs[0].text = h_text
                hdr_cells[i+1].paragraphs[0].runs[0].bold = True
                hdr_cells[i+1].paragraphs[0].runs[0].font.size = Pt(9.5)

            num_cols = len(selected_column_keys)
            max_char_limit = None
            if num_cols > 30:
                max_char_limit = 12
            elif num_cols > 15:
                max_char_limit = 25
            elif num_cols > 8:
                max_char_limit = 40

            def get_cell_plain_text(row_dict: dict, col: str) -> str:
                sf = key_to_sf.get(col)
                if sf and not _is_subfield_satisfied_for_row(sf, row_dict, key_to_sf):
                    return ""
                ft = getattr(sf, "field_type", None) if sf else None
                raw_cell = _display_string_for_pdf_export(row_dict.get(col) if isinstance(row_dict, dict) else None, ft)
                
                limit = max_char_limit or 1000
                if len(raw_cell) > limit:
                    raw_cell = raw_cell[:limit] + "..."
                    
                lines = raw_cell.split("\n")
                if len(lines) > 15:
                    raw_cell = "\n".join(lines[:15]) + "\n[... truncated]"
                    
                return raw_cell

            if not filtered_rows:
                row = table.add_row()
                row.cells[0].paragraphs[0].text = "No records found matching filters."
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                for cell in row.cells[1:]:
                    row.cells[0].merge(cell)
            else:
                for sr, r in enumerate(filtered_rows, start=1):
                    row = table.add_row()
                    row.cells[0].paragraphs[0].text = str(sr)
                    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                    for i, col in enumerate(selected_column_keys):
                        row.cells[i+1].paragraphs[0].text = get_cell_plain_text(r, col)
                        row.cells[i+1].paragraphs[0].runs[0].font.size = Pt(9)

            # Spacing
            doc.add_paragraph()

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


async def background_generate_kpi_pdf(
    job_id: str,
    organization_id: int,
    kpi_id: int,
    year: int,
    period_key: str,
    configuration: dict
) -> None:
    from app.core.database import AsyncSessionLocal
    from app.core.models import KpiReportJob, utc_now
    from app.storage.service import upload_file
    import traceback
    import uuid

    async with AsyncSessionLocal() as db:
        try:
            job_stmt = select(KpiReportJob).where(KpiReportJob.id == job_id)
            job_res = await db.execute(job_stmt)
            job = job_res.scalar_one_or_none()
            if not job:
                return

            job.status = "processing"
            await db.commit()

            fmt = configuration.get("format", "pdf")
            if fmt == "docx":
                doc_bytes = await generate_kpi_docx_report(
                    db=db,
                    organization_id=organization_id,
                    kpi_id=kpi_id,
                    year=year,
                    period_key=period_key,
                    configuration=configuration,
                    requesting_user_id=job.user_id
                )
                filename = f"kpi_report_{kpi_id}_{year}_{period_key or 'full'}_{uuid.uuid4().hex[:8]}.docx"
                relative_path = f"kpi_reports/{filename}"
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_bytes = doc_bytes
            else:
                pdf_bytes = await generate_kpi_pdf_report(
                    db=db,
                    organization_id=organization_id,
                    kpi_id=kpi_id,
                    year=year,
                    period_key=period_key,
                    configuration=configuration,
                    requesting_user_id=job.user_id
                )
                filename = f"kpi_report_{kpi_id}_{year}_{period_key or 'full'}_{uuid.uuid4().hex[:8]}.pdf"
                relative_path = f"kpi_reports/{filename}"
                content_type = "application/pdf"
                file_bytes = pdf_bytes

            stored_path = await upload_file(
                db=db,
                organization_id=organization_id,
                relative_path=relative_path,
                content=file_bytes,
                content_type=content_type
            )

            # Reload and complete
            job_res = await db.execute(job_stmt)
            job = job_res.scalar_one_or_none()
            if job:
                job.status = "completed"
                job.stored_path = stored_path
                job.completed_at = utc_now()
                await db.commit()

        except Exception as e:
            traceback.print_exc()
            try:
                job_res = await db.execute(select(KpiReportJob).where(KpiReportJob.id == job_id))
                job = job_res.scalar_one_or_none()
                if job:
                    job.status = "failed"
                    job.error_message = str(e)
                    job.completed_at = utc_now()
                    await db.commit()
            except Exception:
                pass
